# --- IMPORTS ---
import subprocess
import time
import json
import re
import platform
import os
import threading
import webbrowser
import concurrent.futures
from flask import Flask, jsonify, render_template, abort, request, send_file
import requests
import dns.resolver  # Para Benchmark DNS
from waitress import serve # Servidor de producción
import urllib3 # Para deshabilitar warnings SSL
import logging # Para logging más detallado si es necesario
from datetime import datetime
import socket
from statistics import mean, median, stdev
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- IMPORTS PARA PRUEBAS AVANZADAS ---
try:
    from scapy.all import *
    from scapy.layers.dhcp import DHCP, BOOTP
    from scapy.layers.inet import IP, UDP
    from scapy.layers.l2 import Ether
    SCAPY_AVAILABLE = True
except ImportError:
    SCAPY_AVAILABLE = False
    print("ADVERTENCIA: Scapy no está disponible. Las pruebas DHCP no funcionarán.")

try:
    import netifaces
    NETIFACES_AVAILABLE = True
except ImportError:
    NETIFACES_AVAILABLE = False
    print("ADVERTENCIA: netifaces no está disponible. Algunas funciones de red pueden no funcionar.")

import ipaddress

# --- DESHABILITAR ADVERTENCIAS DE SSL ---
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# --- CONFIGURACIÓN GLOBAL ---
BACKUP_DIR = "data_backups"
JSON_OUTPUT_FILE = "ping_results.json"
SPEEDTEST_RESULTS_FILE = "speedtest_results.json"
TRACEROUTE_HISTORY_FILE = "traceroute_history.json"
WEBSITES_FILE = "websites.txt"
REPORTS_DIR = "reports"
MAX_HISTORY_PER_SITE = 500
PING_INTERVAL_SECONDS = 10
SAVE_INTERVAL_SECONDS = 60 # Guardar datos solo cada 60s
FLASK_PORT = 5000
MAX_PING_WORKERS = 50

# Crear directorio de reportes si no existe
if not os.path.exists(REPORTS_DIR):
    os.makedirs(REPORTS_DIR)

# --- LÓGICA COMPARTIDA Y DE ESTADO ---
ping_results_data = {}
speedtest_results_data = {"latest": None, "history": []}
dns_benchmark_results = {"status": "idle", "results": [], "last_run": None}
traceroute_history = []
data_lock = threading.Lock()
speedtest_lock = threading.Lock()
dns_benchmark_lock = threading.Lock()
traceroute_history_lock = threading.Lock()
websites_lock = threading.Lock()
first_ping_round_done = threading.Event()

# --- Controles para el Speedtest Automático ---
g_speedtest_interval_minutes = 30
auto_speedtest_running_event = threading.Event()
ping_paused_event = threading.Event()  # Para pausar pings durante speedtest

# --- ESTADO PARA PRUEBAS AVANZADAS ---
advanced_tests_results = {
    "dhcp_discovery": {"status": "idle", "results": [], "last_run": None},
    "network_discovery": {"status": "idle", "results": [], "last_run": None},
    "mtr_analysis": {"status": "idle", "results": {}, "last_run": None},
    "netflix_speed": {"status": "idle", "results": {}, "last_run": None},
    "udp_jitter": {"status": "idle", "results": {}, "last_run": None},
    "cdn_dns": {"status": "idle", "results": {}, "last_run": None},
    "mtu_verification": {"status": "idle", "results": {}, "last_run": None},
    "sustained_load": {"status": "idle", "results": {}, "last_run": None}
}
advanced_tests_lock = threading.Lock()

# --- FUNCIONES AUXILIARES PARA REPORTES ---

def check_tool_availability():
    """Verifica la disponibilidad de cada herramienta para las pruebas avanzadas"""
    tools_status = {
        "dhcp_discovery": {
            "available": False,
            "missing_tools": [],
            "install_command": ""
        },
        "network_discovery": {
            "available": False,
            "missing_tools": [],
            "install_command": ""
        },
        "mtr_analysis": {
            "available": False,
            "missing_tools": [],
            "install_command": ""
        },
        "netflix_speed": {
            "available": False,
            "missing_tools": [],
            "install_command": ""
        },
        "udp_jitter": {
            "available": False,
            "missing_tools": [],
            "install_command": ""
        },
        "cdn_dns": {
            "available": True,  # DNS siempre disponible (usa dnspython)
            "missing_tools": [],
            "install_command": ""
        },
        "mtu_verification": {
            "available": True,  # Usa ping nativo del sistema
            "missing_tools": [],
            "install_command": ""
        },
        "sustained_load": {
            "available": True,  # Usa requests (ya instalado)
            "missing_tools": [],
            "install_command": ""
        }
    }

    # 1. DHCP Discovery - Requiere Scapy
    if SCAPY_AVAILABLE:
        tools_status["dhcp_discovery"]["available"] = True
    else:
        tools_status["dhcp_discovery"]["missing_tools"].append("scapy")
        tools_status["dhcp_discovery"]["install_command"] = "pip install scapy"

    # 2. Network Discovery - Requiere nmap
    try:
        result = subprocess.run(["nmap", "--version"], capture_output=True, timeout=5)
        if result.returncode == 0:
            tools_status["network_discovery"]["available"] = True
    except:
        tools_status["network_discovery"]["missing_tools"].append("nmap")
        if platform.system().lower() == "windows":
            tools_status["network_discovery"]["install_command"] = "Descarga desde https://nmap.org/download.html"
        else:
            tools_status["network_discovery"]["install_command"] = "sudo apt-get install nmap (Linux) o brew install nmap (Mac)"

    # 3. PathPing/MTR Analysis - PathPing nativo en Windows, MTR en Linux/Mac
    system = platform.system().lower()
    if system == "windows":
        # Windows tiene pathping por defecto (nativo)
        tools_status["mtr_analysis"]["available"] = True
    else:
        # Linux/Mac: verificar si tienen mtr instalado
        try:
            result = subprocess.run(["mtr", "--version"], capture_output=True, timeout=5)
            if result.returncode == 0:
                tools_status["mtr_analysis"]["available"] = True
            else:
                tools_status["mtr_analysis"]["missing_tools"].append("mtr")
                tools_status["mtr_analysis"]["install_command"] = "sudo apt-get install mtr (Linux) o brew install mtr (Mac)"
        except:
            tools_status["mtr_analysis"]["missing_tools"].append("mtr")
            tools_status["mtr_analysis"]["install_command"] = "sudo apt-get install mtr (Linux) o brew install mtr (Mac)"

    # 4. Netflix Speed - Requiere fast-cli (Node.js)
    try:
        # En Windows necesita shell=True para encontrar comandos npm globales
        result = subprocess.run(["fast", "--version"], capture_output=True, timeout=5, shell=True)
        if result.returncode == 0:
            tools_status["netflix_speed"]["available"] = True
        else:
            tools_status["netflix_speed"]["missing_tools"].append("fast-cli")
            if platform.system().lower() == "windows":
                tools_status["netflix_speed"]["install_command"] = "npm install -g fast-cli (Requiere Node.js instalado desde https://nodejs.org)"
            else:
                tools_status["netflix_speed"]["install_command"] = "npm install -g fast-cli"
    except:
        tools_status["netflix_speed"]["missing_tools"].append("fast-cli")
        if platform.system().lower() == "windows":
            tools_status["netflix_speed"]["install_command"] = "npm install -g fast-cli (Requiere Node.js instalado desde https://nodejs.org)"
        else:
            tools_status["netflix_speed"]["install_command"] = "npm install -g fast-cli"

    # 5. UDP Jitter - Requiere iperf3
    try:
        result = subprocess.run(["iperf3", "--version"], capture_output=True, timeout=5)
        if result.returncode == 0:
            tools_status["udp_jitter"]["available"] = True
    except:
        tools_status["udp_jitter"]["missing_tools"].append("iperf3")
        if platform.system().lower() == "windows":
            tools_status["udp_jitter"]["install_command"] = "Descarga desde https://iperf.fr/iperf-download.php"
        else:
            tools_status["udp_jitter"]["install_command"] = "sudo apt-get install iperf3 (Linux) o brew install iperf3 (Mac)"

    return tools_status

def get_local_ip():
    """Obtiene la IP local del equipo"""
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        local_ip = s.getsockname()[0]
        s.close()
        return local_ip
    except Exception as e:
        return "No disponible"

def get_public_ip():
    """Obtiene la IP pública del equipo"""
    try:
        response = requests.get('https://api.ipify.org?format=json', timeout=5)
        return response.json()['ip']
    except Exception as e:
        try:
            response = requests.get('https://ifconfig.me/ip', timeout=5)
            return response.text.strip()
        except:
            return "No disponible"

def analyze_site_data(site_data):
    """Analiza los datos de un sitio y retorna estadísticas"""
    if not site_data:
        return None
    
    successful_pings = [p for p in site_data if p.get('success', False)]
    failed_pings = [p for p in site_data if not p.get('success', False)]
    
    total_pings = len(site_data)
    successful_count = len(successful_pings)
    failed_count = len(failed_pings)
    
    availability = (successful_count / total_pings * 100) if total_pings > 0 else 0
    
    latencies = [p['time_ms'] for p in successful_pings if 'time_ms' in p]
    
    stats = {
        'total_pings': total_pings,
        'successful': successful_count,
        'failed': failed_count,
        'availability': round(availability, 2),
        'avg_latency': round(mean(latencies), 2) if latencies else None,
        'min_latency': round(min(latencies), 2) if latencies else None,
        'max_latency': round(max(latencies), 2) if latencies else None,
        'median_latency': round(median(latencies), 2) if latencies else None,
        'stdev_latency': round(stdev(latencies), 2) if len(latencies) > 1 else None,
    }
    
    return stats

def classify_site(site):
    """Clasifica un sitio en categorías"""
    if site.startswith('192.168.') or site.startswith('10.') or site.startswith('172.'):
        return 'Red Local'
    elif re.match(r'^\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}$', site):
        return 'IP Pública'
    elif any(keyword in site.lower() for keyword in ['bank', 'banco', 'jep', 'pichincha']):
        return 'Servicios Financieros'
    elif any(keyword in site.lower() for keyword in ['google', 'facebook', 'disney', 'shein']):
        return 'Sitios Externos Populares'
    elif 'dns' in site.lower():
        return 'Servidores DNS'
    else:
        return 'Otros'

def get_status_color(availability, latency):
    """Retorna el color según el estado del sitio"""
    if availability < 95:
        return 'Crítico'
    elif latency and latency > 300:
        return 'Deficiente'
    elif latency and latency > 150:
        return 'Aceptable'
    else:
        return 'Excelente'

def add_table_border(table):
    """Añade bordes a una tabla de Word"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)

# --- FUNCIONES PARA PRUEBAS AVANZADAS ---

def get_network_interfaces():
    """Obtiene lista de interfaces de red disponibles"""
    interfaces = []

    try:
        if platform.system().lower() == "windows":
            # En Windows, usar scapy para obtener interfaces
            if SCAPY_AVAILABLE:
                try:
                    from scapy.arch.windows import get_windows_if_list
                    scapy_interfaces = get_windows_if_list()

                    for iface in scapy_interfaces:
                        if iface.get('name') and iface.get('ips'):
                            # Filtrar interfaces con IP válida
                            valid_ips = [ip for ip in iface.get('ips', []) if not str(ip).startswith('169.254')]
                            if valid_ips:
                                interfaces.append({
                                    'name': iface['name'],
                                    'description': iface.get('description', 'Descripción no disponible'),
                                    'ips': valid_ips,
                                    'display_name': f"{iface.get('description', iface['name'])} ({valid_ips[0]})"
                                })
                except Exception as e:
                    print(f"Error obteniendo interfaces con Scapy: {e}")

            # Fallback usando netifaces
            if not interfaces and NETIFACES_AVAILABLE:
                try:
                    for iface_name in netifaces.interfaces():
                        iface_data = netifaces.ifaddresses(iface_name)
                        if netifaces.AF_INET in iface_data:
                            ipv4_info = iface_data[netifaces.AF_INET][0]
                            ip = ipv4_info.get('addr')
                            if ip and not ip.startswith('169.254') and ip != '127.0.0.1':
                                interfaces.append({
                                    'name': iface_name,
                                    'description': iface_name,
                                    'ips': [ip],
                                    'display_name': f"{iface_name} ({ip})"
                                })
                except Exception as e:
                    print(f"Error obteniendo interfaces con netifaces: {e}")
        else:
            # En Linux/Mac, usar netifaces
            if NETIFACES_AVAILABLE:
                try:
                    for iface_name in netifaces.interfaces():
                        iface_data = netifaces.ifaddresses(iface_name)
                        if netifaces.AF_INET in iface_data:
                            ipv4_info = iface_data[netifaces.AF_INET][0]
                            ip = ipv4_info.get('addr')
                            if ip and not ip.startswith('169.254') and ip != '127.0.0.1':
                                interfaces.append({
                                    'name': iface_name,
                                    'description': iface_name,
                                    'ips': [ip],
                                    'display_name': f"{iface_name} ({ip})"
                                })
                except Exception as e:
                    print(f"Error obteniendo interfaces: {e}")

    except Exception as e:
        print(f"Error general obteniendo interfaces: {e}")

    # Si no se pudieron obtener interfaces, agregar opción por defecto
    if not interfaces:
        interfaces.append({
            'name': 'auto',
            'description': 'Auto-detectar interfaz',
            'ips': ['auto'],
            'display_name': 'Auto-detectar interfaz por defecto'
        })

    return interfaces

def get_installation_instructions(tool_name):
    """Retorna instrucciones de instalación específicas para cada herramienta"""
    instructions = {
        "fast-cli": {
            "description": "Netflix Speed Test Tool",
            "windows": [
                "1. Instalar Node.js desde https://nodejs.org",
                "2. Abrir CMD como administrador",
                "3. Ejecutar: npm install -g fast-cli"
            ],
            "linux": [
                "1. Instalar Node.js: sudo apt install nodejs npm (Ubuntu/Debian) o sudo yum install nodejs npm (CentOS/RHEL)",
                "2. Ejecutar: sudo npm install -g fast-cli"
            ],
            "mac": [
                "1. Instalar Node.js: brew install node",
                "2. Ejecutar: npm install -g fast-cli"
            ]
        },
        "iperf3": {
            "description": "Network Performance Testing Tool",
            "windows": [
                "1. Descargar desde https://iperf.fr/iperf-download.php",
                "2. Extraer y agregar al PATH del sistema",
                "3. O usar: winget install iperf3"
            ],
            "linux": [
                "Ubuntu/Debian: sudo apt install iperf3",
                "CentOS/RHEL: sudo yum install iperf3",
                "Arch: sudo pacman -S iperf3"
            ],
            "mac": [
                "brew install iperf3"
            ]
        },
        "mtr": {
            "description": "Network Diagnostic Tool (My TraceRoute)",
            "windows": [
                "1. Descargar WinMTR desde https://sourceforge.net/projects/winmtr/",
                "2. O usar: winget install WinMTR"
            ],
            "linux": [
                "Ubuntu/Debian: sudo apt install mtr",
                "CentOS/RHEL: sudo yum install mtr",
                "Arch: sudo pacman -S mtr"
            ],
            "mac": [
                "brew install mtr"
            ]
        },
        "nmap": {
            "description": "Network Discovery and Security Auditing Tool",
            "windows": [
                "1. Descargar desde https://nmap.org/download.html",
                "2. Instalar y agregar al PATH",
                "3. O usar: winget install Insecure.Nmap"
            ],
            "linux": [
                "Ubuntu/Debian: sudo apt install nmap",
                "CentOS/RHEL: sudo yum install nmap",
                "Arch: sudo pacman -S nmap"
            ],
            "mac": [
                "brew install nmap"
            ]
        }
    }

    return instructions.get(tool_name, {"description": "Herramienta desconocida", "windows": [], "linux": [], "mac": []})

def format_installation_error(tool_name):
    """Formatea un mensaje de error detallado con instrucciones de instalación"""
    instructions = get_installation_instructions(tool_name)
    system = platform.system().lower()

    system_map = {
        "windows": "windows",
        "linux": "linux",
        "darwin": "mac"
    }

    system_key = system_map.get(system, "linux")
    system_instructions = instructions[system_key]

    error_msg = f"{tool_name} no está instalado.\n\n"
    error_msg += f"📋 {instructions['description']}\n\n"
    error_msg += f"🔧 Instrucciones para {system.title()}:\n"

    for i, instruction in enumerate(system_instructions, 1):
        error_msg += f"   {i}. {instruction}\n"

    error_msg += f"\n💡 También puedes ejecutar el script de instalación: "
    if system == "windows":
        error_msg += "install_dependencies.bat"
    else:
        error_msg += "./install_dependencies.sh"

    return error_msg

def discover_dhcp_servers(interface_name=None):
    """Descubre servidores DHCP en la red local usando Scapy

    Args:
        interface_name: Nombre de la interfaz a usar. Si es None o 'auto', se auto-detecta
    """
    if not SCAPY_AVAILABLE:
        return {"success": False, "error": "Scapy no está disponible. Instala con: pip install scapy"}

    try:
        # Determinar la interfaz a usar
        interface = None

        # Si se especificó una interfaz y no es 'auto', usarla
        if interface_name and interface_name != 'auto':
            interface = interface_name
        else:
            # Auto-detectar interfaz
            # En Windows, intentar obtener la interfaz correcta
            if platform.system().lower() == "windows":
                try:
                    # Usar scapy para obtener interfaces disponibles
                    from scapy.arch.windows import get_windows_if_list
                    interfaces = get_windows_if_list()

                    # Buscar interfaz activa con IP
                    local_ip = get_local_ip()
                    if local_ip != "No disponible":
                        for iface in interfaces:
                            if 'ips' in iface and any(local_ip in str(ip) for ip in iface.get('ips', [])):
                                interface = iface['name']
                                break

                        # Si no se encuentra, usar la primera interfaz activa
                        if not interface:
                            for iface in interfaces:
                                if iface.get('ips') and len(iface['ips']) > 0:
                                    interface = iface['name']
                                    break
                except:
                    # Fallback: usar None para que Scapy auto-detecte
                    interface = None
            else:
                # En Linux/Mac, usar netifaces
                try:
                    gateways = netifaces.gateways()
                    default_gateway = gateways['default'][netifaces.AF_INET]
                    interface = default_gateway[1]
                except:
                    interface = None

        # Crear paquete DHCP Discover
        dhcp_discover = (
            Ether(dst="ff:ff:ff:ff:ff:ff") /
            IP(src="0.0.0.0", dst="255.255.255.255") /
            UDP(sport=68, dport=67) /
            BOOTP(chaddr=RandString(12, "0123456789abcdef")) /
            DHCP(options=[("message-type", "discover"), ("end")])
        )

        # Enviar paquete y esperar respuestas
        print(f"Enviando DHCP Discover en interfaz: {interface or 'auto-detectada'}...")
        start_time = time.time()

        try:
            if interface:
                responses = srp(dhcp_discover, timeout=10, verbose=False, iface=interface)[0]
            else:
                # Dejar que Scapy auto-detecte la interfaz
                responses = srp(dhcp_discover, timeout=10, verbose=False)[0]
        except Exception as send_error:
            # Error al enviar el paquete
            error_detail = str(send_error)
            if "WinPcap" in error_detail or "Npcap" in error_detail or "winpcap" in error_detail.lower():
                return {
                    "success": False,
                    "error": "Falta el driver de captura de paquetes (Npcap/WinPcap)",
                    "help": "Instala Npcap desde https://npcap.com/ - Es necesario para capturar paquetes de red en Windows",
                    "details": error_detail
                }
            elif "permission" in error_detail.lower() or "access" in error_detail.lower():
                return {
                    "success": False,
                    "error": "No tienes permisos suficientes para capturar paquetes de red",
                    "help": "Ejecuta el programa como Administrador (click derecho → Ejecutar como administrador)",
                    "details": error_detail
                }
            else:
                return {
                    "success": False,
                    "error": f"Error al enviar paquete DHCP: {error_detail}",
                    "help": "Verifica que tu interfaz de red esté activa y correctamente configurada"
                }

        elapsed_time = time.time() - start_time
        print(f"Respuestas recibidas: {len(responses)} en {elapsed_time:.2f}s")

        dhcp_servers = []
        for sent, received in responses:
            # Verificar si es un paquete DHCP Offer o DHCP Ack
            if received.haslayer(DHCP):
                dhcp_options = received[DHCP].options
                message_type = None

                # Buscar el tipo de mensaje DHCP
                for opt in dhcp_options:
                    if isinstance(opt, tuple) and opt[0] == 'message-type':
                        message_type = opt[1]
                        break

                # Solo procesar DHCP Offer (2) o DHCP Ack (5)
                if message_type in [2, 5]:
                    server_ip = received[IP].src
                    server_mac = received[Ether].src

                    # Extraer IP ofrecida del paquete DHCP
                    offered_ip = "N/A"
                    if received.haslayer(BOOTP):
                        offered_ip = received[BOOTP].yiaddr if received[BOOTP].yiaddr != '0.0.0.0' else "N/A"

                    dhcp_servers.append({
                        "ip": server_ip,
                        "server_ip": server_ip,  # Compatibilidad con formato de reporte
                        "mac": server_mac,
                        "offered_ip": offered_ip,
                        "response_time_ms": int(elapsed_time * 1000),
                        "timestamp": time.time(),
                        "message_type": "DHCP Offer" if message_type == 2 else "DHCP Ack"
                    })
                    print(f"Servidor DHCP encontrado: {server_ip} (MAC: {server_mac}, IP ofrecida: {offered_ip})")

        # Si no se encontraron servidores, agregar información de diagnóstico
        if len(dhcp_servers) == 0:
            help_msg = "No se detectaron servidores DHCP. Posibles causas:\n"
            help_msg += "1. El programa necesita ejecutarse como Administrador\n"
            help_msg += "2. Npcap no está instalado correctamente (https://npcap.com/)\n"
            help_msg += "3. Tu red usa configuración IP estática en lugar de DHCP\n"
            help_msg += "4. El firewall está bloqueando los paquetes DHCP\n"
            help_msg += "5. La interfaz de red seleccionada no es la correcta"

            return {
                "success": False,
                "error": "No se encontraron servidores DHCP en la red",
                "help": help_msg,
                "interface": interface,
                "packets_sent": 1,
                "packets_received": len(responses),
                "dhcp_responses": 0
            }

        return {
            "success": True,
            "servers": dhcp_servers,
            "dhcp_servers": dhcp_servers,  # Compatibilidad con formato UI
            "interface": interface,
            "total_found": len(dhcp_servers),
            "response_time_ms": int(elapsed_time * 1000)
        }

    except PermissionError:
        return {
            "success": False,
            "error": "Error de permisos: Se requieren privilegios de administrador para ejecutar descubrimiento DHCP.",
            "help": "En Windows: Ejecuta como administrador. En Linux/Mac: Usa sudo"
        }
    except Exception as e:
        error_msg = str(e)
        help_msg = ""

        # Detectar errores comunes y proporcionar ayuda
        if "WinPcap" in error_msg or "Npcap" in error_msg:
            help_msg = "Instala Npcap desde https://npcap.com/ (necesario para Scapy en Windows)"
        elif "No such device" in error_msg or "permission denied" in error_msg.lower():
            help_msg = "Verifica que tienes permisos de administrador y que la interfaz de red está activa"
        elif "timeout" in error_msg.lower():
            help_msg = "No se recibieron respuestas DHCP. Verifica que estás conectado a una red con DHCP activo"

        return {
            "success": False,
            "error": f"Error en descubrimiento DHCP: {error_msg}",
            "help": help_msg
        }

def discover_network_devices(subnet=None):
    """Descubre dispositivos en la red usando nmap"""
    try:
        if not subnet:
            # Detectar subnet automáticamente
            local_ip = get_local_ip()
            if local_ip == "No disponible":
                return {"success": False, "error": "No se pudo obtener IP local"}

            # Calcular subnet /24
            ip_obj = ipaddress.IPv4Address(local_ip)
            network = ipaddress.IPv4Network(f"{local_ip}/24", strict=False)
            subnet = str(network)

        # Ejecutar nmap
        # -sn: Ping scan (no port scan)
        # -T4: Timing template 4 (aggressive, más rápido)
        # --min-rate 1000: Enviar al menos 1000 paquetes por segundo
        cmd = ["nmap", "-sn", "-T4", "--min-rate", "1000", subnet]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=180)

        if result.returncode != 0:
            return {"success": False, "error": f"Error ejecutando nmap: {result.stderr}"}

        # Parsear resultados
        devices = []
        lines = result.stdout.split('\n')
        current_device = {}

        for line in lines:
            if "Nmap scan report for" in line:
                if current_device:
                    devices.append(current_device)
                ip_match = re.search(r'(\d+\.\d+\.\d+\.\d+)', line)
                if ip_match:
                    current_device = {"ip": ip_match.group(1), "hostname": "", "mac": "", "vendor": ""}
                    if "(" in line and ")" in line:
                        hostname = re.search(r'\((.*?)\)', line)
                        if hostname:
                            current_device["hostname"] = hostname.group(1)
            elif "MAC Address:" in line:
                mac_match = re.search(r'([0-9A-Fa-f]{2}[:-]){5}([0-9A-Fa-f]{2})', line)
                if mac_match and current_device:
                    current_device["mac"] = mac_match.group(0)
                    vendor_match = re.search(r'\((.*?)\)', line)
                    if vendor_match:
                        current_device["vendor"] = vendor_match.group(1)

        if current_device:
            devices.append(current_device)

        return {
            "success": True,
            "devices": devices,
            "subnet_scanned": subnet,
            "devices_found": len(devices),
            "subnet": subnet,  # Mantener por compatibilidad
            "total_found": len(devices)  # Mantener por compatibilidad
        }

    except subprocess.TimeoutExpired:
        return {"success": False, "error": "Timeout ejecutando nmap"}
    except FileNotFoundError:
        return {"success": False, "error": format_installation_error("nmap")}
    except Exception as e:
        return {"success": False, "error": f"Error en descubrimiento de red: {str(e)}"}

def run_mtr_analysis(target_host):
    """Ejecuta análisis PathPing (Windows) o MTR (Linux/Mac) a un host específico"""
    try:
        system = platform.system().lower()

        if system == "windows":
            # Usar PathPing en Windows (comando nativo, no requiere instalación)
            # -n: usar direcciones IP en lugar de resolver nombres
            # -q: número de consultas por salto (10 = más rápido que 100 por defecto)
            cmd = ["pathping", "-n", "-q", "10", target_host]
            # En Windows, PathPing usa la codificación del sistema (generalmente cp850 o cp1252)
            import locale
            system_encoding = locale.getpreferredencoding()
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=300, encoding=system_encoding, errors='replace')

            if result.returncode != 0 and not result.stdout:
                return {"success": False, "error": f"Error ejecutando PathPing: {result.stderr if result.stderr else 'Sin salida'}"}

            # Parsear resultados de PathPing
            output = result.stdout
            lines = output.strip().split('\n')
            hops = []

            # PathPing tiene dos fases:
            # 1. Fase de traceroute (primeras líneas)
            # 2. Fase de estadísticas (después de "Computing statistics for X seconds...")

            parsing_route = False
            parsing_stats = False
            route_hops = {}  # Mapeo de hop_num a hostname

            for line in lines:
                line_stripped = line.strip()

                # Detectar inicio de ruta de traceo
                if "Tracing route" in line or "Rastreando ruta" in line or "Seguimiento de ruta" in line:
                    parsing_route = True
                    continue

                # Detectar inicio de estadísticas
                if "Computing statistics" in line or "Calculando estadísticas" in line or "Procesamiento de estad" in line:
                    parsing_route = False
                    parsing_stats = True
                    continue

                # Parsear fase de traceroute para obtener hostnames/IPs
                if parsing_route and line_stripped:
                    # Formato: "  1  2800:4f0:500:9745:294b:9d73:a767:695d"
                    # O:       "  1  192.168.1.1"
                    # IPv4 y IPv6
                    match = re.match(r'\s*(\d+)\s+(\S+)\s*$', line)
                    if match:
                        hop_num = int(match.group(1))
                        ip = match.group(2)
                        # Filtrar líneas inválidas
                        if ip not in ["*", "---", "timed", "out."]:
                            route_hops[hop_num] = ip

                # Parsear fase de estadísticas
                if parsing_stats and line_stripped:
                    # Formato de línea de estadísticas en Windows (español):
                    # "  1    0ms     0/  10 =  0%     0/  10 =  0%  2800:4f0:500:9745:..."
                    # Salto  RTT    Perdido/Enviado = Pct  Perdido/Enviado = Pct  Dirección

                    # Buscar líneas que empiecen con número de salto seguido de RTT
                    match = re.match(r'\s*(\d+)\s+([\d\-]+)ms\s+(\d+)/\s*(\d+)\s*=\s*(\d+)%\s+(\d+)/\s*(\d+)\s*=\s*(\d+)%\s+(\S+)', line)

                    if not match:
                        # Intentar formato alternativo sin RTT (cuando muestra ---)
                        match = re.match(r'\s*(\d+)\s+(---)\s+(\d+)/\s*(\d+)\s*=\s*(\d+)%\s+(\d+)/\s*(\d+)\s*=\s*(\d+)%\s+(\S+)', line)

                    if match:
                        hop_num = int(match.group(1))
                        rtt_str = match.group(2)
                        rtt_ms = int(rtt_str) if rtt_str != "---" else 0

                        # Primera columna: Origen hasta aquí (acumulado)
                        lost_cumul = int(match.group(3))
                        sent_cumul = int(match.group(4))
                        loss_pct_cumul = int(match.group(5))

                        # Segunda columna: Este nodo/vínculo (específico del salto)
                        lost = int(match.group(6))
                        sent = int(match.group(7))
                        loss_pct = int(match.group(8))

                        hostname = match.group(9)

                        # Usar el hostname de la fase de traceroute si está disponible
                        if hop_num in route_hops:
                            hostname = route_hops[hop_num]

                        # Saltar el hop 0 (origen) y hops sin respuesta válida
                        if hop_num > 0 and hostname not in ["---", "*"]:
                            hops.append({
                                "hop": hop_num,
                                "hostname": hostname,
                                "loss_percent": loss_pct,  # Pérdida específica del salto
                                "packets_sent": sent,
                                "packets_lost": lost,
                                "last_ms": rtt_ms,
                                "avg_ms": rtt_ms,
                                "best_ms": rtt_ms if rtt_ms > 0 else 0,
                                "worst_ms": rtt_ms if rtt_ms > 0 else 0
                            })

            # Si no se parsearon hops, intentar con un enfoque más simple
            if not hops:
                # Buscar líneas que contengan estadísticas de forma más permisiva
                for line in lines:
                    if "/" in line and "=" in line and "%" in line:
                        parts = line.split()
                        for i, part in enumerate(parts):
                            if "/" in part and i > 0:
                                try:
                                    hop_num = int(parts[0])
                                    hostname = parts[1] if len(parts) > 1 else f"hop_{hop_num}"

                                    # Parsear "lost/sent"
                                    lost_sent = part.split("/")
                                    lost = int(lost_sent[0])
                                    sent = int(lost_sent[1])

                                    # Buscar pérdida "%"
                                    loss_pct = 0
                                    for p in parts[i:]:
                                        if "%" in p:
                                            loss_pct = int(p.replace("%", "").replace("=", ""))
                                            break

                                    if hop_num > 0 and hostname not in ["---", "*"]:
                                        hops.append({
                                            "hop": hop_num,
                                            "hostname": hostname,
                                            "loss_percent": loss_pct,
                                            "packets_sent": sent,
                                            "packets_lost": lost,
                                            "last_ms": 0,
                                            "avg_ms": 0,
                                            "best_ms": 0,
                                            "worst_ms": 0
                                        })
                                except (ValueError, IndexError):
                                    continue

            return {
                "success": True,
                "target": target_host,
                "method": "pathping",
                "hops": hops,
                "total_hops": len(hops),
                "raw_output": output[:500] if not hops else "",  # Incluir salida raw si falla el parseo
                "timestamp": time.time()
            }

        else:
            # Usar MTR en Linux/Mac
            cmd = ["mtr", "--report", "--report-cycles", "10", target_host]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)

            if result.returncode != 0:
                return {"success": False, "error": f"Error ejecutando MTR: {result.stderr}"}

            # Parsear resultados MTR
            lines = result.stdout.strip().split('\n')
            hops = []

            for line in lines:
                if line.strip() and "HOST:" not in line and "Start:" not in line:
                    parts = line.split()
                    if len(parts) >= 7:
                        try:
                            hop_num = int(parts[0].rstrip('.'))
                            hostname = parts[1]
                            loss_pct = float(parts[2].rstrip('%'))
                            snt = int(parts[3])
                            last_ms = float(parts[4])
                            avg_ms = float(parts[5])
                            best_ms = float(parts[6])
                            worst_ms = float(parts[7]) if len(parts) > 7 else avg_ms

                            hops.append({
                                "hop": hop_num,
                                "hostname": hostname,
                                "loss_percent": loss_pct,
                                "packets_sent": snt,
                                "last_ms": last_ms,
                                "avg_ms": avg_ms,
                                "best_ms": best_ms,
                                "worst_ms": worst_ms
                            })
                        except (ValueError, IndexError):
                            continue

            return {
                "success": True,
                "target": target_host,
                "method": "mtr",
                "hops": hops,
                "total_hops": len(hops),
                "timestamp": time.time()
            }

    except subprocess.TimeoutExpired:
        return {"success": False, "error": "Timeout ejecutando análisis de ruta"}
    except FileNotFoundError:
        system = platform.system().lower()
        if system == "windows":
            return {"success": False, "error": "PathPing no encontrado (debería estar disponible en Windows)"}
        else:
            return {"success": False, "error": format_installation_error("mtr")}
    except Exception as e:
        return {"success": False, "error": f"Error en análisis de ruta: {str(e)}"}

def run_netflix_speed_test():
    """Ejecuta test de velocidad específico de Netflix usando fast-cli"""
    try:
        # Ejecutar fast-cli (shell=True necesario en Windows para comandos npm globales)
        cmd = ["fast", "--json"]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60, shell=True)

        if result.returncode != 0:
            return {"success": False, "error": f"Error ejecutando fast-cli: {result.stderr}"}

        # Parsear JSON de fast-cli
        try:
            data = json.loads(result.stdout)
            return {
                "success": True,
                "download_mbps": data.get("downloadSpeed", 0),
                "upload_mbps": data.get("uploadSpeed", 0),
                "latency_ms": data.get("latency", 0),
                "timestamp": time.time(),
                "server": data.get("server", "Netflix CDN")
            }
        except json.JSONDecodeError:
            # Fallback para parseo de texto
            output = result.stdout.strip()
            speed_match = re.search(r'(\d+\.?\d*)\s*Mbps', output)
            if speed_match:
                return {
                    "success": True,
                    "download_mbps": float(speed_match.group(1)),
                    "upload_mbps": 0,
                    "latency_ms": 0,
                    "timestamp": time.time(),
                    "server": "Netflix CDN"
                }
            else:
                return {"success": False, "error": "No se pudo parsear resultado de fast-cli"}

    except subprocess.TimeoutExpired:
        return {"success": False, "error": "Timeout ejecutando fast-cli"}
    except FileNotFoundError:
        return {"success": False, "error": format_installation_error("fast-cli")}
    except Exception as e:
        return {"success": False, "error": f"Error en test de Netflix: {str(e)}"}

def run_udp_jitter_test(server_host, server_port=5201, duration=10):
    """Ejecuta test de jitter UDP usando iperf3"""
    try:
        # Ejecutar iperf3 en modo UDP cliente
        cmd = [
            "iperf3", "-c", server_host, "-p", str(server_port),
            "-u", "-t", str(duration), "-J"  # -J para output JSON
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=duration + 30)

        if result.returncode != 0:
            error_msg = result.stderr.strip()

            # Detectar errores comunes y dar ayuda
            help_msg = ""
            if "unable to connect" in error_msg.lower() or "connection refused" in error_msg.lower():
                help_msg = "No se pudo conectar al servidor iperf3. Asegúrate de que haya un servidor iperf3 ejecutándose en " + server_host + ":" + str(server_port) + ". Inicia un servidor con: iperf3 -s"
            elif "no route to host" in error_msg.lower():
                help_msg = "No hay ruta al host " + server_host + ". Verifica que el servidor sea accesible desde tu red."
            elif not error_msg:
                help_msg = "iperf3 se ejecutó pero no produjo salida. Verifica que el servidor iperf3 esté corriendo en " + server_host + ":" + str(server_port)

            return {
                "success": False,
                "error": f"Error ejecutando iperf3: {error_msg if error_msg else 'Sin salida de error'}",
                "help": help_msg
            }

        # Parsear JSON de iperf3
        try:
            data = json.loads(result.stdout)
            end_data = data.get("end", {})

            return {
                "success": True,
                "server": server_host,
                "port": server_port,
                "duration": duration,
                "jitter_ms": end_data.get("sum", {}).get("jitter_ms", 0),
                "lost_packets": end_data.get("sum", {}).get("lost_packets", 0),
                "packets": end_data.get("sum", {}).get("packets", 0),
                "lost_percent": end_data.get("sum", {}).get("lost_percent", 0),
                "bytes": end_data.get("sum", {}).get("bytes", 0),
                "timestamp": time.time()
            }
        except (json.JSONDecodeError, KeyError):
            return {"success": False, "error": "No se pudo parsear resultado de iperf3"}

    except subprocess.TimeoutExpired:
        return {"success": False, "error": "Timeout ejecutando iperf3"}
    except FileNotFoundError:
        return {"success": False, "error": format_installation_error("iperf3")}
    except Exception as e:
        return {"success": False, "error": f"Error en test UDP: {str(e)}"}

def run_cdn_dns_test(domain):
    """Ejecuta test de resolución DNS y latencia para dominios CDN"""
    try:
        start_time = time.time()

        # Resolver DNS
        try:
            answers = dns.resolver.resolve(domain, 'A')
            dns_time = (time.time() - start_time) * 1000  # en ms

            ips = [str(answer) for answer in answers]
        except Exception as e:
            return {"success": False, "error": f"Error resolviendo DNS: {str(e)}"}

        # Hacer ping a cada IP
        ping_results = []
        for ip in ips[:5]:  # Limitar a 5 IPs
            try:
                if platform.system().lower() == "windows":
                    cmd = ["ping", "-n", "4", ip]
                else:
                    cmd = ["ping", "-c", "4", ip]

                result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)

                if result.returncode == 0:
                    # Parsear latencia
                    if platform.system().lower() == "windows":
                        times = re.findall(r'time[<=](\d+)ms', result.stdout)
                    else:
                        times = re.findall(r'time=(\d+\.?\d*) ms', result.stdout)

                    if times:
                        latencies = [float(t) for t in times]
                        avg_latency = sum(latencies) / len(latencies)
                        ping_results.append({
                            "ip": ip,
                            "method": "ICMP ping",
                            "avg_latency": avg_latency,
                            "min_latency": min(latencies),
                            "max_latency": max(latencies),
                            "success": True
                        })
                    else:
                        ping_results.append({
                            "ip": ip,
                            "method": "ICMP ping (no times parsed)",
                            "avg_latency": 0,
                            "min_latency": 0,
                            "max_latency": 0,
                            "success": False
                        })
                else:
                    ping_results.append({
                        "ip": ip,
                        "method": "ICMP ping failed",
                        "avg_latency": 0,
                        "min_latency": 0,
                        "max_latency": 0,
                        "success": False
                    })
            except subprocess.TimeoutExpired:
                ping_results.append({
                    "ip": ip,
                    "method": "ICMP timeout",
                    "avg_latency": 0,
                    "min_latency": 0,
                    "max_latency": 0,
                    "success": False
                })
            except Exception as e:
                # Si ping ICMP falla, intentar conexión TCP (puerto 80)
                try:
                    import socket
                    tcp_latencies = []

                    for _ in range(3):  # 3 intentos TCP
                        start_tcp = time.time()
                        sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                        sock.settimeout(5)

                        try:
                            result = sock.connect_ex((ip, 80))
                            end_tcp = time.time()

                            if result == 0:  # Conexión exitosa
                                tcp_latency = (end_tcp - start_tcp) * 1000
                                tcp_latencies.append(tcp_latency)
                        finally:
                            sock.close()

                    if tcp_latencies:
                        avg_tcp = sum(tcp_latencies) / len(tcp_latencies)
                        ping_results.append({
                            "ip": ip,
                            "method": "TCP connect (port 80)",
                            "avg_latency": avg_tcp,
                            "min_latency": min(tcp_latencies),
                            "max_latency": max(tcp_latencies),
                            "success": True
                        })
                    else:
                        ping_results.append({
                            "ip": ip,
                            "method": "ICMP/TCP failed",
                            "avg_latency": 0,
                            "min_latency": 0,
                            "max_latency": 0,
                            "success": False
                        })
                except:
                    ping_results.append({
                        "ip": ip,
                        "method": "ICMP/TCP failed",
                        "avg_latency": 0,
                        "min_latency": 0,
                        "max_latency": 0,
                        "success": False
                    })

        return {
            "success": True,
            "domain": domain,
            "dns_resolution_time": dns_time,
            "resolved_ips": ips,
            "ping_results": ping_results,
            "timestamp": time.time()
        }

    except Exception as e:
        return {"success": False, "error": f"Error en test CDN DNS: {str(e)}"}

def verify_mtu(target_host="8.8.8.8"):
    """Verifica el MTU óptimo probando diferentes tamaños de paquetes

    Args:
        target_host: Host de destino para pruebas (por defecto Google DNS)

    Returns:
        dict: Resultados de la verificación de MTU
    """
    try:
        results = []
        # Tamaños comunes de MTU a probar
        # 1500 es el estándar Ethernet, 1492 para PPPoE, 1280 IPv6 mínimo
        mtu_sizes = [1500, 1492, 1472, 1400, 1280, 1200, 1024, 576]

        optimal_mtu = None
        system = platform.system().lower()

        for mtu in mtu_sizes:
            # Tamaño del payload = MTU - 28 (20 bytes IP header + 8 bytes ICMP header)
            payload_size = mtu - 28

            try:
                if system == "windows":
                    # En Windows: ping -n 1 -l <size> -f <host>
                    # -f establece el flag "Don't Fragment"
                    cmd = ["ping", "-n", "1", "-l", str(payload_size), "-f", target_host]
                else:
                    # En Linux/Mac: ping -c 1 -M do -s <size> <host>
                    # -M do establece "Don't Fragment"
                    cmd = ["ping", "-c", "1", "-M", "do", "-s", str(payload_size), target_host]

                result = subprocess.run(cmd, capture_output=True, text=True, timeout=5)

                # Verificar si el paquete fue enviado sin fragmentación
                if result.returncode == 0:
                    # Parsear tiempo de respuesta
                    if system == "windows":
                        time_match = re.search(r'time[<=](\d+)ms', result.stdout)
                    else:
                        time_match = re.search(r'time=(\d+\.?\d*) ms', result.stdout)

                    response_time = float(time_match.group(1)) if time_match else 0

                    results.append({
                        "mtu": mtu,
                        "payload_size": payload_size,
                        "success": True,
                        "fragmented": False,
                        "response_time_ms": response_time,
                        "status": "OK"
                    })

                    # El primer MTU que funciona es el óptimo
                    if optimal_mtu is None:
                        optimal_mtu = mtu
                else:
                    # Verificar si el error es por fragmentación
                    error_output = result.stderr + result.stdout
                    fragmentation_error = (
                        "Packet needs to be fragmented" in error_output or
                        "Message too long" in error_output or
                        "packet needs to be fragmented but DF set" in error_output
                    )

                    results.append({
                        "mtu": mtu,
                        "payload_size": payload_size,
                        "success": False,
                        "fragmented": fragmentation_error,
                        "response_time_ms": 0,
                        "status": "FRAGMENTATION_NEEDED" if fragmentation_error else "FAILED"
                    })
            except subprocess.TimeoutExpired:
                results.append({
                    "mtu": mtu,
                    "payload_size": payload_size,
                    "success": False,
                    "fragmented": False,
                    "response_time_ms": 0,
                    "status": "TIMEOUT"
                })
            except Exception as e:
                results.append({
                    "mtu": mtu,
                    "payload_size": payload_size,
                    "success": False,
                    "fragmented": False,
                    "response_time_ms": 0,
                    "status": f"ERROR: {str(e)}"
                })

        # Determinar recomendaciones
        recommendations = []
        if optimal_mtu:
            if optimal_mtu < 1500:
                recommendations.append(f"Tu conexión soporta MTU máximo de {optimal_mtu} bytes.")
                recommendations.append("Considera configurar tu interfaz de red con este MTU para evitar fragmentación.")
            else:
                recommendations.append("Tu conexión soporta el MTU estándar de 1500 bytes (óptimo).")

            if optimal_mtu == 1492:
                recommendations.append("MTU de 1492 es típico de conexiones PPPoE (DSL/ADSL).")
            elif optimal_mtu == 1280:
                recommendations.append("MTU de 1280 es el mínimo para IPv6.")
        else:
            recommendations.append("No se pudo determinar el MTU óptimo. Verifica tu conexión de red.")

        return {
            "success": True,
            "target_host": target_host,
            "optimal_mtu": optimal_mtu,
            "results": results,
            "recommendations": recommendations,
            "timestamp": time.time()
        }

    except Exception as e:
        return {"success": False, "error": f"Error verificando MTU: {str(e)}"}

def run_sustained_load_test(duration_seconds=60, test_server="http://speedtest.tele2.net"):
    """Ejecuta un test de carga sostenida descargando archivos de diferentes tamaños

    Args:
        duration_seconds: No usado (mantenido por compatibilidad)
        test_server: Servidor base para descargas

    Returns:
        dict: Resultados del test de carga sostenida con múltiples archivos
    """
    try:
        import requests

        # Archivos de prueba de diferentes tamaños
        test_files = [
            {"size": "1MB", "url": f"{test_server}/1MB.zip", "size_bytes": 1 * 1024 * 1024},
            {"size": "10MB", "url": f"{test_server}/10MB.zip", "size_bytes": 10 * 1024 * 1024},
            {"size": "50MB", "url": f"{test_server}/50MB.zip", "size_bytes": 50 * 1024 * 1024},
            {"size": "100MB", "url": f"{test_server}/100MB.zip", "size_bytes": 100 * 1024 * 1024},
            {"size": "250MB", "url": f"{test_server}/250MB.zip", "size_bytes": 250 * 1024 * 1024},
            {"size": "1GB", "url": f"{test_server}/1GB.zip", "size_bytes": 1024 * 1024 * 1024}
        ]

        print(f"Iniciando test de carga sostenida con {len(test_files)} archivos...")

        overall_start = time.time()
        file_results = []
        total_bytes_downloaded = 0
        all_samples = []

        for file_info in test_files:
            print(f"Descargando {file_info['size']}...")

            file_start = time.time()
            bytes_downloaded = 0
            samples = []
            chunk_size = 8192  # 8KB chunks

            try:
                response = requests.get(file_info['url'], stream=True, timeout=30)

                if response.status_code != 200:
                    file_results.append({
                        "size": file_info['size'],
                        "success": False,
                        "error": f"HTTP {response.status_code}"
                    })
                    continue

                last_sample_time = file_start
                sample_bytes = 0

                # Descargar archivo completo
                for chunk in response.iter_content(chunk_size=chunk_size):
                    if chunk:
                        chunk_len = len(chunk)
                        bytes_downloaded += chunk_len
                        sample_bytes += chunk_len

                        current_time = time.time()

                        # Tomar muestra cada segundo
                        if current_time - last_sample_time >= 1.0:
                            interval = current_time - last_sample_time
                            speed_mbps = (sample_bytes * 8) / (interval * 1_000_000)  # Mbps

                            sample = {
                                "timestamp": current_time - overall_start,
                                "file": file_info['size'],
                                "speed_mbps": speed_mbps,
                                "bytes_downloaded": bytes_downloaded
                            }
                            samples.append(sample)
                            all_samples.append(sample)

                            last_sample_time = current_time
                            sample_bytes = 0

                file_end = time.time()
                file_duration = file_end - file_start

                # Calcular estadísticas del archivo
                if samples:
                    speeds = [s["speed_mbps"] for s in samples]
                    avg_speed = sum(speeds) / len(speeds)
                    min_speed = min(speeds)
                    max_speed = max(speeds)

                    # Calcular estabilidad
                    if len(speeds) > 1:
                        variance = sum((x - avg_speed) ** 2 for x in speeds) / (len(speeds) - 1)
                        std_dev = variance ** 0.5
                        stability_percent = ((avg_speed - std_dev) / avg_speed * 100) if avg_speed > 0 else 0
                    else:
                        std_dev = 0
                        stability_percent = 100
                else:
                    avg_speed = (bytes_downloaded * 8) / (file_duration * 1_000_000) if file_duration > 0 else 0
                    min_speed = avg_speed
                    max_speed = avg_speed
                    std_dev = 0
                    stability_percent = 100

                file_results.append({
                    "size": file_info['size'],
                    "success": True,
                    "duration_seconds": file_duration,
                    "bytes_downloaded": bytes_downloaded,
                    "mb_downloaded": bytes_downloaded / (1024 * 1024),
                    "avg_speed_mbps": avg_speed,
                    "min_speed_mbps": min_speed,
                    "max_speed_mbps": max_speed,
                    "std_deviation_mbps": std_dev,
                    "stability_percent": stability_percent,
                    "samples": samples[:30]  # Limitar muestras por archivo
                })

                total_bytes_downloaded += bytes_downloaded

            except requests.exceptions.Timeout:
                file_results.append({
                    "size": file_info['size'],
                    "success": False,
                    "error": "Timeout"
                })
                continue
            except Exception as e:
                file_results.append({
                    "size": file_info['size'],
                    "success": False,
                    "error": str(e)
                })
                continue

        overall_end = time.time()
        total_duration = overall_end - overall_start

        # Calcular estadísticas generales
        successful_downloads = [f for f in file_results if f.get("success")]

        if successful_downloads:
            all_speeds = [f["avg_speed_mbps"] for f in successful_downloads]
            overall_avg_speed = sum(all_speeds) / len(all_speeds)
            overall_min_speed = min([f["min_speed_mbps"] for f in successful_downloads])
            overall_max_speed = max([f["max_speed_mbps"] for f in successful_downloads])

            # Detectar degradación de velocidad
            speed_trend = []
            for i, f in enumerate(successful_downloads):
                speed_trend.append(f["avg_speed_mbps"])

            # Analizar tendencia
            recommendations = []
            if len(speed_trend) >= 3:
                first_avg = sum(speed_trend[:2]) / 2
                last_avg = sum(speed_trend[-2:]) / 2
                degradation = ((first_avg - last_avg) / first_avg * 100) if first_avg > 0 else 0

                if degradation > 20:
                    recommendations.append("⚠️ Degradación significativa de velocidad detectada en archivos grandes.")
                    recommendations.append("Posible throttling del ISP o saturación de red.")
                elif degradation > 10:
                    recommendations.append("⚠️ Ligera degradación de velocidad en archivos grandes.")
                else:
                    recommendations.append("✅ Velocidad consistente a través de diferentes tamaños de archivo.")

            # Verificar estabilidad general
            avg_stability = sum([f["stability_percent"] for f in successful_downloads]) / len(successful_downloads)
            if avg_stability < 70:
                recommendations.append("⚠️ Variabilidad significativa detectada durante las descargas.")
            elif avg_stability >= 85:
                recommendations.append("✅ Conexión muy estable durante todas las descargas.")

            return {
                "success": True,
                "total_duration_seconds": total_duration,
                "total_bytes": total_bytes_downloaded,
                "total_mb": total_bytes_downloaded / (1024 * 1024),
                "total_gb": total_bytes_downloaded / (1024 * 1024 * 1024),
                "files_tested": len(test_files),
                "files_successful": len(successful_downloads),
                "files_failed": len(test_files) - len(successful_downloads),
                "overall_avg_speed_mbps": overall_avg_speed,
                "overall_min_speed_mbps": overall_min_speed,
                "overall_max_speed_mbps": overall_max_speed,
                "speed_degradation_percent": degradation if len(speed_trend) >= 3 else 0,
                "avg_stability_percent": avg_stability,
                "file_results": file_results,
                "recommendations": recommendations,
                "timestamp": time.time()
            }
        else:
            return {
                "success": False,
                "error": "No se pudo completar ninguna descarga exitosamente"
            }

    except Exception as e:
        return {"success": False, "error": f"Error en test de carga sostenida: {str(e)}"}

def add_advanced_tests_analysis(doc, advanced_data):
    """Agrega análisis detallado de las pruebas avanzadas al documento"""

    tests_performed = 0
    total_tests = 8

    # DHCP Discovery Analysis
    doc.add_heading('5.1. Descubrimiento de Servidores DHCP', 2)
    dhcp_data = advanced_data.get('dhcp_discovery', {}).get('results', {})
    if dhcp_data and dhcp_data.get('success'):
        tests_performed += 1
        servers = dhcp_data.get('dhcp_servers', [])
        if servers:
            doc.add_paragraph(f'✅ Se detectaron {len(servers)} servidor(es) DHCP en la red local:')
            for i, server in enumerate(servers, 1):
                doc.add_paragraph(f'  {i}. Servidor: {server.get("server_ip", "N/A")} | IP Ofrecida: {server.get("offered_ip", "N/A")} | Tiempo: {server.get("response_time_ms", 0):.1f}ms', style='List Bullet')

            # Análisis
            if len(servers) > 1:
                doc.add_paragraph('⚠️ ADVERTENCIA: Se detectaron múltiples servidores DHCP. Esto puede causar conflictos de IP.', style='Intense Quote')
            else:
                doc.add_paragraph('✅ Configuración DHCP normal detectada.', style='Intense Quote')
        else:
            doc.add_paragraph('❌ No se detectaron servidores DHCP en la red local.')
    else:
        doc.add_paragraph('❌ Prueba DHCP no ejecutada o falló.')

    # Network Discovery Analysis
    doc.add_heading('5.2. Descubrimiento de Dispositivos de Red', 2)
    network_data = advanced_data.get('network_discovery', {}).get('results', {})
    if network_data and network_data.get('success'):
        tests_performed += 1
        devices = network_data.get('devices', [])
        subnet = network_data.get('subnet_scanned', 'N/A')

        doc.add_paragraph(f'✅ Escaneo completado en subred: {subnet}')
        doc.add_paragraph(f'🖥️ Total de dispositivos activos encontrados: {len(devices)}')

        if devices:
            # Crear tabla de dispositivos
            device_table = doc.add_table(rows=1, cols=4)
            device_table.style = 'Light Grid Accent 1'
            add_table_border(device_table)

            headers = ['IP', 'Hostname', 'MAC Address', 'Fabricante']
            header_cells = device_table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].font.bold = True

            for device in devices[:10]:  # Limitar a 10 dispositivos
                row_cells = device_table.add_row().cells
                row_cells[0].text = device.get('ip', 'N/A')
                row_cells[1].text = device.get('hostname', 'N/A')
                row_cells[2].text = device.get('mac', 'N/A')
                row_cells[3].text = device.get('vendor', 'N/A')

            # Análisis de seguridad
            devices_with_mac = len([d for d in devices if d.get('mac')])
            if devices_with_mac > 0:
                doc.add_paragraph(f'🔍 Análisis: {devices_with_mac}/{len(devices)} dispositivos exponen información MAC.', style='Intense Quote')
    else:
        doc.add_paragraph('❌ Prueba de descubrimiento de red no ejecutada o falló.')

    # MTR Analysis
    # Netflix Speed Test
    doc.add_heading('5.4. Test de Velocidad Netflix (Fast.com)', 2)
    netflix_data = advanced_data.get('netflix_speed', {}).get('results', {})
    if netflix_data and netflix_data.get('success'):
        tests_performed += 1
        download_speed = netflix_data.get('download_mbps', 0)
        doc.add_paragraph(f'✅ Velocidad de descarga Netflix: {download_speed:.1f} Mbps')

        # Análisis de calidad
        if download_speed >= 25:
            quality = "4K Ultra HD"
            status = "✅ Excelente"
        elif download_speed >= 15:
            quality = "Full HD (1080p)"
            status = "✅ Buena"
        elif download_speed >= 5:
            quality = "HD (720p)"
            status = "⚠️ Aceptable"
        else:
            quality = "SD (480p)"
            status = "❌ Limitada"

        doc.add_paragraph(f'📺 Calidad de streaming soportada: {quality} ({status})', style='Intense Quote')
    else:
        doc.add_paragraph('❌ Prueba de velocidad Netflix no ejecutada o falló.')

    # UDP Jitter Test
    doc.add_heading('5.5. Test de Jitter UDP (iperf3)', 2)
    udp_data = advanced_data.get('udp_jitter', {}).get('results', {})
    if udp_data and udp_data.get('success'):
        tests_performed += 1
        jitter = udp_data.get('jitter_ms', 0)
        packet_loss = udp_data.get('packet_loss_percent', 0)

        doc.add_paragraph(f'✅ Test UDP completado')
        doc.add_paragraph(f'📊 Jitter promedio: {jitter:.2f} ms')
        doc.add_paragraph(f'📊 Pérdida de paquetes: {packet_loss:.2f}%')

        # Análisis de calidad para VoIP/videoconferencia
        if jitter <= 20 and packet_loss <= 1:
            quality_status = "✅ Excelente para VoIP/videoconferencia"
        elif jitter <= 50 and packet_loss <= 3:
            quality_status = "⚠️ Aceptable para VoIP básico"
        else:
            quality_status = "❌ No recomendado para aplicaciones de tiempo real"

        doc.add_paragraph(f'🎯 Evaluación: {quality_status}', style='Intense Quote')
    else:
        doc.add_paragraph('❌ Prueba de jitter UDP no ejecutada o falló.')

    # CDN DNS Test
    doc.add_heading('5.6. Test de DNS y Latencia CDN', 2)
    cdn_data = advanced_data.get('cdn_dns', {}).get('results', {})
    if cdn_data and cdn_data.get('success'):
        tests_performed += 1
        domain = cdn_data.get('domain', 'N/A')
        dns_time = cdn_data.get('dns_resolution_time', 0)
        ips = cdn_data.get('resolved_ips', [])

        doc.add_paragraph(f'✅ Test CDN DNS completado para: {domain}')
        doc.add_paragraph(f'🕐 Tiempo de resolución DNS: {dns_time:.1f} ms')
        doc.add_paragraph(f'🌐 IPs resueltas: {len(ips)}')

        # Análisis de conectividad
        ping_results = cdn_data.get('ping_results', [])
        successful_pings = [p for p in ping_results if p.get('success')]
        if successful_pings:
            avg_latency = sum(p.get('avg_latency', 0) for p in successful_pings) / len(successful_pings)
            doc.add_paragraph(f'📡 Latencia promedio a CDN: {avg_latency:.1f} ms', style='Intense Quote')
        else:
            doc.add_paragraph('⚠️ No se pudo establecer conectividad con el CDN via ICMP/TCP', style='Intense Quote')
    else:
        doc.add_paragraph('❌ Prueba CDN DNS no ejecutada o falló.')

    # PathPing/MTR Analysis
    doc.add_heading('5.3. Análisis PathPing/Traceroute', 2)
    mtr_data = advanced_data.get('mtr_analysis', {}).get('results', {})
    if mtr_data and mtr_data.get('success'):
        tests_performed += 1
        target = mtr_data.get('target', 'N/A')
        method = mtr_data.get('method', 'N/A')
        total_hops = mtr_data.get('total_hops', 0)

        doc.add_paragraph(f'✅ Análisis completado hacia: {target}')
        doc.add_paragraph(f'🛣️ Método: {method.upper()} ({("Windows" if method == "pathping" else "Linux/Mac")})')
        doc.add_paragraph(f'📍 Total de saltos: {total_hops}')

        hops = mtr_data.get('hops', [])
        if hops:
            # Análisis de pérdida de paquetes
            hops_with_loss = [h for h in hops if h.get('loss_percent', 0) > 0]
            if hops_with_loss:
                doc.add_paragraph(f'⚠️ {len(hops_with_loss)} salto(s) con pérdida de paquetes detectada:', style='Intense Quote')
                for hop in hops_with_loss[:5]:  # Limitar a 5
                    doc.add_paragraph(f'  • Salto {hop.get("hop")}: {hop.get("hostname")} - {hop.get("loss_percent")}% pérdida', style='List Bullet')
            else:
                doc.add_paragraph('✅ Sin pérdida de paquetes en ningún salto.', style='Intense Quote')

            # Análisis de latencia
            avg_latencies = [h.get('avg_ms', 0) for h in hops if h.get('avg_ms', 0) > 0]
            if avg_latencies:
                max_latency = max(avg_latencies)
                doc.add_paragraph(f'📊 Latencia máxima en ruta: {max_latency:.1f} ms', style='Intense Quote')
        else:
            doc.add_paragraph('⚠️ No se obtuvieron estadísticas de saltos.', style='Intense Quote')
    else:
        doc.add_paragraph('❌ Prueba PathPing/Traceroute no ejecutada o falló.')

    # MTU Verification Analysis
    doc.add_heading('5.7. Verificación MTU', 2)
    mtu_data = advanced_data.get('mtu_verification', {}).get('results', {})
    if mtu_data and mtu_data.get('success'):
        tests_performed += 1
        target = mtu_data.get('target_host', 'N/A')
        optimal_mtu = mtu_data.get('optimal_mtu')

        doc.add_paragraph(f'✅ Verificación MTU completada hacia: {target}')

        if optimal_mtu:
            doc.add_paragraph(f'🎯 MTU óptimo detectado: {optimal_mtu} bytes')

            # Análisis del MTU
            if optimal_mtu < 1500:
                doc.add_paragraph(f'⚠️ Tu conexión soporta MTU máximo de {optimal_mtu} bytes (menor al estándar de 1500).', style='Intense Quote')
                if optimal_mtu == 1492:
                    doc.add_paragraph('ℹ️ MTU de 1492 es típico de conexiones PPPoE (DSL/ADSL).', style='Intense Quote')
            else:
                doc.add_paragraph('✅ Tu conexión soporta el MTU estándar de 1500 bytes.', style='Intense Quote')

            # Recomendaciones
            recommendations = mtu_data.get('recommendations', [])
            if recommendations:
                doc.add_paragraph('💡 Recomendaciones:')
                for rec in recommendations[:3]:  # Limitar a 3 recomendaciones
                    doc.add_paragraph(f'  • {rec}', style='List Bullet')
        else:
            doc.add_paragraph('⚠️ No se pudo determinar el MTU óptimo.', style='Intense Quote')
    else:
        doc.add_paragraph('❌ Prueba de verificación MTU no ejecutada o falló.')

    # Sustained Load Test Analysis
    doc.add_heading('5.8. Test de Carga Sostenida', 2)
    load_data = advanced_data.get('sustained_load', {}).get('results', {})
    if load_data and load_data.get('success'):
        tests_performed += 1
        duration = load_data.get('duration_seconds', 0)
        avg_speed = load_data.get('avg_speed_mbps', 0)
        stability = load_data.get('stability_percent', 0)
        drops = load_data.get('drops_detected', 0)

        doc.add_paragraph(f'✅ Test de carga sostenida completado')
        doc.add_paragraph(f'⏱️ Duración: {duration:.1f} segundos')
        doc.add_paragraph(f'📊 Velocidad promedio: {avg_speed:.2f} Mbps')
        doc.add_paragraph(f'📈 Estabilidad: {stability:.1f}%')
        doc.add_paragraph(f'⚠️ Caídas detectadas: {drops}')

        # Análisis de estabilidad
        if stability >= 80:
            doc.add_paragraph('✅ La conexión es estable y consistente durante cargas sostenidas.', style='Intense Quote')
        elif stability >= 60:
            doc.add_paragraph('⚠️ La conexión muestra variabilidad moderada bajo carga sostenida.', style='Intense Quote')
        else:
            doc.add_paragraph('❌ La conexión muestra variabilidad significativa bajo carga sostenida.', style='Intense Quote')

        if drops > 0:
            doc.add_paragraph(f'⚠️ Se detectaron {drops} caídas de velocidad durante el test.', style='Intense Quote')

        # Recomendaciones
        recommendations = load_data.get('recommendations', [])
        if recommendations:
            doc.add_paragraph('💡 Recomendaciones:')
            for rec in recommendations[:3]:  # Limitar a 3 recomendaciones
                # Remover emojis para el documento
                clean_rec = rec.replace('⚠️', '').replace('✅', '').strip()
                doc.add_paragraph(f'  • {clean_rec}', style='List Bullet')
    else:
        doc.add_paragraph('❌ Test de carga sostenida no ejecutado o falló.')

    # Resumen de pruebas avanzadas
    doc.add_paragraph()
    summary_para = doc.add_paragraph()
    summary_para.add_run('📋 Resumen de Pruebas Avanzadas: ').bold = True
    summary_para.add_run(f'{tests_performed}/{total_tests} pruebas ejecutadas exitosamente.')

def add_traceroute_analysis(doc, traceroute_data):
    """Agrega análisis detallado del historial de traceroute/pathping"""

    if not traceroute_data:
        doc.add_paragraph('❌ No hay datos de traceroute/pathping disponibles.')
        return

    # Estadísticas generales
    total_traces = len(traceroute_data)
    successful_traces = len([t for t in traceroute_data if not t.get('error', False)])
    pathping_traces = len([t for t in traceroute_data if t.get('method') == 'pathping'])
    traceroute_traces = len([t for t in traceroute_data if t.get('method') == 'traceroute'])

    doc.add_paragraph(f'📊 Total de análisis de rutas realizados: {total_traces}')
    doc.add_paragraph(f'✅ Análisis exitosos: {successful_traces}')
    doc.add_paragraph(f'🛤️ Traceroutes: {traceroute_traces}')
    doc.add_paragraph(f'📊 PathPings: {pathping_traces}')

    if successful_traces == 0:
        return

    # Análisis por sitios más analizados
    site_counts = {}
    for trace in traceroute_data:
        if not trace.get('error', False):
            site = trace.get('website', 'N/A')
            site_counts[site] = site_counts.get(site, 0) + 1

    if site_counts:
        doc.add_paragraph()
        doc.add_paragraph('🎯 Sitios más analizados:')
        sorted_sites = sorted(site_counts.items(), key=lambda x: x[1], reverse=True)
        for site, count in sorted_sites[:5]:
            doc.add_paragraph(f'  • {site}: {count} análisis', style='List Bullet')

    # Análisis de métodos utilizados
    if pathping_traces > 0:
        doc.add_paragraph()
        pathping_para = doc.add_paragraph()
        pathping_para.add_run('📊 Análisis PathPing: ').bold = True
        pathping_para.add_run('Se utilizó PathPing para obtener estadísticas detalladas de pérdida de paquetes y latencia por salto.')

    # Mostrar últimos análisis (hasta 3)
    recent_traces = sorted(traceroute_data, key=lambda x: x.get('timestamp', 0), reverse=True)[:3]

    doc.add_paragraph()
    doc.add_paragraph('🕐 Análisis recientes:')

    for trace in recent_traces:
        if trace.get('error', False):
            continue

        date = datetime.fromtimestamp(trace.get('timestamp', 0))
        method_icon = '📊' if trace.get('method') == 'pathping' else '🛤️'

        trace_para = doc.add_paragraph()
        trace_para.add_run(f'{method_icon} {trace.get("method", "N/A").upper()} ').bold = True
        trace_para.add_run(f'a {trace.get("website", "N/A")} ({date.strftime("%d/%m/%Y %H:%M")})')

def add_dns_benchmark_analysis(doc, dns_data):
    """Agrega análisis del benchmark DNS"""

    if dns_data.get('status') != 'completed':
        doc.add_paragraph('❌ Benchmark DNS no ejecutado.')
        return

    results = dns_data.get('results', [])
    if not results:
        doc.add_paragraph('❌ No hay resultados de benchmark DNS disponibles.')
        return

    doc.add_paragraph(f'✅ Benchmark DNS completado con {len(results)} servidores.')

    # Crear tabla de resultados
    dns_table = doc.add_table(rows=1, cols=4)
    dns_table.style = 'Light Grid Accent 1'
    add_table_border(dns_table)

    headers = ['Servidor DNS', 'IP', 'Tiempo Promedio (ms)', 'Tasa de Éxito (%)']
    header_cells = dns_table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True

    # Ordenar por tiempo de respuesta
    sorted_results = sorted(results, key=lambda x: x.get('avg_time', 999))

    for dns_server in sorted_results:
        row_cells = dns_table.add_row().cells
        row_cells[0].text = dns_server.get('name', 'N/A')
        row_cells[1].text = dns_server.get('ip', 'N/A')
        row_cells[2].text = f"{dns_server.get('avg_time', 0):.1f}"
        row_cells[3].text = f"{dns_server.get('success_rate', 0):.1f}"

    # Análisis y recomendaciones
    if sorted_results:
        fastest = sorted_results[0]
        slowest = sorted_results[-1]

        doc.add_paragraph()
        analysis_para = doc.add_paragraph()
        analysis_para.add_run('🏆 Servidor DNS más rápido: ').bold = True
        analysis_para.add_run(f'{fastest.get("name", "N/A")} ({fastest.get("avg_time", 0):.1f}ms)')

        if fastest.get('avg_time', 0) <= 20:
            performance = "✅ Excelente"
        elif fastest.get('avg_time', 0) <= 50:
            performance = "⚠️ Aceptable"
        else:
            performance = "❌ Lento"

        doc.add_paragraph(f'📊 Rendimiento DNS general: {performance}', style='Intense Quote')

def add_advanced_recommendations(recommendations, advanced_data, dns_data, traceroute_data):
    """Agrega recomendaciones basadas en los resultados de pruebas avanzadas"""

    # Análisis DHCP
    dhcp_data = advanced_data.get('dhcp_discovery', {}).get('results', {})
    if dhcp_data and dhcp_data.get('success'):
        servers = dhcp_data.get('dhcp_servers', [])
        if len(servers) > 1:
            recommendations.append({
                'priority': 'ALTA',
                'action': 'Resolver conflicto de múltiples servidores DHCP',
                'detail': f'Se detectaron {len(servers)} servidores DHCP. Esto puede causar conflictos de asignación de IP y problemas de conectividad.'
            })
        elif len(servers) == 0:
            recommendations.append({
                'priority': 'MEDIA',
                'action': 'Verificar configuración DHCP',
                'detail': 'No se detectaron servidores DHCP. Verificar configuración de red y disponibilidad del servicio.'
            })

    # Análisis Network Discovery
    network_data = advanced_data.get('network_discovery', {}).get('results', {})
    if network_data and network_data.get('success'):
        devices = network_data.get('devices', [])
        if len(devices) > 50:
            recommendations.append({
                'priority': 'MEDIA',
                'action': 'Revisar segmentación de red',
                'detail': f'Se detectaron {len(devices)} dispositivos en la red. Considerar segmentación para mejorar rendimiento y seguridad.'
            })

        # Análisis de seguridad básico
        devices_with_mac = len([d for d in devices if d.get('mac')])
        if devices_with_mac > 20:
            recommendations.append({
                'priority': 'BAJA',
                'action': 'Implementar monitoreo de dispositivos',
                'detail': 'Alto número de dispositivos detectados. Considerar implementar sistema de inventario y monitoreo de red.'
            })

    # Análisis Netflix/Streaming
    netflix_data = advanced_data.get('netflix_speed', {}).get('results', {})
    if netflix_data and netflix_data.get('success'):
        download_speed = netflix_data.get('download_mbps', 0)
        if download_speed < 5:
            recommendations.append({
                'priority': 'ALTA',
                'action': 'Optimizar conexión para streaming',
                'detail': f'Velocidad Netflix: {download_speed:.1f} Mbps. Insuficiente para streaming HD. Revisar configuración de red o plan de internet.'
            })
        elif download_speed < 15:
            recommendations.append({
                'priority': 'MEDIA',
                'action': 'Considerar upgrade de velocidad',
                'detail': f'Velocidad Netflix: {download_speed:.1f} Mbps. Suficiente para HD básico, pero limitado para 4K o múltiples streams.'
            })

    # Análisis Jitter UDP
    udp_data = advanced_data.get('udp_jitter', {}).get('results', {})
    if udp_data and udp_data.get('success'):
        jitter = udp_data.get('jitter_ms', 0)
        packet_loss = udp_data.get('packet_loss_percent', 0)

        if jitter > 50 or packet_loss > 3:
            recommendations.append({
                'priority': 'ALTA',
                'action': 'Optimizar calidad para aplicaciones de tiempo real',
                'detail': f'Jitter: {jitter:.1f}ms, Pérdida: {packet_loss:.1f}%. No apto para VoIP/videoconferencia. Revisar configuración QoS.'
            })
        elif jitter > 20 or packet_loss > 1:
            recommendations.append({
                'priority': 'MEDIA',
                'action': 'Ajustar configuración para VoIP',
                'detail': f'Jitter: {jitter:.1f}ms, Pérdida: {packet_loss:.1f}%. Calidad límite para VoIP. Considerar optimizaciones de red.'
            })

    # Análisis DNS
    if dns_data.get('status') == 'completed':
        results = dns_data.get('results', [])
        if results:
            fastest = min(results, key=lambda x: x.get('avg_time', 999))
            if fastest.get('avg_time', 0) > 50:
                recommendations.append({
                    'priority': 'MEDIA',
                    'action': 'Optimizar configuración DNS',
                    'detail': f'DNS más rápido: {fastest.get("avg_time", 0):.1f}ms. Considerar cambiar a servidores DNS más rápidos (ej: 1.1.1.1, 8.8.8.8).'
                })

    # Análisis Traceroute/PathPing
    if traceroute_data:
        pathping_traces = [t for t in traceroute_data if t.get('method') == 'pathping' and not t.get('error', False)]
        if len(pathping_traces) > 0:
            recommendations.append({
                'priority': 'BAJA',
                'action': 'Continuar análisis detallado de rutas',
                'detail': f'Se realizaron {len(pathping_traces)} análisis PathPing. Revisar resultados para optimizar rutas críticas.'
            })

    # Análisis MTR
    mtr_data = advanced_data.get('mtr_analysis', {}).get('results', {})
    if mtr_data and mtr_data.get('success'):
        hops = mtr_data.get('hops', [])
        packet_loss_hops = [h for h in hops if h.get('packet_loss', 0) > 0]

        if len(packet_loss_hops) > 0:
            recommendations.append({
                'priority': 'ALTA',
                'action': 'Investigar pérdida de paquetes en ruta',
                'detail': f'MTR detectó pérdida de paquetes en {len(packet_loss_hops)} saltos. Contactar ISP si el problema persiste.'
            })

    # Recomendación general de pruebas avanzadas
    tests_performed = sum(1 for test_data in advanced_data.values() if test_data.get('results'))
    if tests_performed < 3:
        recommendations.append({
            'priority': 'BAJA',
            'action': 'Ejecutar más pruebas avanzadas',
            'detail': f'Solo {tests_performed}/6 pruebas avanzadas ejecutadas. Realizar análisis completo para diagnóstico más preciso.'
        })

def generate_technical_report():
    """Genera un reporte técnico completo en formato DOCX incluyendo pruebas básicas y avanzadas"""

    # Obtener información del sistema
    local_ip = get_local_ip()
    public_ip = get_public_ip()
    report_date = datetime.now()
    
    # Obtener datos básicos
    with data_lock:
        ping_data = dict(ping_results_data)

    with speedtest_lock:
        speedtest_data = dict(speedtest_results_data)

    # Obtener datos de pruebas avanzadas
    with advanced_tests_lock:
        advanced_data = {
            'dhcp_discovery': dict(advanced_tests_results.get('dhcp_discovery', {})),
            'network_discovery': dict(advanced_tests_results.get('network_discovery', {})),
            'mtr_analysis': dict(advanced_tests_results.get('mtr_analysis', {})),
            'netflix_speed': dict(advanced_tests_results.get('netflix_speed', {})),
            'udp_jitter': dict(advanced_tests_results.get('udp_jitter', {})),
            'cdn_dns': dict(advanced_tests_results.get('cdn_dns', {}))
        }

    # Obtener historial de traceroute/pathping
    with traceroute_history_lock:
        traceroute_data = list(traceroute_history)

    # Obtener benchmark DNS
    with dns_benchmark_lock:
        dns_data = dict(dns_benchmark_results)
    
    # Calcular periodo de análisis
    all_timestamps = []
    for site, data in ping_data.items():
        if data:
            all_timestamps.extend([p['timestamp'] for p in data if 'timestamp' in p])
    
    if all_timestamps:
        start_time = datetime.fromtimestamp(min(all_timestamps))
        end_time = datetime.fromtimestamp(max(all_timestamps))
    else:
        start_time = report_date
        end_time = report_date
    
    # Crear documento
    doc = Document()
    
    # --- PORTADA ---
    title = doc.add_heading('REPORTE TÉCNICO DE MONITOREO', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Infraestructura de Red y Conectividad', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Contar pruebas avanzadas realizadas
    advanced_tests_performed = sum(1 for test_data in advanced_data.values() if test_data.get('results'))
    total_traceroutes = len([t for t in traceroute_data if not t.get('error', False)])
    dns_benchmark_done = dns_data.get('status') == 'completed'

    # Información del sistema
    info_table = doc.add_table(rows=10, cols=2)
    info_table.style = 'Light Grid Accent 1'
    add_table_border(info_table)

    info_data = [
        ('Fecha de Generación:', report_date.strftime('%d/%m/%Y %H:%M:%S')),
        ('Periodo Analizado:', f'{start_time.strftime("%d/%m/%Y %H:%M")} - {end_time.strftime("%d/%m/%Y %H:%M")}'),
        ('IP Local del Monitor:', local_ip),
        ('IP Pública de Salida:', public_ip),
        ('Total de Sitios Monitoreados:', str(len(ping_data))),
        ('Sistema Operativo:', platform.system() + ' ' + platform.release()),
        ('Intervalo de Monitoreo:', f'{PING_INTERVAL_SECONDS} segundos'),
        ('Pruebas Avanzadas Realizadas:', f'{advanced_tests_performed}/6'),
        ('Análisis de Rutas (Traceroute/PathPing):', str(len(traceroute_data))),
        ('Benchmark DNS:', 'Completado' if dns_benchmark_done else 'No ejecutado'),
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # --- RESUMEN EJECUTIVO ---
    doc.add_heading('1. RESUMEN EJECUTIVO', 1)
    
    # Calcular estadísticas generales
    total_sites = len(ping_data)
    sites_ok = sum(1 for site, data in ping_data.items() if data and analyze_site_data(data) and analyze_site_data(data)['availability'] >= 95)
    sites_warning = sum(1 for site, data in ping_data.items() if data and analyze_site_data(data) and 80 <= analyze_site_data(data)['availability'] < 95)
    sites_critical = sum(1 for site, data in ping_data.items() if data and analyze_site_data(data) and analyze_site_data(data)['availability'] < 80)
    
    all_latencies = []
    for site, data in ping_data.items():
        if data:
            stats = analyze_site_data(data)
            if stats and stats['avg_latency']:
                all_latencies.append(stats['avg_latency'])
    
    avg_network_latency = round(mean(all_latencies), 2) if all_latencies and len(all_latencies) > 0 else 0
    
    # Estado general
    if sites_critical > 0:
        general_status = 'CRÍTICO'
        status_color = RGBColor(231, 76, 60)
    elif sites_warning > 0:
        general_status = 'ACEPTABLE'
        status_color = RGBColor(241, 196, 15)
    else:
        general_status = 'ÓPTIMO'
        status_color = RGBColor(46, 204, 113)
    
    p = doc.add_paragraph()
    p.add_run('Estado General del Sistema: ').bold = True
    run = p.add_run(general_status)
    run.bold = True
    run.font.color.rgb = status_color
    run.font.size = Pt(14)
    
    doc.add_paragraph(f'• Total de sitios monitoreados: {total_sites}')
    doc.add_paragraph(f'• Sitios con disponibilidad óptima (≥95%): {sites_ok}')
    doc.add_paragraph(f'• Sitios con advertencias (80-95%): {sites_warning}')
    doc.add_paragraph(f'• Sitios críticos (<80%): {sites_critical}')
    doc.add_paragraph(f'• Latencia promedio de red: {avg_network_latency} ms')
    
    # Speedtest info
    if speedtest_data.get('latest'):
        latest_speed = speedtest_data['latest']
        doc.add_paragraph(f'• Velocidad de descarga actual: {latest_speed.get("download_mbps", "N/A")} Mbps')
        doc.add_paragraph(f'• Velocidad de subida actual: {latest_speed.get("upload_mbps", "N/A")} Mbps')
    
    doc.add_page_break()
    
    # --- METODOLOGÍA ---
    doc.add_heading('2. METODOLOGÍA DE MONITOREO', 1)
    
    doc.add_paragraph('El sistema de monitoreo implementa las siguientes técnicas:')
    doc.add_paragraph('• Pruebas ICMP (ping) para medir latencia y disponibilidad', style='List Bullet')
    doc.add_paragraph('• Solicitudes HTTP/HTTPS para verificar servicios web', style='List Bullet')
    doc.add_paragraph('• Pruebas de velocidad mediante Speedtest CLI', style='List Bullet')
    doc.add_paragraph('• Análisis de DNS mediante resolución de nombres', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Parámetros medidos:')
    doc.add_paragraph('• Latencia (tiempo de respuesta en milisegundos)', style='List Bullet')
    doc.add_paragraph('• Jitter (variación de latencia)', style='List Bullet')
    doc.add_paragraph('• Tasa de pérdida de paquetes', style='List Bullet')
    doc.add_paragraph('• Disponibilidad del servicio (porcentaje de éxito)', style='List Bullet')
    doc.add_paragraph('• Códigos de estado HTTP', style='List Bullet')
    
    doc.add_page_break()
    
    # --- ANÁLISIS DETALLADO POR SITIO ---
    doc.add_heading('3. ANÁLISIS DETALLADO POR SITIO', 1)
    
    # Tabla principal de sitios
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Light Grid Accent 1'
    add_table_border(table)
    
    header_cells = table.rows[0].cells
    headers = ['Sitio/IP', 'Disponibilidad (%)', 'Latencia Prom. (ms)', 'Latencia Mín/Máx', 'Jitter (ms)', 'Paquetes (OK/Total)', 'Estado']
    
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Ordenar sitios por categoría
    sites_by_category = {}
    for site in ping_data.keys():
        category = classify_site(site)
        if category not in sites_by_category:
            sites_by_category[category] = []
        sites_by_category[category].append(site)
    
    # Añadir datos de cada sitio
    for category in sorted(sites_by_category.keys()):
        for site in sorted(sites_by_category[category]):
            data = ping_data[site]
            if not data:
                continue
            
            stats = analyze_site_data(data)
            if not stats:
                continue
            
            row_cells = table.add_row().cells
            row_cells[0].text = site
            row_cells[1].text = f"{stats['availability']:.1f}%"
            row_cells[2].text = f"{stats['avg_latency']:.1f}" if stats['avg_latency'] else 'N/A'
            
            if stats['min_latency'] and stats['max_latency']:
                row_cells[3].text = f"{stats['min_latency']:.1f} / {stats['max_latency']:.1f}"
            else:
                row_cells[3].text = 'N/A'
            
            row_cells[4].text = f"{stats['stdev_latency']:.1f}" if stats['stdev_latency'] else 'N/A'
            row_cells[5].text = f"{stats['successful']} / {stats['total_pings']}"
            
            status = get_status_color(stats['availability'], stats['avg_latency'])
            row_cells[6].text = status
    
    # --- ANÁLISIS DE RUTAS DE RED (TRACEROUTE) ---
    doc.add_heading('3.1. Análisis de Rutas de Red (Traceroute)', 2)
    
    doc.add_paragraph('A continuación se presenta el análisis de traceroute para los sitios principales monitoreados. El traceroute muestra la ruta que siguen los paquetes desde el equipo monitor hasta el destino.')
    doc.add_paragraph()
    
    # Ejecutar traceroute para los primeros 5 sitios más relevantes
    sites_for_traceroute = []
    for site in ping_data.keys():
        if ping_data[site]:
            stats = analyze_site_data(ping_data[site])
            if stats and stats['availability'] > 0:
                sites_for_traceroute.append((site, stats['avg_latency'] or 999))
    
    # Ordenar por latencia y tomar los primeros 5
    sites_for_traceroute.sort(key=lambda x: x[1])
    sites_for_traceroute = [s[0] for s in sites_for_traceroute[:5]]
    
    if sites_for_traceroute:
        for site in sites_for_traceroute:
            doc.add_paragraph(f'Traceroute a {site}:', style='Heading 3')
            
            # Ejecutar traceroute
            trace_result = run_traceroute(site)
            
            if trace_result and 'output' in trace_result:
                # Añadir output del traceroute en fuente monospace
                trace_para = doc.add_paragraph()
                trace_run = trace_para.add_run(trace_result['output'])
                trace_run.font.name = 'Courier New'
                trace_run.font.size = Pt(9)
                
                # Análisis del traceroute
                output_lines = trace_result['output'].split('\n')
                hop_count = 0
                for line in output_lines:
                    if re.match(r'^\s*\d+', line):
                        hop_count += 1
                
                if hop_count > 0:
                    doc.add_paragraph(f'Total de saltos detectados: {hop_count}')
                    
                    if hop_count > 15:
                        doc.add_paragraph('⚠️ Ruta extensa: El sitio requiere muchos saltos, puede indicar una ruta subóptima o geográficamente distante.')
                    elif hop_count < 5:
                        doc.add_paragraph('✅ Ruta eficiente: El sitio está relativamente cerca en términos de red.')
                    else:
                        doc.add_paragraph('✓ Ruta normal: Número típico de saltos para este tipo de conexión.')
            else:
                doc.add_paragraph('No se pudo completar el traceroute para este sitio.')
            
            doc.add_paragraph()
    else:
        doc.add_paragraph('No hay sitios disponibles para análisis de traceroute.')
    
    doc.add_page_break()
    
    # --- ANÁLISIS DE VELOCIDAD ---
    doc.add_heading('4. ANÁLISIS DE RENDIMIENTO DE CONEXIÓN', 1)
    
    if speedtest_data.get('history') and len(speedtest_data['history']) > 0:
        history = speedtest_data['history']
        
        downloads = [s['download_mbps'] for s in history if 'download_mbps' in s]
        uploads = [s['upload_mbps'] for s in history if 'upload_mbps' in s]
        pings = [s['ping_ms'] for s in history if 'ping_ms' in s]
        
        if downloads and uploads and pings:
            doc.add_paragraph(f'Total de pruebas realizadas: {len(history)}')
            doc.add_paragraph()
            
            speed_table = doc.add_table(rows=4, cols=4)
            speed_table.style = 'Light Grid Accent 1'
            add_table_border(speed_table)
            
            speed_table.rows[0].cells[0].text = 'Métrica'
            speed_table.rows[0].cells[1].text = 'Promedio'
            speed_table.rows[0].cells[2].text = 'Mínimo'
            speed_table.rows[0].cells[3].text = 'Máximo'
            
            for cell in speed_table.rows[0].cells:
                cell.paragraphs[0].runs[0].font.bold = True
            
            speed_table.rows[1].cells[0].text = 'Descarga (Mbps)'
            speed_table.rows[1].cells[1].text = f'{mean(downloads):.2f}'
            speed_table.rows[1].cells[2].text = f'{min(downloads):.2f}'
            speed_table.rows[1].cells[3].text = f'{max(downloads):.2f}'
            
            speed_table.rows[2].cells[0].text = 'Subida (Mbps)'
            speed_table.rows[2].cells[1].text = f'{mean(uploads):.2f}'
            speed_table.rows[2].cells[2].text = f'{min(uploads):.2f}'
            speed_table.rows[2].cells[3].text = f'{max(uploads):.2f}'
            
            speed_table.rows[3].cells[0].text = 'Latencia (ms)'
            speed_table.rows[3].cells[1].text = f'{mean(pings):.2f}'
            speed_table.rows[3].cells[2].text = f'{min(pings):.2f}'
            speed_table.rows[3].cells[3].text = f'{max(pings):.2f}'
            
            doc.add_paragraph()
            doc.add_paragraph('Historial de pruebas de velocidad:', style='Heading 3')
            
            history_table = doc.add_table(rows=1, cols=4)
            history_table.style = 'Light Grid Accent 1'
            add_table_border(history_table)
            
            history_table.rows[0].cells[0].text = 'Fecha y Hora'
            history_table.rows[0].cells[1].text = 'Descarga (Mbps)'
            history_table.rows[0].cells[2].text = 'Subida (Mbps)'
            history_table.rows[0].cells[3].text = 'Latencia (ms)'
            
            for cell in history_table.rows[0].cells:
                cell.paragraphs[0].runs[0].font.bold = True
            
            for test in reversed(history[-10:]):  # Últimas 10 pruebas
                row = history_table.add_row()
                test_date = datetime.fromtimestamp(test['timestamp']).strftime('%d/%m/%Y %H:%M:%S')
                row.cells[0].text = test_date
                row.cells[1].text = str(test.get('download_mbps', 'N/A'))
                row.cells[2].text = str(test.get('upload_mbps', 'N/A'))
                row.cells[3].text = str(test.get('ping_ms', 'N/A'))
    else:
        doc.add_paragraph('No hay datos de pruebas de velocidad disponibles en este periodo.')
    
    doc.add_page_break()
    
    # --- CLASIFICACIÓN POR CATEGORÍAS ---
    doc.add_heading('5. ANÁLISIS POR CATEGORÍAS', 1)
    
    for category in sorted(sites_by_category.keys()):
        doc.add_heading(f'5.{list(sorted(sites_by_category.keys())).index(category) + 1}. {category}', 2)
        
        sites = sites_by_category[category]
        category_latencies = []
        category_availability = []
        
        for site in sites:
            data = ping_data.get(site, [])
            if data:
                stats = analyze_site_data(data)
                if stats:
                    if stats['avg_latency']:
                        category_latencies.append(stats['avg_latency'])
                    category_availability.append(stats['availability'])
        
        if category_latencies:
            doc.add_paragraph(f'• Sitios en esta categoría: {len(sites)}')
            doc.add_paragraph(f'• Latencia promedio: {mean(category_latencies):.2f} ms')
            doc.add_paragraph(f'• Disponibilidad promedio: {mean(category_availability):.2f}%')
        else:
            doc.add_paragraph(f'• Sitios en esta categoría: {len(sites)}')
            doc.add_paragraph('• Sin datos suficientes para análisis')
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # --- DETECCIÓN DE ANOMALÍAS ---
    doc.add_heading('6. DETECCIÓN DE ANOMALÍAS Y PROBLEMAS', 1)
    
    anomalies_found = False
    
    # Sitios con baja disponibilidad
    low_availability = []
    for site, data in ping_data.items():
        if data:
            stats = analyze_site_data(data)
            if stats and stats['availability'] < 95:
                low_availability.append((site, stats))
    
    if low_availability:
        anomalies_found = True
        doc.add_heading('6.1. Sitios con Disponibilidad Inferior a 95%', 2)
        for site, stats in sorted(low_availability, key=lambda x: x[1]['availability']):
            p = doc.add_paragraph()
            p.add_run(f'• {site}: ').bold = True
            p.add_run(f"{stats['availability']:.1f}% disponibilidad ({stats['failed']} de {stats['total_pings']} paquetes perdidos)")
    
    # Sitios con alta latencia
    high_latency = []
    for site, data in ping_data.items():
        if data:
            stats = analyze_site_data(data)
            if stats and stats['avg_latency'] and stats['avg_latency'] > 200:
                high_latency.append((site, stats))
    
    if high_latency:
        anomalies_found = True
        doc.add_heading('6.2. Sitios con Latencia Superior a 200ms', 2)
        for site, stats in sorted(high_latency, key=lambda x: x[1]['avg_latency'], reverse=True):
            p = doc.add_paragraph()
            p.add_run(f'• {site}: ').bold = True
            p.add_run(f"{stats['avg_latency']:.1f} ms promedio")
    
    # Sitios con alto jitter
    high_jitter = []
    for site, data in ping_data.items():
        if data:
            stats = analyze_site_data(data)
            if stats and stats['stdev_latency'] and stats['stdev_latency'] > 15:
                high_jitter.append((site, stats))
    
    if high_jitter:
        anomalies_found = True
        doc.add_heading('6.3. Sitios con Jitter Elevado (>15ms)', 2)
        for site, stats in sorted(high_jitter, key=lambda x: x[1]['stdev_latency'], reverse=True):
            p = doc.add_paragraph()
            p.add_run(f'• {site}: ').bold = True
            p.add_run(f"{stats['stdev_latency']:.1f} ms de variación")
    
    if not anomalies_found:
        doc.add_paragraph('✓ No se detectaron anomalías significativas en el periodo analizado.')
        doc.add_paragraph('Todos los sitios monitoreados operan dentro de parámetros normales.')
    
    doc.add_page_break()
    
    # --- CONCLUSIONES ---
    doc.add_heading('7. CONCLUSIONES TÉCNICAS', 1)
    
    conclusions = []
    
    if sites_critical > 0:
        conclusions.append(f'Se identificaron {sites_critical} sitios con disponibilidad crítica (<80%), requiriendo atención inmediata.')
    
    if avg_network_latency > 150:
        conclusions.append(f'La latencia promedio de red ({avg_network_latency} ms) está por encima del rango óptimo, sugiriendo posibles problemas de conectividad o saturación.')
    elif avg_network_latency < 50:
        conclusions.append(f'La latencia promedio de red ({avg_network_latency} ms) se encuentra en un rango excelente.')
    
    local_sites = [s for s in ping_data.keys() if s.startswith('192.168.')]
    if local_sites:
        local_stats = [analyze_site_data(ping_data[s]) for s in local_sites if ping_data.get(s)]
        latency_values = [s['avg_latency'] for s in local_stats if s and s.get('avg_latency')]
        local_avg = mean(latency_values) if latency_values else 0
        if local_avg > 10:
            conclusions.append(f'Los dispositivos en la red local presentan latencia elevada ({local_avg:.1f} ms), sugiriendo congestión en la LAN.')
    
    if speedtest_data.get('history') and len(speedtest_data['history']) > 1:
        history = speedtest_data['history']
        downloads = [s['download_mbps'] for s in history if 'download_mbps' in s]
        if downloads:
            download_variation = (max(downloads) - min(downloads)) / mean(downloads) * 100
            if download_variation > 30:
                conclusions.append(f'Se observa alta variabilidad en la velocidad de descarga ({download_variation:.1f}%), indicando posible inestabilidad del ISP.')
    
    if not conclusions:
        conclusions.append('La infraestructura de red opera de manera estable y dentro de los parámetros esperados.')
        conclusions.append('No se identificaron problemas críticos que requieran intervención inmediata.')
    
    for i, conclusion in enumerate(conclusions, 1):
        doc.add_paragraph(f'{i}. {conclusion}')

    doc.add_page_break()

    # --- PRUEBAS AVANZADAS DE RED ---
    doc.add_heading('5. ANÁLISIS DE PRUEBAS AVANZADAS', 1)

    # Agregar análisis de pruebas avanzadas
    add_advanced_tests_analysis(doc, advanced_data)

    doc.add_page_break()

    # --- ANÁLISIS DE RUTAS AVANZADO (TRACEROUTE/PATHPING) ---
    doc.add_heading('6. ANÁLISIS AVANZADO DE RUTAS Y RENDIMIENTO', 1)

    # Agregar análisis de traceroute/pathping
    add_traceroute_analysis(doc, traceroute_data)

    doc.add_page_break()

    # --- BENCHMARK DNS ---
    doc.add_heading('7. ANÁLISIS DE RENDIMIENTO DNS', 1)

    # Agregar análisis de DNS
    add_dns_benchmark_analysis(doc, dns_data)

    doc.add_page_break()

    # --- RECOMENDACIONES ---
    doc.add_heading('8. RECOMENDACIONES Y PLAN DE ACCIÓN', 1)
    
    recommendations = []
    
    if sites_critical > 0:
        recommendations.append({
            'priority': 'ALTA',
            'action': 'Revisar sitios con disponibilidad crítica',
            'detail': 'Verificar conectividad, configuración de firewall y estado de los servidores destino.'
        })
    
    if high_latency:
        recommendations.append({
            'priority': 'MEDIA',
            'action': 'Optimizar rutas de red para sitios con alta latencia',
            'detail': 'Considerar CDN, cambio de proveedor, o revisar configuración de red.'
        })
    
    if high_jitter:
        recommendations.append({
            'priority': 'MEDIA',
            'action': 'Investigar causa de jitter elevado',
            'detail': 'Verificar saturación de ancho de banda, interferencias WiFi, o problemas con el ISP.'
        })
    
    if avg_network_latency > 100:
        recommendations.append({
            'priority': 'MEDIA',
            'action': 'Evaluar plan de conectividad',
            'detail': 'Considerar upgrade de velocidad o cambio de proveedor de internet.'
        })
    
    recommendations.append({
        'priority': 'BAJA',
        'action': 'Continuar monitoreo regular',
        'detail': 'Mantener el sistema de monitoreo activo para detectar tendencias y problemas emergentes.'
    })
    
    recommendations.append({
        'priority': 'BAJA',
        'action': 'Implementar alertas automáticas',
        'detail': 'Configurar notificaciones cuando los umbrales críticos sean superados.'
    })

    # Recomendaciones basadas en pruebas avanzadas
    add_advanced_recommendations(recommendations, advanced_data, dns_data, traceroute_data)

    for i, rec in enumerate(recommendations, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. [{rec["priority"]}] ').bold = True
        p.add_run(f'{rec["action"]}: ')
        p.add_run(rec["detail"])
    
    # Guardar documento
    filename = f"Reporte_Red_{report_date.strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(REPORTS_DIR, filename)
    doc.save(filepath)
    
    return filepath, filename

# --- HILOS DE TRABAJO EN SEGUNDO PLANO ---

def get_websites():
    with websites_lock:
        if not os.path.exists(WEBSITES_FILE):
            default_websites = ["google.com", "8.8.8.8", "facebook.com", "1.1.1.1"]
            try:
                with open(WEBSITES_FILE, "w") as f:
                    for site in default_websites:
                        f.write(f"{site}\n")
                print(f"Archivo {WEBSITES_FILE} no encontrado. Creado con sitios por defecto.")
                return default_websites
            except Exception as e:
                print(f"Error crítico al crear {WEBSITES_FILE}: {e}")
                return []
        try:
            with open(WEBSITES_FILE, "r", encoding='utf-8') as f:
                sites = []
                for i, line in enumerate(f):
                    site = line.strip()
                    if site and not site.startswith('#'):
                        if re.match(r"^(?:(?:https?://)?(?:[a-zA-Z0-9](?:[a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?\.)+[a-zA-Z]{2,}|(?:\d{1,3}\.){3}\d{1,3})(?::\d+)?(?:/.*)?$", site, re.IGNORECASE):
                            sites.append(site)
                        else:
                            print(f"Advertencia: Línea {i+1} inválida omitida en {WEBSITES_FILE}: '{site}'")
                return sites
        except FileNotFoundError:
             print(f"Advertencia: El archivo {WEBSITES_FILE} desapareció. Intentando recrear.")
             return []
        except Exception as e:
            print(f"Error al leer {WEBSITES_FILE}: {e}")
            return []

def ping(website):
    """Ejecuta ping a un sitio web con detección mejorada para Windows/Linux"""
    try:
        param = "-n" if platform.system().lower() == "windows" else "-c"
        host_match = re.match(r"^(?:https?://)?([^/:]+)", website)
        if not host_match:
            return {"website": website, "timestamp": time.time(), "success": False, "error": "Formato inválido"}
        host = host_match.group(1)
        
        # Construir comando según sistema operativo
        if platform.system().lower() == "windows":
            command = ["ping", "-n", "1", "-w", "2000", host]
        else:
            command = ["ping", "-c", "1", "-W", "2", host]
        
        result = subprocess.run(command, capture_output=True, text=True, timeout=5)
        output = result.stdout + result.stderr
        
        # Detección mejorada - múltiples indicadores de éxito
        success_indicators = [
            "ttl=", "TTL=",  # Time To Live
            "time=", "Time=", "TIME=",  # Tiempo de respuesta
            "tiempo=", "Tiempo=",  # Windows en español
            "bytes=" # Otro indicador común
        ]
        
        ping_success = any(indicator in output for indicator in success_indicators)
        
        # Verificar errores explícitos
        error_indicators = [
            "unreachable", "timed out", "timeout", "failed",
            "inaccesible", "tiempo de espera agotado", "error"
        ]
        has_error = any(error.lower() in output.lower() for error in error_indicators)
        
        if ping_success and not has_error and result.returncode == 0:
            # Extraer tiempo de respuesta con múltiples patrones
            time_patterns = [
                r"time[=<\s]+(\d+\.?\d*)\s*ms",  # time=15ms
                r"tiempo[=<\s]+(\d+\.?\d*)\s*ms",  # tiempo=15ms (español)
                r"Time[=<\s]+(\d+\.?\d*)\s*ms",  # Time=15ms
            ]
            
            time_ms = None
            for pattern in time_patterns:
                time_match = re.search(pattern, output, re.IGNORECASE)
                if time_match:
                    time_ms = float(time_match.group(1))
                    break
            
            return {"website": website, "timestamp": time.time(), "success": True, "time_ms": time_ms if time_ms is not None else 0}
        else:
            return {"website": website, "timestamp": time.time(), "success": False}
            
    except subprocess.TimeoutExpired:
        return {"website": website, "timestamp": time.time(), "success": False, "error": "Timeout"}
    except Exception as e:
        return {"website": website, "timestamp": time.time(), "success": False, "error": str(e)}

def http_check(website):
    try:
        if not website.startswith("http://") and not website.startswith("https://"):
            url = f"https://{website}"
        else:
            url = website
        response = requests.get(url, timeout=5, allow_redirects=True, verify=False)
        return {"status_code": response.status_code, "success": True}
    except Exception as e:
        return {"status_code": None, "success": False, "error": str(e)}

def calculate_jitter(website_data):
    if not website_data or len(website_data) < 2:
        return 0
    successful_pings = [p for p in website_data if p.get("success", False) and "time_ms" in p]
    if len(successful_pings) < 2:
        return 0
    latencies = [p["time_ms"] for p in successful_pings[-20:]]
    if len(latencies) < 2:
        return 0
    try:
        return round(stdev(latencies), 2)
    except:
        return 0

def calculate_packet_loss(website_data):
    if not website_data:
        return 0
    total = len(website_data)
    failed = sum(1 for p in website_data if not p.get("success", False))
    return round((failed / total) * 100, 2) if total > 0 else 0

def main_check_loop():
    global ping_results_data
    last_save_time = time.time()
    try:
        if os.path.exists(JSON_OUTPUT_FILE):
            with open(JSON_OUTPUT_FILE, "r", encoding='utf-8') as f:
                temp_data = json.load(f)
                with data_lock:
                    ping_results_data = temp_data
                print(f"Datos cargados desde {JSON_OUTPUT_FILE}")
    except Exception as e:
        print(f"No se pudieron cargar datos previos: {e}")
        ping_results_data = {}
    
    iteration = 0
    while True:
        iteration += 1
        
        # Pausar si se está ejecutando speedtest
        if ping_paused_event.is_set():
            print("⏸️ Pings pausados - Speedtest en ejecución...")
            time.sleep(2)
            continue
        
        websites = get_websites()
        if not websites:
            print("No hay sitios para monitorear. Esperando...")
            time.sleep(PING_INTERVAL_SECONDS)
            continue
        
        print(f"\n--- Ronda {iteration} - Monitoreando {len(websites)} sitios ---")
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_PING_WORKERS) as executor:
            ping_futures = {executor.submit(ping, site): site for site in websites}
            http_futures = {executor.submit(http_check, site): site for site in websites}
            
            ping_res = {}
            for future in concurrent.futures.as_completed(ping_futures):
                result = future.result()
                ping_res[result["website"]] = result
            
            http_res = {}
            for future in concurrent.futures.as_completed(http_futures):
                site = http_futures[future]
                try:
                    http_res[site] = future.result()
                except Exception as e:
                    http_res[site] = {"status_code": None, "success": False, "error": str(e)}
        
        with data_lock:
            for site in websites:
                if site not in ping_results_data:
                    ping_results_data[site] = []
                
                ping_result = ping_res.get(site, {"website": site, "timestamp": time.time(), "success": False})
                http_result = http_res.get(site, {})
                
                combined_result = {**ping_result}
                if http_result.get("status_code"):
                    combined_result["http_status"] = http_result["status_code"]
                
                ping_results_data[site].append(combined_result)
                
                if len(ping_results_data[site]) > MAX_HISTORY_PER_SITE:
                    ping_results_data[site] = ping_results_data[site][-MAX_HISTORY_PER_SITE:]
        
        current_time = time.time()
        if current_time - last_save_time >= SAVE_INTERVAL_SECONDS:
            try:
                with data_lock:
                    with open(JSON_OUTPUT_FILE, "w", encoding='utf-8') as f:
                        json.dump(ping_results_data, f, indent=2)
                last_save_time = current_time
            except Exception as e:
                print(f"Error al guardar datos: {e}")
        
        if iteration == 1:
            first_ping_round_done.set()
        
        time.sleep(PING_INTERVAL_SECONDS)

def speedtest_loop():
    global speedtest_results_data, g_speedtest_interval_minutes
    try:
        if os.path.exists(SPEEDTEST_RESULTS_FILE):
            with open(SPEEDTEST_RESULTS_FILE, "r", encoding='utf-8') as f:
                temp_data = json.load(f)
                with speedtest_lock:
                    speedtest_results_data = temp_data
                print(f"Datos de speedtest cargados desde {SPEEDTEST_RESULTS_FILE}")
    except Exception as e:
        print(f"No se pudieron cargar datos previos de speedtest: {e}")
        speedtest_results_data = {"latest": None, "history": []}
    
    last_run_time = 0
    
    while True:
        current_time = time.time()
        
        # Verificar si debe ejecutarse speedtest
        if auto_speedtest_running_event.is_set():
            time_since_last_run = current_time - last_run_time
            interval_seconds = g_speedtest_interval_minutes * 60
            
            # Ejecutar si nunca se ha ejecutado o si ha pasado el intervalo
            if last_run_time == 0 or time_since_last_run >= interval_seconds:
                print(f"Ejecutando prueba de velocidad automática (intervalo: {g_speedtest_interval_minutes} min)...")
                result = run_speedtest_internal()
                if result and 'error' not in result:
                    print(f"Speedtest completado: {result.get('download_mbps', 'N/A')} Mbps down, {result.get('upload_mbps', 'N/A')} Mbps up")
                    last_run_time = current_time
                elif result:
                    print(f"Speedtest falló: {result.get('error', 'Error desconocido')}")
                    last_run_time = current_time  # Marcar como ejecutado para no reintentar inmediatamente
            else:
                # Mostrar tiempo restante
                time_remaining = interval_seconds - time_since_last_run
                minutes_remaining = int(time_remaining / 60)
                if minutes_remaining > 0 and minutes_remaining % 5 == 0:  # Log cada 5 minutos
                    print(f"Próximo speedtest automático en ~{minutes_remaining} minutos")
        else:
            # Si el auto speedtest está desactivado, resetear el timer
            last_run_time = 0
        
        # Sleep corto para ser reactivo a cambios de configuración
        time.sleep(10)  # Chequear cada 10 segundos


def run_speedtest_internal():
    """Ejecuta speedtest con manejo robusto de errores y detección automática de versión"""
    
    # ========================================
    # PAUSAR PINGS DURANTE SPEEDTEST
    # ========================================
    print("\n" + "="*60)
    print("🚀 INICIANDO SPEEDTEST - PAUSANDO PINGS")
    print("="*60)
    ping_paused_event.set()  # Pausar pings
    
    try:
        # Detectar qué versión de speedtest está instalada
        speedtest_cli_mode = False
        try:
            # Intentar con speedtest-cli (versión Python antigua)
            test_result = subprocess.run(["speedtest", "--version"], capture_output=True, text=True, timeout=5)
            if "speedtest-cli" in test_result.stdout.lower() or "--json" in test_result.stdout:
                speedtest_cli_mode = True
                print("Detectado: speedtest-cli (Python)")
            else:
                print("Detectado: speedtest (Ookla oficial)")
        except:
            print("No se pudo detectar versión, intentando speedtest-cli primero")
            speedtest_cli_mode = True
    
        # Construir comando según versión
        if speedtest_cli_mode:
            # speedtest-cli (versión Python antigua): usa --json
            command = ["speedtest", "--json"]
        else:
            # speedtest (Ookla oficial): usa --format=json
            command = ["speedtest", "--accept-license", "--accept-gdpr", "--format=json"]
        
        print(f"Ejecutando: {' '.join(command)}")
        result = subprocess.run(command, capture_output=True, text=True, timeout=120)
        
        # Verificar que hay output
        if not result.stdout or not result.stdout.strip():
            error_msg = result.stderr.strip() if result.stderr else "Sin output de speedtest"
            print(f"Speedtest sin output: {error_msg}")
            
            # Si falló con un método, intentar el otro
            if speedtest_cli_mode:
                print("Reintentando con speedtest oficial de Ookla...")
                return run_speedtest_ookla()
            else:
                return {"error": f"Speedtest no retornó datos. Instala con: pip install speedtest-cli"}
        
        # Verificar código de retorno
        if result.returncode != 0:
            print(f"Speedtest falló con código {result.returncode}")
            print(f"Error: {result.stderr}")
            
            # Si falló, intentar el otro método
            if speedtest_cli_mode:
                print("Reintentando con speedtest oficial de Ookla...")
                return run_speedtest_ookla()
            else:
                return {"error": f"Speedtest falló: {result.stderr or 'Error desconocido'}"}
        
        # Intentar parsear JSON
        try:
            data = json.loads(result.stdout)
        except json.JSONDecodeError as e:
            print(f"Error parseando JSON: {e}")
            print(f"Output (primeros 500 chars): {result.stdout[:500]}")
            return {"error": "Error parseando resultado. Verifica instalación de speedtest."}
        
        # Parsear según formato (diferente entre versiones)
        speedtest_result = None
        
        if speedtest_cli_mode:
            # Formato speedtest-cli: download/upload en bps
            if "download" in data and "upload" in data and "ping" in data:
                speedtest_result = {
                    "timestamp": time.time(),
                    "download_mbps": round(data["download"] / 1_000_000, 2),
                    "upload_mbps": round(data["upload"] / 1_000_000, 2),
                    "ping_ms": round(data["ping"], 2),
                    "server": data.get("server", {}).get("sponsor", "Unknown") if isinstance(data.get("server"), dict) else "Unknown"
                }
        else:
            # Formato Ookla: download.bandwidth en bytes/s
            if "download" in data and "upload" in data and "ping" in data:
                speedtest_result = {
                    "timestamp": time.time(),
                    "download_mbps": round(data["download"]["bandwidth"] * 8 / 1_000_000, 2),
                    "upload_mbps": round(data["upload"]["bandwidth"] * 8 / 1_000_000, 2),
                    "ping_ms": round(data["ping"]["latency"], 2),
                    "server": data.get("server", {}).get("name", "Unknown")
                }
        
        if not speedtest_result:
            print(f"Formato JSON no reconocido: {data}")
            return {"error": "Formato de respuesta no reconocido"}
        
        # Guardar con lock
        with speedtest_lock:
            speedtest_results_data["latest"] = speedtest_result
            speedtest_results_data["history"].append(speedtest_result)
            if len(speedtest_results_data["history"]) > 100:
                speedtest_results_data["history"] = speedtest_results_data["history"][-100:]
            
            try:
                with open(SPEEDTEST_RESULTS_FILE, "w", encoding='utf-8') as f:
                    json.dump(speedtest_results_data, f, indent=2)
            except Exception as e:
                print(f"Error guardando speedtest: {e}")
        
        return speedtest_result
        
    except subprocess.TimeoutExpired:
        print("Speedtest timeout después de 120 segundos")
        return {"error": "Timeout ejecutando speedtest (>2 minutos)"}
    except FileNotFoundError:
        print("Comando speedtest no encontrado")
        return {"error": "speedtest-cli no está instalado. Instala con: pip install speedtest-cli"}
    except Exception as e:
        print(f"Error inesperado en speedtest: {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        return {"error": f"Error inesperado: {str(e)}"}
    
    finally:
        # ========================================
        # REANUDAR PINGS SIEMPRE
        # ========================================
        ping_paused_event.clear()  # Reanudar pings
        print("="*60)
        print("✅ SPEEDTEST FINALIZADO - REANUDANDO PINGS")
        print("="*60 + "\n")

def run_speedtest_ookla():
    """Intenta ejecutar speedtest oficial de Ookla"""
    try:
        result = subprocess.run(
            ["speedtest", "--accept-license", "--accept-gdpr", "--format=json"],
            capture_output=True,
            text=True,
            timeout=120
        )
        
        if result.returncode != 0 or not result.stdout:
            return {"error": "Speedtest Ookla no disponible"}
        
        data = json.loads(result.stdout)
        return {
            "timestamp": time.time(),
            "download_mbps": round(data["download"]["bandwidth"] * 8 / 1_000_000, 2),
            "upload_mbps": round(data["upload"]["bandwidth"] * 8 / 1_000_000, 2),
            "ping_ms": round(data["ping"]["latency"], 2),
            "server": data.get("server", {}).get("name", "Unknown")
        }
    except:
        return {"error": "Speedtest Ookla falló"}


def load_traceroute_history():
    """Carga el historial de traceroute desde archivo"""
    global traceroute_history
    try:
        if os.path.exists(TRACEROUTE_HISTORY_FILE):
            with open(TRACEROUTE_HISTORY_FILE, "r", encoding='utf-8') as f:
                with traceroute_history_lock:
                    traceroute_history = json.load(f)
                print(f"Historial de traceroute cargado: {len(traceroute_history)} entradas")
    except Exception as e:
        print(f"Error cargando historial de traceroute: {e}")
        traceroute_history = []

def save_traceroute_to_history(website, result, use_pathping):
    """Guarda un resultado de traceroute en el historial"""
    global traceroute_history

    entry = {
        "timestamp": time.time(),
        "date": datetime.now().isoformat(),
        "website": website,
        "method": "pathping" if use_pathping else "traceroute",
        "output": result.get("output", ""),
        "error": result.get("error", False)
    }

    with traceroute_history_lock:
        traceroute_history.append(entry)

        # Mantener solo los últimos 100 resultados
        if len(traceroute_history) > 100:
            traceroute_history = traceroute_history[-100:]

        # Guardar a archivo
        try:
            with open(TRACEROUTE_HISTORY_FILE, "w", encoding='utf-8') as f:
                json.dump(traceroute_history, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f"Error guardando historial de traceroute: {e}")

def get_traceroute_history(website=None, limit=10):
    """Obtiene el historial de traceroute, opcionalmente filtrado por website"""
    with traceroute_history_lock:
        history = traceroute_history.copy()

    if website:
        history = [entry for entry in history if entry["website"] == website]

    # Ordenar por timestamp descendente (más reciente primero)
    history.sort(key=lambda x: x["timestamp"], reverse=True)

    return history[:limit]

def run_traceroute(website, use_pathping=False):
    try:
        host_match = re.match(r"^(?:https?://)?([^/:]+)", website)
        if not host_match:
            return {"output": "Formato de sitio inválido"}
        host = host_match.group(1)

        if platform.system().lower() == "windows":
            if use_pathping:
                # PathPing: Combina ping y traceroute con estadísticas de pérdida
                command = ["pathping", "-n", "-q", "100", "-p", "500", host]
                timeout = 180  # PathPing puede tardar más tiempo
            else:
                # Tracert tradicional
                command = ["tracert", "-w", "2000", host]
                timeout = 120
        else:
            # En Linux/Mac no hay PathPing, usar traceroute o mtr
            if use_pathping:
                # Usar MTR como equivalente más cercano a PathPing
                command = ["mtr", "--report", "--report-cycles", "10", "--no-dns", host]
                timeout = 120
            else:
                command = ["traceroute", "-m", "20", "-w", "2", host]
                timeout = 120

        result = subprocess.run(command, capture_output=True, text=True, timeout=timeout)

        output = result.stdout if result.stdout else result.stderr

        # Agregar información del método usado
        method_info = ""
        if use_pathping:
            if platform.system().lower() == "windows":
                method_info = "=== PATHPING (Windows) ===\n"
                method_info += "PathPing combina traceroute con estadísticas de pérdida de paquetes por salto.\n"
                method_info += "Muestra latencia y % de pérdida para cada router en la ruta.\n\n"
            else:
                method_info = "=== MTR (Equivalente a PathPing en Linux/Mac) ===\n"
                method_info += "MTR proporciona estadísticas similares a PathPing con pérdida de paquetes.\n\n"
        else:
            method_info = f"=== TRACEROUTE ({platform.system()}) ===\n"
            method_info += "Muestra la ruta que siguen los paquetes hasta el destino.\n\n"

        result = {"output": method_info + output, "method": "pathping" if use_pathping else "traceroute"}

        # Guardar en historial
        save_traceroute_to_history(website, result, use_pathping)

        return result

    except subprocess.TimeoutExpired:
        method = "PathPing" if use_pathping else "Traceroute"
        result = {"output": f"Timeout: El {method} tardó demasiado.", "error": True}
        save_traceroute_to_history(website, result, use_pathping)
        return result
    except FileNotFoundError:
        if use_pathping and platform.system().lower() != "windows":
            result = {"output": "MTR no está instalado. Instala con: sudo apt install mtr (Ubuntu/Debian) o brew install mtr (Mac)", "error": True}
        else:
            result = {"output": "Comando no encontrado. Verifica que las herramientas de red estén instaladas.", "error": True}
        save_traceroute_to_history(website, result, use_pathping)
        return result
    except Exception as e:
        result = {"output": f"Error: {str(e)}", "error": True}
        save_traceroute_to_history(website, result, use_pathping)
        return result

# --- BENCHMARK DNS ---
def dns_benchmark():
    dns_servers = [
        ("Google DNS 1", "8.8.8.8"),
        ("Google DNS 2", "8.8.4.4"),
        ("Cloudflare 1", "1.1.1.1"),
        ("Cloudflare 2", "1.0.0.1"),
        ("OpenDNS 1", "208.67.222.222"),
        ("OpenDNS 2", "208.67.220.220"),
    ]
    
    test_domains = ["google.com", "facebook.com", "amazon.com", "youtube.com", "wikipedia.org"]
    
    results = []
    
    for name, ip in dns_servers:
        resolver = dns.resolver.Resolver(configure=False)
        resolver.nameservers = [ip]
        resolver.timeout = 2
        resolver.lifetime = 2
        
        times = []
        for domain in test_domains:
            try:
                start = time.time()
                resolver.resolve(domain, 'A')
                elapsed = (time.time() - start) * 1000
                times.append(elapsed)
            except:
                pass
        
        if times:
            avg_time = round(sum(times) / len(times), 2)
            results.append({"name": name, "ip": ip, "avg_time_ms": avg_time, "success_rate": len(times) / len(test_domains) * 100})
        else:
            results.append({"name": name, "ip": ip, "avg_time_ms": None, "success_rate": 0})
    
    results.sort(key=lambda x: x["avg_time_ms"] if x["avg_time_ms"] is not None else float('inf'))
    
    return results

def dns_benchmark_thread():
    global dns_benchmark_results
    with dns_benchmark_lock:
        dns_benchmark_results["status"] = "running"
    
    try:
        results = dns_benchmark()
        with dns_benchmark_lock:
            dns_benchmark_results["results"] = results
            dns_benchmark_results["last_run"] = time.time()
            dns_benchmark_results["status"] = "completed"
    except Exception as e:
        with dns_benchmark_lock:
            dns_benchmark_results["status"] = "error"
            dns_benchmark_results["error"] = str(e)

# --- CONFIGURACIÓN FLASK ---
app = Flask(__name__)

@app.route("/")
def index():
    return render_template("index.html", ping_interval=PING_INTERVAL_SECONDS)

@app.route("/api/ping-results")
def api_ping_results():
    with data_lock:
        return jsonify(ping_results_data)

@app.route("/api/ping-results/<path:website>")
def api_ping_results_single(website):
    with data_lock:
        if website in ping_results_data:
            return jsonify({website: ping_results_data[website]})
        else:
            abort(404)

@app.route("/api/speedtest-results")
def api_speedtest_results():
    with speedtest_lock:
        return jsonify({"data": speedtest_results_data, "auto_enabled": auto_speedtest_running_event.is_set(), "interval_minutes": g_speedtest_interval_minutes})

@app.route("/api/run-speedtest", methods=["POST"])
def api_run_speedtest():
    result = run_speedtest_internal()
    if result and "error" not in result:
        return jsonify({"success": True, "data": result})
    else:
        return jsonify({"success": False, "error": result.get("error", "Unknown error")}), 500

@app.route("/api/traceroute/<path:website>")
def api_traceroute(website):
    # Obtener parámetro de método (traceroute o pathping)
    use_pathping = request.args.get('method', 'traceroute').lower() == 'pathping'
    result = run_traceroute(website, use_pathping)
    return jsonify(result)

@app.route("/api/traceroute-history")
def api_traceroute_history():
    """API para obtener historial de traceroute"""
    website = request.args.get('website')
    limit = int(request.args.get('limit', 10))

    history = get_traceroute_history(website, limit)
    return jsonify({"history": history, "total": len(history)})

@app.route("/api/traceroute-history/<int:entry_id>")
def api_traceroute_history_detail(entry_id):
    """API para obtener detalles de una entrada específica del historial"""
    with traceroute_history_lock:
        if 0 <= entry_id < len(traceroute_history):
            return jsonify(traceroute_history[entry_id])

    return jsonify({"error": "Entrada no encontrada"}), 404

@app.route("/api/traceroute-history", methods=["DELETE"])
def api_traceroute_history_clear():
    """API para limpiar el historial de traceroute"""
    global traceroute_history

    with traceroute_history_lock:
        traceroute_history = []

        # Eliminar archivo
        try:
            if os.path.exists(TRACEROUTE_HISTORY_FILE):
                os.remove(TRACEROUTE_HISTORY_FILE)
        except Exception as e:
            return jsonify({"success": False, "error": str(e)}), 500

    return jsonify({"success": True, "message": "Historial limpiado"})

@app.route("/api/dns-benchmark")
def api_dns_benchmark():
    with dns_benchmark_lock:
        return jsonify(dns_benchmark_results)

@app.route("/api/dns-benchmark/run", methods=["POST"])
def api_dns_benchmark_run():
    with dns_benchmark_lock:
        if dns_benchmark_results["status"] == "running":
            return jsonify({"success": False, "message": "Ya hay un benchmark en ejecución"}), 400
    
    thread = threading.Thread(target=dns_benchmark_thread, daemon=True)
    thread.start()
    return jsonify({"success": True, "message": "Benchmark DNS iniciado"})

@app.route("/api/toggle-auto-speedtest", methods=["POST"])
def api_toggle_auto_speedtest():
    global auto_speedtest_running_event
    if auto_speedtest_running_event.is_set():
        auto_speedtest_running_event.clear()
        return jsonify({"success": True, "enabled": False, "message": "Speedtest automático desactivado"})
    else:
        auto_speedtest_running_event.set()
        return jsonify({"success": True, "enabled": True, "message": "Speedtest automático activado"})

@app.route("/api/set-speedtest-interval", methods=["POST"])
def api_set_speedtest_interval():
    global g_speedtest_interval_minutes
    data = request.get_json()
    minutes = data.get("minutes")
    if minutes and isinstance(minutes, (int, float)) and 1 <= minutes <= 1440:
        g_speedtest_interval_minutes = int(minutes)
        print(f"Intervalo de speedtest actualizado a {g_speedtest_interval_minutes} minutos")
        return jsonify({"success": True, "interval_minutes": g_speedtest_interval_minutes})
    else:
        return jsonify({"success": False, "error": "Intervalo inválido (1-1440 minutos)"}), 400

@app.route("/api/get-speedtest-interval", methods=["GET"])
def api_get_speedtest_interval():
    """Obtener el intervalo actual de speedtest automático"""
    return jsonify({
        "success": True,
        "interval_minutes": g_speedtest_interval_minutes,
        "auto_enabled": auto_speedtest_running_event.is_set()
    })

@app.route("/api/dashboard-stats")
def api_dashboard_stats():
    with data_lock:
        websites = list(ping_results_data.keys())
        total_sites = len(websites)
        online_sites = 0
        offline_sites = 0
        avg_latency_list = []
        
        for site in websites:
            site_data = ping_results_data[site]
            if site_data:
                last_ping = site_data[-1]
                if last_ping.get("success", False):
                    online_sites += 1
                    if "time_ms" in last_ping:
                        avg_latency_list.append(last_ping["time_ms"])
                else:
                    offline_sites += 1
        
        avg_latency = round(sum(avg_latency_list) / len(avg_latency_list), 2) if avg_latency_list else 0
        
    with speedtest_lock:
        latest_speed = speedtest_results_data.get("latest")
    
    # Manejar caso cuando latest_speed es None
    if latest_speed is None:
        latest_speed = {}
    
    return jsonify({
        "total_sites": total_sites,
        "online_sites": online_sites,
        "offline_sites": offline_sites,
        "avg_latency": avg_latency,
        "latest_download": latest_speed.get("download_mbps", 0),
        "latest_upload": latest_speed.get("upload_mbps", 0),
        "ping_paused": ping_paused_event.is_set()  # Indicar si pings están pausados
    })

@app.route("/api/websites", methods=["GET"])
def api_get_websites():
    sites = get_websites()
    return jsonify(sites)

@app.route("/api/websites", methods=["POST"])
def api_save_websites():
    data = request.get_json()
    new_websites = data.get("websites", [])
    
    if not isinstance(new_websites, list):
        return jsonify({"success": False, "message": "Formato inválido"}), 400
    
    with websites_lock:
        try:
            with open(WEBSITES_FILE, "w", encoding='utf-8') as f:
                for site in new_websites:
                    f.write(f"{site}\n")
            return jsonify({"success": True, "message": "Sitios actualizados correctamente"})
        except Exception as e:
            return jsonify({"success": False, "message": str(e)}), 500

@app.route("/history/<path:website>")
def history(website):
    with data_lock:
        if website not in ping_results_data:
            abort(404)
    return render_template("history.html", website=website)

@app.route("/speedtest-history")
def speedtest_history():
    return render_template("speedtest_history.html")

# --- NUEVO ENDPOINT PARA GENERAR REPORTE ---
@app.route("/api/generate-report", methods=["POST"])
def api_generate_report():
    """Endpoint para generar reporte técnico"""
    try:
        print("Iniciando generación de reporte...")
        filepath, filename = generate_technical_report()
        print(f"Reporte generado exitosamente: {filename}")
        return jsonify({
            "success": True,
            "message": "Reporte generado exitosamente",
            "filename": filename,
            "download_url": f"/api/download-report/{filename}"
        })
    except Exception as e:
        print(f"Error generando reporte: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@app.route("/api/download-report/<filename>")
def api_download_report(filename):
    """Endpoint para descargar el reporte generado"""
    try:
        filepath = os.path.join(REPORTS_DIR, filename)
        if not os.path.exists(filepath):
            abort(404)
        return send_file(filepath, as_attachment=True, download_name=filename)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# --- PLANTILLAS HTML ---
def create_template_files():
    if not os.path.exists("templates"):
        os.makedirs("templates")
    
    with open("templates/index.html", "w", encoding="utf-8") as f:
        f.write(r'''
<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Monitor de Red Avanzado</title>
    <style>
        :root {
            --bg-color: #f4f7f9;
            --text-color: #333;
            --card-bg: #fff;
            --table-header-bg: #34495e;
            --table-row-hover: #f1f3f5;
            --success-color: #2ecc71;
            --danger-color: #e74c3c;
            --warning-color: #f39c12;
            --info-color: #3498db;
            --primary-color: #3498db;
            --border-color: #ddd;
            --link-color: #2980b9;
        }
        [data-theme="dark"] {
            --bg-color: #2c3e50;
            --text-color: #ecf0f1;
            --card-bg: #34495e;
            --table-header-bg: #2c3e50;
            --table-row-hover: #3e5771;
            --success-color: #27ae60;
            --danger-color: #c0392b;
            --warning-color: #e67e22;
            --info-color: #5dade2;
            --primary-color: #5dade2;
            --border-color: #4a627a;
            --link-color: #5dade2;
        }
        * { box-sizing: border-box; margin: 0; padding: 0; }
        body {
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
            background-color: var(--bg-color);
            color: var(--text-color);
            padding: 20px;
            transition: background-color 0.3s, color 0.3s;
        }
        header {
            text-align: center;
            margin-bottom: 30px;
        }
        header h1 {
            font-size: 2.5rem;
            color: var(--info-color);
            margin-bottom: 10px;
        }
        .header-actions {
            display: flex;
            justify-content: center;
            gap: 10px;
            flex-wrap: wrap;
            margin-top: 15px;
        }
        button {
            background-color: var(--link-color);
            color: #fff;
            border: none;
            padding: 10px 15px;
            border-radius: 5px;
            cursor: pointer;
            font-weight: bold;
            transition: opacity 0.3s;
        }
        button:hover { opacity: 0.8; }
        button:disabled {
            background-color: #95a5a6;
            cursor: not-allowed;
        }
        .btn-danger { background-color: var(--danger-color); }
        .btn-success { background-color: var(--success-color); }
        .btn-warning { background-color: var(--warning-color); }
        .dashboard {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 20px;
            margin-bottom: 30px;
        }
        .card {
            background-color: var(--card-bg);
            padding: 20px;
            border-radius: 8px;
            box-shadow: 0 2px 15px rgba(0,0,0,0.1);
            text-align: center;
        }
        .card h3 {
            font-size: 0.9rem;
            color: var(--text-color);
            margin-bottom: 10px;
            text-transform: uppercase;
        }
        .card .value {
            font-size: 2rem;
            font-weight: bold;
            color: var(--info-color);
        }
        table {
            width: 100%;
            border-collapse: collapse;
            background-color: var(--card-bg);
            box-shadow: 0 2px 15px rgba(0,0,0,0.1);
            border-radius: 8px;
            overflow: hidden;
        }
        thead {
            background-color: var(--table-header-bg);
            color: #fff;
        }
        th, td {
            padding: 12px 15px;
            text-align: left;
            border-bottom: 1px solid var(--border-color);
        }
        th.sortable {
            cursor: pointer;
            user-select: none;
        }
        th.sortable:hover {
            background-color: rgba(255,255,255,0.1);
        }
        tbody tr:hover {
            background-color: var(--table-row-hover);
        }
        .status-ok { color: var(--success-color); font-weight: bold; }
        .status-fail { color: var(--danger-color); font-weight: bold; }
        .http-ok { color: var(--success-color); }
        .http-error { color: var(--danger-color); }
        .http-warn { color: var(--warning-color); }
        .clickable-row {
            cursor: pointer;
        }
        .clickable-row:hover {
            background-color: var(--table-row-hover);
        }
        .modal {
            display: none;
            position: fixed;
            z-index: 1;
            left: 0;
            top: 0;
            width: 100%;
            height: 100%;
            overflow: auto;
            background-color: rgba(0,0,0,0.6);
        }
        .modal-content {
            background-color: var(--card-bg);
            margin: 5% auto;
            padding: 20px;
            border: 1px solid var(--border-color);
            width: 90%;
            max-width: 600px;
            border-radius: 8px;
            box-shadow: 0 4px 20px rgba(0,0,0,0.3);
        }
        .close {
            color: #aaa;
            float: right;
            font-size: 28px;
            font-weight: bold;
            cursor: pointer;
        }
        .close:hover,
        .close:focus {
            color: var(--text-color);
        }
        #sitesList {
            list-style: none;
            max-height: 300px;
            overflow-y: auto;
            border: 1px solid var(--border-color);
            border-radius: 5px;
            padding: 10px;
            margin: 10px 0;
        }
        #sitesList li {
            padding: 8px;
            border-bottom: 1px solid var(--border-color);
            display: flex;
            justify-content: space-between;
            align-items: center;
        }
        #sitesList li:last-child {
            border-bottom: none;
        }
        .delete-site {
            color: var(--danger-color);
            cursor: pointer;
            font-size: 1.5rem;
            font-weight: bold;
        }
        .delete-site:hover {
            opacity: 0.7;
        }
        input[type="text"] {
            width: 100%;
            padding: 10px;
            border: 1px solid var(--border-color);
            border-radius: 5px;
            background-color: var(--card-bg);
            color: var(--text-color);
            margin: 10px 0;
        }
        .speedtest-section {
            background-color: var(--card-bg);
            padding: 20px;
            border-radius: 8px;
            box-shadow: 0 2px 15px rgba(0,0,0,0.1);
            margin-bottom: 20px;
        }
        .speedtest-section h2 {
            margin-bottom: 15px;
        }
        .speedtest-result {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(150px, 1fr));
            gap: 15px;
            margin: 15px 0;
        }
        .speedtest-metric {
            text-align: center;
            padding: 10px;
            background-color: var(--bg-color);
            border-radius: 5px;
        }
        .speedtest-metric .label {
            font-size: 0.9rem;
            color: var(--text-color);
            opacity: 0.8;
        }
        .speedtest-metric .value {
            font-size: 1.5rem;
            font-weight: bold;
            color: var(--info-color);
            margin-top: 5px;
        }
        .dns-section {
            background-color: var(--card-bg);
            padding: 20px;
            border-radius: 8px;
            box-shadow: 0 2px 15px rgba(0,0,0,0.1);
            margin-bottom: 20px;
        }
        .dns-table {
            width: 100%;
            margin-top: 15px;
        }
        /* Estilos para gestión de backups */
        #backupsList table {
            width: 100%;
            border-collapse: collapse;
            margin-top: 15px;
        }
        #backupsList th,
        #backupsList td {
            padding: 10px;
            text-align: left;
            border-bottom: 1px solid var(--border-color);
        }
        #backupsList th {
            background-color: var(--table-header-bg);
            color: white;
        }
        #backupsList tr:hover {
            background-color: var(--table-row-hover);
        }
        .btn-small {
            padding: 5px 10px;
            margin: 2px;
            font-size: 0.85em;
        }
    </style>
</head>
<body>
    <header>
        <h1>🌐 Monitor de Red Avanzado</h1>
        <p>Monitoreo en tiempo real de conectividad y rendimiento</p>
        <div class="header-actions">
            <button id="themeToggle">🌓 Cambiar Tema</button>
            <button id="manageSitesBtn">⚙️ Gestionar Sitios</button>
            <button id="advancedTestsBtn" class="btn-warning">🔬 Pruebas Avanzadas</button>
            <button id="generateReportBtn" class="btn-success">📄 Generar Reporte Técnico</button>
        </div>
        
        <!-- Indicador de pausa de pings -->
        <div id="pingPausedIndicator" style="display:none; background-color:#fff3cd; color:#856404; padding:10px; margin-top:15px; border-radius:5px; border:2px solid #ffc107; text-align:center; font-weight:bold;">
            ⏸️ PINGS PAUSADOS - Speedtest en ejecución...
        </div>
    </header>

    <div class="dashboard">
        <div class="card">
            <h3>📊 Total de Sitios</h3>
            <div class="value" id="totalSites">0</div>
        </div>
        <div class="card">
            <h3>✅ En Línea</h3>
            <div class="value" id="onlineSites" style="color: var(--success-color);">0</div>
        </div>
        <div class="card">
            <h3>❌ Fuera de Línea</h3>
            <div class="value" id="offlineSites" style="color: var(--danger-color);">0</div>
        </div>
        <div class="card">
            <h3>⏱️ Latencia Promedio</h3>
            <div class="value" id="avgLatency">0 ms</div>
        </div>
    </div>

    <div class="speedtest-section">
        <h2>🚀 Prueba de Velocidad</h2>
        <button id="runSpeedtest" class="btn-success">▶️ Ejecutar Prueba</button>
        <button id="viewHistoryBtn" class="btn-warning">📜 Ver Historial</button>
        <button id="toggleAutoSpeedtest">🔄 Auto: OFF</button>
        
        <!-- Configuración de Intervalo Speedtest -->
        <div id="autoSpeedtestConfig" style="display:none; margin-top:15px; padding:15px; background-color:var(--bg-color); border-radius:5px; border:1px solid var(--border-color);">
            <label for="speedtestInterval" style="margin-right:10px; font-weight:500;">
                ⏱️ Intervalo de pruebas automáticas:
            </label>
            <input type="number" id="speedtestInterval" min="1" max="1440" value="30" style="width:80px; padding:5px; border-radius:3px; border:1px solid var(--border-color); background-color:var(--card-bg); color:var(--text-color);">
            <span style="margin:0 10px;">minutos</span>
            <button id="saveIntervalBtn" class="btn-success" style="padding:8px 15px;">💾 Guardar</button>
            <span id="intervalStatus" style="margin-left:10px; font-style:italic;"></span>
        </div>
        
        <div class="speedtest-result" id="speedtestResult" style="display:none;">
            <div class="speedtest-metric">
                <div class="label">Descarga</div>
                <div class="value" id="downloadSpeed">-- Mbps</div>
            </div>
            <div class="speedtest-metric">
                <div class="label">Subida</div>
                <div class="value" id="uploadSpeed">-- Mbps</div>
            </div>
            <div class="speedtest-metric">
                <div class="label">Latencia</div>
                <div class="value" id="pingSpeed">-- ms</div>
            </div>
        </div>
    </div>

    <!-- GESTIÓN DE DATOS -->
    <div class="data-management-section" style="background-color:var(--card-bg); padding:20px; border-radius:8px; box-shadow:0 2px 15px rgba(0,0,0,.1); margin-bottom:20px;">
        <h2>💾 Gestión de Datos</h2>
        
        <div style="display:flex; gap:10px; flex-wrap:wrap; margin:15px 0;">
            <button id="createBackupBtn" class="btn-success">📦 Crear Backup</button>
            <button id="manageBackupsBtn" class="btn-warning">📂 Gestionar Backups</button>
            <button id="clearDataBtn" class="btn-danger">🗑️ Limpiar Datos</button>
        </div>
        
        <div id="backupStatus" style="margin-top:10px; font-style:italic; padding:10px; border-radius:5px;"></div>
    </div>

    <!-- Modal Gestión de Backups -->
    <div id="backupModal" class="modal">
        <div class="modal-content">
            <span class="close" id="closeBackupModal">&times;</span>
            <h2>📂 Gestión de Backups</h2>
            <div id="backupsList" style="max-height:400px; overflow-y:auto;">
                <p>Cargando...</p>
            </div>
        </div>
    </div>

    <div class="dns-section">
        <h2>🔍 Benchmark DNS</h2>
        <button id="runDnsBenchmark" class="btn-info">▶️ Ejecutar Benchmark DNS</button>
        <div id="dnsBenchmarkStatus"></div>
        <table class="dns-table" id="dnsTable" style="display:none;">
            <thead>
                <tr>
                    <th>Servidor DNS</th>
                    <th>IP</th>
                    <th>Tiempo Promedio (ms)</th>
                    <th>Tasa de Éxito</th>
                </tr>
            </thead>
            <tbody id="dnsTableBody"></tbody>
        </table>
    </div>

    <h2 style="margin: 30px 0 15px 0;">📡 Estado de Conectividad</h2>
    <table id="pingTable">
        <thead>
            <tr>
                <th class="sortable" data-sort="website">Sitio Web / IP ▲</th>
                <th class="sortable" data-sort="ping_status">Estado Ping</th>
                <th class="sortable" data-sort="http_status">Estado HTTP</th>
                <th class="sortable" data-sort="latency">Latencia (ms)</th>
                <th class="sortable" data-sort="jitter">Jitter (ms)</th>
                <th class="sortable" data-sort="packet_loss">Pérdida (%)</th>
                <th>Enviados / Perdidos</th>
            </tr>
        </thead>
        <tbody id="pingTableBody">
            <tr><td colspan="7" style="text-align:center;">Cargando datos...</td></tr>
        </tbody>
    </table>

    <!-- Modal para gestionar sitios -->
    <div id="sitesModal" class="modal">
        <div class="modal-content">
            <span class="close" id="closeSitesModal">&times;</span>
            <h2>⚙️ Gestionar Sitios Monitoreados</h2>
            <input type="text" id="newSiteInput" placeholder="Ingresa un sitio o IP (ej: google.com, 8.8.8.8)">
            <button id="addSiteBtn" class="btn-success">➕ Agregar Sitio</button>
            <ul id="sitesList"></ul>
            <button id="saveSitesBtn" class="btn-success">💾 Guardar Cambios</button>
        </div>
    </div>

    <script>
        const PING_INTERVAL_SECONDS = {{ ping_interval }};
        let sortColumn = 'website';
        let sortDirection = 'asc';
        let autoSpeedtestEnabled = false;

        // Theme Toggle
        const themeToggle = document.getElementById('themeToggle');
        const currentTheme = localStorage.getItem('theme') || 'light';
        document.body.dataset.theme = currentTheme;

        themeToggle.addEventListener('click', () => {
            const newTheme = document.body.dataset.theme === 'dark' ? 'light' : 'dark';
            document.body.dataset.theme = newTheme;
            localStorage.setItem('theme', newTheme);
        });

        // Advanced Tests
        const advancedTestsBtn = document.getElementById('advancedTestsBtn');
        advancedTestsBtn.addEventListener('click', () => {
            window.location.href = '/advanced-tests';
        });

        // Generate Report
        const generateReportBtn = document.getElementById('generateReportBtn');
        generateReportBtn.addEventListener('click', async () => {
            generateReportBtn.disabled = true;
            generateReportBtn.textContent = '⏳ Generando reporte...';
            
            try {
                const response = await fetch('/api/generate-report', { method: 'POST' });
                const result = await response.json();
                
                if (result.success) {
                    alert('✅ Reporte generado exitosamente: ' + result.filename);
                    // Descargar automáticamente
                    window.location.href = result.download_url;
                } else {
                    alert('❌ Error al generar reporte: ' + result.error);
                }
            } catch (error) {
                alert('❌ Error de conexión: ' + error.message);
            } finally {
                generateReportBtn.disabled = false;
                generateReportBtn.textContent = '📄 Generar Reporte Técnico';
            }
        });

        // Dashboard Update
        async function updateDashboard() {
            try {
                const response = await fetch('/api/dashboard-stats');
                const data = await response.json();
                
                document.getElementById('totalSites').textContent = data.total_sites;
                document.getElementById('onlineSites').textContent = data.online_sites;
                document.getElementById('offlineSites').textContent = data.offline_sites;
                document.getElementById('avgLatency').textContent = data.avg_latency + ' ms';
                
                // Actualizar indicador de pausa de pings
                const pingPausedIndicator = document.getElementById('pingPausedIndicator');
                if (data.ping_paused) {
                    pingPausedIndicator.style.display = 'block';
                } else {
                    pingPausedIndicator.style.display = 'none';
                }
            } catch (error) {
                console.error('Error actualizando dashboard:', error);
            }
        }

        // Ping Results Update
        async function updatePingResults() {
            try {
                const response = await fetch('/api/ping-results');
                const data = await response.json();
                
                const tbody = document.getElementById('pingTableBody');
                tbody.innerHTML = '';
                
                const sites = Object.keys(data);
                if (sites.length === 0) {
                    tbody.innerHTML = '<tr><td colspan="7" style="text-align:center;">No hay datos disponibles</td></tr>';
                    return;
                }

                // Sort
                sites.sort((a, b) => {
                    const dataA = data[a];
                    const dataB = data[b];
                    
                    let valA, valB;
                    
                    if (sortColumn === 'website') {
                        valA = a.toLowerCase();
                        valB = b.toLowerCase();
                    } else if (sortColumn === 'ping_status') {
                        valA = dataA.length > 0 && dataA[dataA.length - 1].success ? 1 : 0;
                        valB = dataB.length > 0 && dataB[dataB.length - 1].success ? 1 : 0;
                    } else if (sortColumn === 'latency') {
                        const lastA = dataA.length > 0 ? dataA[dataA.length - 1] : {};
                        const lastB = dataB.length > 0 ? dataB[dataB.length - 1] : {};
                        valA = lastA.success && lastA.time_ms ? lastA.time_ms : 999999;
                        valB = lastB.success && lastB.time_ms ? lastB.time_ms : 999999;
                    } else {
                        valA = 0;
                        valB = 0;
                    }
                    
                    if (sortDirection === 'asc') {
                        return valA < valB ? -1 : valA > valB ? 1 : 0;
                    } else {
                        return valA > valB ? -1 : valA < valB ? 1 : 0;
                    }
                });

                sites.forEach(site => {
                    const siteData = data[site];
                    if (siteData.length === 0) return;
                    
                    const lastPing = siteData[siteData.length - 1];
                    const pingStatus = lastPing.success ? 'OK' : 'FALLIDO';
                    const pingClass = lastPing.success ? 'status-ok' : 'status-fail';
                    
                    const latency = lastPing.success && lastPing.time_ms ? lastPing.time_ms.toFixed(1) : 'N/A';
                    
                    // Calculate jitter
                    let jitter = 0;
                    const successfulPings = siteData.filter(p => p.success && p.time_ms);
                    if (successfulPings.length >= 2) {
                        const latencies = successfulPings.slice(-20).map(p => p.time_ms);
                        const mean = latencies.reduce((a, b) => a + b, 0) / latencies.length;
                        const variance = latencies.reduce((a, b) => a + Math.pow(b - mean, 2), 0) / latencies.length;
                        jitter = Math.sqrt(variance).toFixed(2);
                    }
                    
                    // Packet loss
                    const totalPings = siteData.length;
                    const failedPings = siteData.filter(p => !p.success).length;
                    const packetLoss = ((failedPings / totalPings) * 100).toFixed(1);
                    
                    const httpStatus = lastPing.http_status || 'N/A';
                    let httpClass = '';
                    if (httpStatus >= 200 && httpStatus < 300) httpClass = 'http-ok';
                    else if (httpStatus >= 400) httpClass = 'http-error';
                    else if (httpStatus >= 300) httpClass = 'http-warn';
                    
                    const row = document.createElement('tr');
                    row.className = 'clickable-row';
                    row.onclick = () => window.location.href = `/history/${encodeURIComponent(site)}`;
                    
                    row.innerHTML = `
                        <td>${site}</td>
                        <td class="${pingClass}">${pingStatus}</td>
                        <td class="${httpClass}">${httpStatus}</td>
                        <td>${latency}</td>
                        <td>${jitter}</td>
                        <td>${packetLoss}%</td>
                        <td>${totalPings - failedPings} / ${failedPings}</td>
                    `;
                    
                    tbody.appendChild(row);
                });
                
            } catch (error) {
                console.error('Error actualizando tabla:', error);
            }
        }

        // Speedtest
        const runSpeedtestBtn = document.getElementById('runSpeedtest');
        const speedtestResult = document.getElementById('speedtestResult');
        
        runSpeedtestBtn.addEventListener('click', async () => {
            runSpeedtestBtn.disabled = true;
            runSpeedtestBtn.textContent = '⏳ Ejecutando...';
            speedtestResult.style.display = 'none';
            
            try {
                const response = await fetch('/api/run-speedtest', { method: 'POST' });
                const result = await response.json();
                
                if (result.success) {
                    document.getElementById('downloadSpeed').textContent = result.data.download_mbps + ' Mbps';
                    document.getElementById('uploadSpeed').textContent = result.data.upload_mbps + ' Mbps';
                    document.getElementById('pingSpeed').textContent = result.data.ping_ms + ' ms';
                    speedtestResult.style.display = 'grid';
                } else {
                    alert('Error: ' + result.error);
                }
            } catch (error) {
                alert('Error de conexión: ' + error.message);
            } finally {
                runSpeedtestBtn.disabled = false;
                runSpeedtestBtn.textContent = '▶️ Ejecutar Prueba';
            }
        });

        // Auto Speedtest Toggle
        const toggleAutoSpeedtestBtn = document.getElementById('toggleAutoSpeedtest');
        toggleAutoSpeedtestBtn.addEventListener('click', async () => {
            try {
                const response = await fetch('/api/toggle-auto-speedtest', { method: 'POST' });
                const result = await response.json();
                
                autoSpeedtestEnabled = result.enabled;
                toggleAutoSpeedtestBtn.textContent = autoSpeedtestEnabled ? '🔄 Auto: ON' : '🔄 Auto: OFF';
                toggleAutoSpeedtestBtn.className = autoSpeedtestEnabled ? 'btn-success' : '';
                
                // Mostrar/ocultar configuración de intervalo
                const autoSpeedtestConfig = document.getElementById('autoSpeedtestConfig');
                if (autoSpeedtestConfig) {
                    autoSpeedtestConfig.style.display = autoSpeedtestEnabled ? 'block' : 'none';
                }
            } catch (error) {
                console.error('Error toggling auto speedtest:', error);
            }
        });

        // DNS Benchmark
        const runDnsBenchmarkBtn = document.getElementById('runDnsBenchmark');
        const dnsBenchmarkStatus = document.getElementById('dnsBenchmarkStatus');
        const dnsTable = document.getElementById('dnsTable');
        const dnsTableBody = document.getElementById('dnsTableBody');

        runDnsBenchmarkBtn.addEventListener('click', async () => {
            runDnsBenchmarkBtn.disabled = true;
            dnsBenchmarkStatus.textContent = '⏳ Ejecutando benchmark...';
            dnsTable.style.display = 'none';
            
            try {
                const response = await fetch('/api/dns-benchmark/run', { method: 'POST' });
                const result = await response.json();
                
                if (result.success) {
                    dnsBenchmarkStatus.textContent = '✅ Benchmark iniciado. Actualizando resultados...';
                    setTimeout(updateDnsBenchmark, 3000);
                }
            } catch (error) {
                dnsBenchmarkStatus.textContent = '❌ Error: ' + error.message;
            } finally {
                runDnsBenchmarkBtn.disabled = false;
            }
        });

        async function updateDnsBenchmark() {
            try {
                const response = await fetch('/api/dns-benchmark');
                const data = await response.json();
                
                if (data.status === 'completed' && data.results && data.results.length > 0) {
                    dnsTableBody.innerHTML = '';
                    
                    data.results.forEach(result => {
                        const row = document.createElement('tr');
                        row.innerHTML = `
                            <td>${result.name}</td>
                            <td>${result.ip}</td>
                            <td>${result.avg_time_ms !== null ? result.avg_time_ms + ' ms' : 'N/A'}</td>
                            <td>${result.success_rate.toFixed(0)}%</td>
                        `;
                        dnsTableBody.appendChild(row);
                    });
                    
                    dnsTable.style.display = 'table';
                    dnsBenchmarkStatus.textContent = '✅ Benchmark completado';
                } else if (data.status === 'running') {
                    dnsBenchmarkStatus.textContent = '⏳ Ejecutando benchmark...';
                    setTimeout(updateDnsBenchmark, 2000);
                } else if (data.status === 'error') {
                    dnsBenchmarkStatus.textContent = '❌ Error: ' + data.error;
                }
            } catch (error) {
                console.error('Error actualizando DNS benchmark:', error);
            }
        }

        // Manage Sites Modal
        const sitesModal = document.getElementById('sitesModal');
        const manageSitesBtn = document.getElementById('manageSitesBtn');
        const closeSitesModal = document.getElementById('closeSitesModal');
        const sitesList = document.getElementById('sitesList');
        const newSiteInput = document.getElementById('newSiteInput');
        const addSiteBtn = document.getElementById('addSiteBtn');
        const saveSitesBtn = document.getElementById('saveSitesBtn');

        function renderSiteList(sites) {
            if (!sitesList) return;
            sitesList.innerHTML = '';
            
            sites.forEach(site => {
                const li = document.createElement('li');
                li.textContent = site;
                li.dataset.site = site;
                
                const deleteBtn = document.createElement('span');
                deleteBtn.className = 'delete-site';
                deleteBtn.innerHTML = '&times;';
                deleteBtn.onclick = (e) => {
                    e.stopPropagation();
                    li.remove();
                };
                
                li.appendChild(deleteBtn);
                sitesList.appendChild(li);
            });
        }

        manageSitesBtn.addEventListener('click', async () => {
            try {
                const response = await fetch('/api/websites');
                const sites = await response.json();
                renderSiteList(sites);
                sitesModal.style.display = 'block';
            } catch (error) {
                alert('Error cargando sitios: ' + error.message);
            }
        });

        addSiteBtn.addEventListener('click', () => {
            const site = newSiteInput.value.trim();
            if (site) {
                const current = Array.from(sitesList.querySelectorAll('li')).map(li => li.dataset.site);
                if (current.includes(site)) {
                    alert('El sitio ya existe');
                    return;
                }
                renderSiteList([site, ...current]);
                newSiteInput.value = '';
            }
        });

        saveSitesBtn.addEventListener('click', async () => {
            const sites = Array.from(sitesList.querySelectorAll('li')).map(li => li.dataset.site);
            
            try {
                saveSitesBtn.disabled = true;
                saveSitesBtn.textContent = 'Guardando...';
                
                const response = await fetch('/api/websites', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ websites: sites })
                });
                
                const result = await response.json();
                
                if (result.success) {
                    alert('✅ Sitios guardados correctamente');
                    sitesModal.style.display = 'none';
                    updatePingResults();
                } else {
                    alert('❌ Error: ' + result.message);
                }
            } catch (error) {
                alert('❌ Error de conexión: ' + error.message);
            } finally {
                saveSitesBtn.disabled = false;
                saveSitesBtn.textContent = '💾 Guardar Cambios';
            }
        });

        closeSitesModal.addEventListener('click', () => {
            sitesModal.style.display = 'none';
        });

        window.addEventListener('click', (e) => {
            if (e.target === sitesModal) {
                sitesModal.style.display = 'none';
            }
        });

        // View History
        const viewHistoryBtn = document.getElementById('viewHistoryBtn');
        viewHistoryBtn.addEventListener('click', () => {
            window.location.href = '/speedtest-history';
        });

        // Sorting
        document.querySelectorAll('th.sortable').forEach(th => {
            th.addEventListener('click', () => {
                const newCol = th.dataset.sort;
                if (sortColumn === newCol) {
                    sortDirection = sortDirection === 'asc' ? 'desc' : 'asc';
                } else {
                    sortColumn = newCol;
                    sortDirection = 'asc';
                }
                updatePingResults();
            });
        });

        
        // ========================================
        // GESTIÓN DE BACKUPS
        // ========================================
        
        const createBackupBtn = document.getElementById('createBackupBtn');
        const manageBackupsBtn = document.getElementById('manageBackupsBtn');
        const clearDataBtn = document.getElementById('clearDataBtn');
        const backupStatus = document.getElementById('backupStatus');
        const backupModal = document.getElementById('backupModal');
        const closeBackupModal = document.getElementById('closeBackupModal');
        const backupsList = document.getElementById('backupsList');
        
        // Crear backup
        if (createBackupBtn) {
            createBackupBtn.addEventListener('click', async () => {
                createBackupBtn.disabled = true;
                createBackupBtn.textContent = '⏳ Creando...';
                
                try {
                    const response = await fetch('/api/backup/create', { method: 'POST' });
                    const result = await response.json();
                    
                    if (result.success) {
                        backupStatus.textContent = `✅ ${result.message}`;
                        backupStatus.style.backgroundColor = '#d4edda';
                        backupStatus.style.color = '#155724';
                    } else {
                        backupStatus.textContent = `❌ Error: ${result.error}`;
                        backupStatus.style.backgroundColor = '#f8d7da';
                        backupStatus.style.color = '#721c24';
                    }
                } catch (error) {
                    backupStatus.textContent = '❌ Error de conexión';
                    backupStatus.style.backgroundColor = '#f8d7da';
                    backupStatus.style.color = '#721c24';
                } finally {
                    createBackupBtn.disabled = false;
                    createBackupBtn.textContent = '📦 Crear Backup';
                    setTimeout(() => { backupStatus.textContent = ''; backupStatus.style.backgroundColor = ''; }, 5000);
                }
            });
        }
        
        // Gestionar backups
        if (manageBackupsBtn) {
            manageBackupsBtn.addEventListener('click', async () => {
                backupModal.style.display = 'block';
                await loadBackups();
            });
        }
        
        async function loadBackups() {
            backupsList.innerHTML = '<p>Cargando...</p>';
            
            try {
                const response = await fetch('/api/backup/list');
                const result = await response.json();
                
                if (result.backups && result.backups.length > 0) {
                    let html = '<table style="width:100%; border-collapse:collapse;"><thead><tr><th style="padding:10px; text-align:left; background-color:var(--table-header-bg); color:white;">Archivo</th><th style="padding:10px; text-align:left; background-color:var(--table-header-bg); color:white;">Fecha</th><th style="padding:10px; text-align:left; background-color:var(--table-header-bg); color:white;">Tamaño</th><th style="padding:10px; text-align:left; background-color:var(--table-header-bg); color:white;">Acciones</th></tr></thead><tbody>';
                    
                    result.backups.forEach(backup => {
                        html += `<tr style="border-bottom: 1px solid var(--border-color);">
                            <td style="padding:10px;">${backup.filename}</td>
                            <td style="padding:10px;">${backup.date}</td>
                            <td style="padding:10px;">${(backup.size / 1024).toFixed(2)} KB</td>
                            <td style="padding:10px;">
                                <button onclick="restoreBackup('${backup.filename}')" class="btn-success btn-small">↻ Restaurar</button>
                                <button onclick="deleteBackup('${backup.filename}')" class="btn-danger btn-small">🗑️</button>
                            </td>
                        </tr>`;
                    });
                    
                    html += '</tbody></table>';
                    backupsList.innerHTML = html;
                } else {
                    backupsList.innerHTML = '<p style="padding:20px; text-align:center;">No hay backups disponibles</p>';
                }
            } catch (error) {
                backupsList.innerHTML = '<p style="padding:20px; text-align:center; color:#e74c3c;">Error al cargar backups</p>';
            }
        }
        
        async function restoreBackup(filename) {
            if (!confirm(`¿Restaurar el backup ${filename}?\\n\\nEsto reemplazará los datos actuales.`)) return;
            
            try {
                const response = await fetch(`/api/backup/restore/${filename}`, { method: 'POST' });
                const result = await response.json();
                
                if (result.success) {
                    alert(`✅ ${result.message}\\n\\nRecarga la página para ver los cambios.`);
                    backupModal.style.display = 'none';
                    setTimeout(() => location.reload(), 1000);
                } else {
                    alert(`❌ Error: ${result.error}`);
                }
            } catch (error) {
                alert('❌ Error de conexión');
            }
        }
        
        async function deleteBackup(filename) {
            if (!confirm(`¿Eliminar el backup ${filename}?\\n\\nEsta acción no se puede deshacer.`)) return;
            
            try {
                const response = await fetch(`/api/backup/delete/${filename}`, { method: 'DELETE' });
                const result = await response.json();
                
                if (result.success) {
                    await loadBackups();
                } else {
                    alert(`❌ Error: ${result.error}`);
                }
            } catch (error) {
                alert('❌ Error de conexión');
            }
        }
        
        // Limpiar datos
        if (clearDataBtn) {
            clearDataBtn.addEventListener('click', async () => {
                if (!confirm('⚠️ ¿Estás seguro de limpiar TODOS los datos de monitoreo?\\n\\nSe creará un backup automático antes de limpiar.')) return;
                
                clearDataBtn.disabled = true;
                clearDataBtn.textContent = '⏳ Limpiando...';
                
                try {
                    const response = await fetch('/api/data/clear', { method: 'POST' });
                    const result = await response.json();
                    
                    if (result.success) {
                        alert(`✅ ${result.message}\\n\\nBackup guardado como: ${result.backup}`);
                        setTimeout(() => location.reload(), 1000);
                    } else {
                        alert(`❌ Error: ${result.error}`);
                    }
                } catch (error) {
                    alert('❌ Error de conexión');
                } finally {
                    clearDataBtn.disabled = false;
                    clearDataBtn.textContent = '🗑️ Limpiar Datos';
                }
            });
        }
        
        // Cerrar modal de backups
        if (closeBackupModal) {
            closeBackupModal.onclick = () => { backupModal.style.display = 'none'; };
        }
        
        window.addEventListener('click', (e) => {
            if (e.target === backupModal) {
                backupModal.style.display = 'none';
            }
        });
        
        // ========================================
        // CONFIGURACIÓN SPEEDTEST
        // ========================================
        
        const autoSpeedtestConfig = document.getElementById('autoSpeedtestConfig');
        const speedtestInterval = document.getElementById('speedtestInterval');
        const saveIntervalBtn = document.getElementById('saveIntervalBtn');
        const intervalStatus = document.getElementById('intervalStatus');
        
        // Cargar intervalo actual al inicio
        async function loadCurrentInterval() {
            try {
                const response = await fetch('/api/get-speedtest-interval');
                const result = await response.json();
                
                if (result.success && speedtestInterval) {
                    speedtestInterval.value = result.interval_minutes;
                    console.log(`Intervalo cargado: ${result.interval_minutes} minutos`);
                    
                    // Si el auto está habilitado, mostrar la configuración
                    if (result.auto_enabled) {
                        autoSpeedtestEnabled = true;
                        const toggleBtn = document.getElementById('toggleAutoSpeedtest');
                        if (toggleBtn) {
                            toggleBtn.textContent = '🔄 Auto: ON';
                            toggleBtn.className = 'btn-success';
                        }
                        if (autoSpeedtestConfig) {
                            autoSpeedtestConfig.style.display = 'block';
                        }
                    }
                }
            } catch (error) {
                console.error('Error cargando intervalo:', error);
            }
        }
        
        // Guardar intervalo
        if (saveIntervalBtn) {
            saveIntervalBtn.addEventListener('click', async () => {
                const minutes = parseInt(speedtestInterval.value);
                
                if (isNaN(minutes) || minutes < 1 || minutes > 1440) {
                    alert('⚠️ Intervalo debe estar entre 1 y 1440 minutos');
                    return;
                }
                
                saveIntervalBtn.disabled = true;
                saveIntervalBtn.textContent = '⏳ Guardando...';
                
                try {
                    const response = await fetch('/api/set-speedtest-interval', {
                        method: 'POST',
                        headers: { 'Content-Type': 'application/json' },
                        body: JSON.stringify({ minutes: minutes })
                    });
                    
                    const result = await response.json();
                    
                    if (result.success) {
                        intervalStatus.textContent = `✅ Guardado: ${minutes} min`;
                        intervalStatus.style.color = '#2ecc71';
                        setTimeout(() => { intervalStatus.textContent = ''; }, 3000);
                    } else {
                        alert(`❌ Error: ${result.error}`);
                    }
                } catch (error) {
                    alert('❌ Error de conexión');
                } finally {
                    saveIntervalBtn.disabled = false;
                    saveIntervalBtn.textContent = '💾 Guardar';
                }
            });
        }
        
        // Initialize
        document.addEventListener('DOMContentLoaded', () => {
            updatePingResults();
            updateDashboard();
            updateDnsBenchmark();
            loadCurrentInterval();  // Cargar intervalo de speedtest
            
            setInterval(updatePingResults, PING_INTERVAL_SECONDS * 1000);
            setInterval(() => {
                updateDashboard();
                updateDnsBenchmark();
            }, 5000);
        });
    </script>
</body>
</html>
''')

    with open("templates/history.html", "w", encoding="utf-8") as f:
        f.write(r'''<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Historial - {{ website }}</title>
    <style>
        :root {
            --bg-color: #f4f7f9;
            --text-color: #333;
            --card-bg: #fff;
            --table-header-bg: #34495e;
            --table-row-hover: #f1f3f5;
            --success-color: #2ecc71;
            --danger-color: #e74c3c;
            --warning-color: #f39c12;
            --info-color: #3498db;
            --primary-color: #3498db;
            --border-color: #ddd;
            --link-color: #2980b9;
        }
        [data-theme="dark"] {
            --bg-color: #2c3e50;
            --text-color: #ecf0f1;
            --card-bg: #34495e;
            --table-header-bg: #2c3e50;
            --table-row-hover: #3e5771;
            --success-color: #27ae60;
            --danger-color: #c0392b;
            --warning-color: #e67e22;
            --info-color: #5dade2;
            --primary-color: #5dade2;
            --border-color: #4a627a;
            --link-color: #5dade2;
        }
        body {
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
            background-color: var(--bg-color);
            color: var(--text-color);
            margin: 0;
            padding: 20px;
            transition: background-color 0.3s, color 0.3s;
        }
        h1, h2 {
            text-align: center;
            color: var(--text-color);
        }
        .chart-container, .traceroute-container {
            width: 100%;
            max-width: 1200px;
            margin: 20px auto;
            padding: 20px;
            background-color: var(--card-bg);
            box-shadow: 0 2px 15px rgba(0,0,0,.1);
            border-radius: 8px;
        }
        button {
            background-color: var(--link-color);
            color: #fff;
            border: 0;
            padding: 10px 15px;
            border-radius: 5px;
            cursor: pointer;
            font-weight: 700;
            margin: 5px;
            transition: opacity 0.2s;
        }
        button:hover {
            opacity: .8;
        }
        button:disabled {
            opacity: 0.5;
            cursor: not-allowed;
        }
        .btn-success {
            background-color: var(--success-color);
        }
        .btn-danger {
            background-color: var(--danger-color);
        }
        .traceroute-output {
            background-color: #1e1e1e;
            color: #d4d4d4;
            padding: 15px;
            border-radius: 5px;
            font-family: 'Courier New', monospace;
            font-size: 13px;
            line-height: 1.5;
            max-height: 500px;
            overflow-y: auto;
            white-space: pre-wrap;
            word-wrap: break-word;
            margin-top: 15px;
        }
        .traceroute-controls {
            display: flex;
            gap: 10px;
            align-items: center;
            margin-bottom: 15px;
            flex-wrap: wrap;
        }
        .traceroute-status {
            margin-left: 15px;
            font-style: italic;
            color: var(--text-color);
        }
        .hop-line {
            margin: 5px 0;
            padding: 5px;
            border-left: 3px solid #3498db;
            padding-left: 10px;
        }
        .hop-line:hover {
            background-color: rgba(52, 152, 219, 0.1);
        }
        .loading {
            display: inline-block;
            animation: spin 1s linear infinite;
        }
        @keyframes spin {
            0% { transform: rotate(0deg); }
            100% { transform: rotate(360deg); }
        }
    </style>
    <script src="https://cdn.jsdelivr.net/npm/chart.js"></script>
    <script src="https://cdn.jsdelivr.net/npm/chartjs-adapter-date-fns"></script>
</head>
<body>
    <div style="padding-left: 20px;">
        <button onclick="window.location.href='/'">← Volver al Inicio</button>
    </div>
    
    <h1>📈 Historial de Latencia para: {{ website }}</h1>
    
    <div class="chart-container">
        <canvas id="pingChart"></canvas>
    </div>
    
    <div class="traceroute-container">
        <h2>🛤️ Traceroute / PathPing - Análisis de Ruta</h2>
        <div class="traceroute-controls">
            <div class="method-selection" style="margin-bottom: 15px;">
                <label style="font-weight: bold; margin-right: 20px;">Método de análisis:</label>
                <label style="margin-right: 15px;">
                    <input type="radio" name="traceMethod" value="traceroute" checked style="margin-right: 5px;">
                    🛤️ Traceroute (Rápido)
                </label>
                <label>
                    <input type="radio" name="traceMethod" value="pathping" style="margin-right: 5px;">
                    📊 PathPing (Detallado - Estadísticas de pérdida)
                </label>
            </div>
            <div class="method-description" id="methodDescription" style="background: #f8f9fa; padding: 10px; border-radius: 4px; margin-bottom: 15px; font-size: 14px;">
                <strong>Traceroute:</strong> Muestra la ruta que siguen los paquetes hasta el destino. Ejecución rápida (~30-60 segundos).
            </div>
            <button id="runTracerouteBtn" class="btn-success">▶️ Ejecutar Análisis</button>
            <button id="stopTracerouteBtn" class="btn-danger" style="display:none;" disabled>⏹️ Detener</button>
            <button id="viewHistoryBtn" class="btn-info" style="margin-left: 10px;">📜 Ver Historial</button>
            <button id="clearHistoryBtn" class="btn-danger" style="margin-left: 5px;">🗑️ Limpiar Historial</button>
            <span id="tracerouteStatus" class="traceroute-status"></span>
        </div>
        <div id="tracerouteOutput" class="traceroute-output" style="display:none;">
            Selecciona un método y ejecuta el análisis para ver la ruta de red...
        </div>
    </div>

    <!-- Modal para historial de traceroute -->
    <div id="tracerouteHistoryModal" class="modal" style="display:none;">
        <div class="modal-content" style="max-width: 90%; width: 1000px;">
            <span class="close" id="closeHistoryModal">&times;</span>
            <h2>📜 Historial de Traceroute/PathPing - {{ website }}</h2>

            <div class="history-controls" style="margin-bottom: 15px;">
                <button id="refreshHistoryBtn" class="btn-info">🔄 Actualizar</button>
                <button id="exportHistoryBtn" class="btn-success">💾 Exportar</button>
                <select id="historyFilterMethod" style="margin-left: 10px; padding: 5px;">
                    <option value="">Todos los métodos</option>
                    <option value="traceroute">Solo Traceroute</option>
                    <option value="pathping">Solo PathPing</option>
                </select>
            </div>

            <div id="historyContent" style="max-height: 600px; overflow-y: auto;">
                <p>Cargando historial...</p>
            </div>
        </div>
    </div>
    
    <script>
        const websiteName = "{{ website }}";
        const ctx = document.getElementById("pingChart").getContext("2d");
        let tracerouteRunning = false;
        
        // Función para cargar datos del gráfico
        async function fetchData() {
            try {
                const response = await fetch(`/api/ping-results/${encodeURIComponent(websiteName)}`);
                const data = await response.json();
                const siteData = data[websiteName] || [];
                
                const timestamps = siteData.map(d => new Date(d.timestamp * 1000));
                const latencies = siteData.map(d => d.success ? d.time_ms : null);
                
                new Chart(ctx, {
                    type: "line",
                    data: {
                        labels: timestamps,
                        datasets: [{
                            label: "Latencia (ms)",
                            data: latencies,
                            borderColor: "#3498db",
                            backgroundColor: "rgba(52, 152, 219, 0.1)",
                            fill: true,
                            tension: 0.2
                        }]
                    },
                    options: {
                        responsive: true,
                        scales: {
                            x: {
                                type: "time",
                                time: {
                                    unit: "minute"
                                }
                            },
                            y: {
                                beginAtZero: true,
                                title: {
                                    display: true,
                                    text: 'Latencia (ms)'
                                }
                            }
                        },
                        plugins: {
                            legend: {
                                display: true
                            },
                            tooltip: {
                                mode: 'index',
                                intersect: false
                            }
                        }
                    }
                });
            } catch (error) {
                console.error("Error cargando datos:", error);
            }
        }
        
        // Funciones de traceroute/pathping
        const runTracerouteBtn = document.getElementById('runTracerouteBtn');
        const stopTracerouteBtn = document.getElementById('stopTracerouteBtn');
        const tracerouteStatus = document.getElementById('tracerouteStatus');
        const tracerouteOutput = document.getElementById('tracerouteOutput');
        const methodDescription = document.getElementById('methodDescription');

        // Manejar cambio de método
        const methodRadios = document.querySelectorAll('input[name="traceMethod"]');
        methodRadios.forEach(radio => {
            radio.addEventListener('change', updateMethodDescription);
        });

        function updateMethodDescription() {
            const selectedMethod = document.querySelector('input[name="traceMethod"]:checked').value;
            const isWindows = navigator.platform.toLowerCase().includes('win');

            if (selectedMethod === 'pathping') {
                if (isWindows) {
                    methodDescription.innerHTML = '<strong>PathPing (Windows):</strong> Combina traceroute con estadísticas detalladas de pérdida de paquetes por salto. Proporciona latencia promedio, mínima, máxima y % de pérdida para cada router. Ejecución más lenta (~2-3 minutos).';
                } else {
                    methodDescription.innerHTML = '<strong>MTR (Equivalente a PathPing):</strong> En Linux/Mac usa MTR que proporciona estadísticas similares a PathPing con datos de pérdida de paquetes y latencia por salto. Ejecución moderada (~1-2 minutos).';
                }
            } else {
                methodDescription.innerHTML = '<strong>Traceroute:</strong> Muestra la ruta que siguen los paquetes hasta el destino. Ejecución rápida (~30-60 segundos).';
            }
        }

        async function runTraceroute() {
            if (tracerouteRunning) return;

            const selectedMethod = document.querySelector('input[name="traceMethod"]:checked').value;
            const methodName = selectedMethod === 'pathping' ? 'PathPing' : 'Traceroute';

            tracerouteRunning = true;
            runTracerouteBtn.disabled = true;
            stopTracerouteBtn.style.display = 'inline-block';
            stopTracerouteBtn.disabled = false;

            tracerouteStatus.innerHTML = `<span class="loading">⏳</span> Ejecutando ${methodName}...`;
            tracerouteOutput.style.display = 'block';
            tracerouteOutput.textContent = `Iniciando ${methodName}...\n\n`;

            try {
                const url = `/api/traceroute/${encodeURIComponent(websiteName)}?method=${selectedMethod}`;
                const response = await fetch(url);
                const result = await response.json();

                if (result.output) {
                    // Formatear salida con colores
                    const formattedOutput = formatTracerouteOutput(result.output, selectedMethod);
                    tracerouteOutput.innerHTML = formattedOutput;
                    tracerouteStatus.innerHTML = `✅ ${methodName} completado`;
                    tracerouteStatus.style.color = 'var(--success-color)';
                } else if (result.error) {
                    tracerouteOutput.textContent = `❌ Error: ${result.error}`;
                    tracerouteStatus.innerHTML = '❌ Error al ejecutar';
                    tracerouteStatus.style.color = 'var(--danger-color)';
                }
            } catch (error) {
                tracerouteOutput.textContent = `❌ Error de conexión: ${error.message}`;
                tracerouteStatus.innerHTML = '❌ Error';
                tracerouteStatus.style.color = 'var(--danger-color)';
            } finally {
                tracerouteRunning = false;
                runTracerouteBtn.disabled = false;
                stopTracerouteBtn.style.display = 'none';

                setTimeout(() => {
                    tracerouteStatus.textContent = '';
                }, 5000);
            }
        }
        
        function formatTracerouteOutput(output, method = 'traceroute') {
            // Dividir por líneas y formatear
            const lines = output.split('\n');
            let formatted = '';
            let inStatsSection = false;

            lines.forEach(line => {
                if (line.trim() === '') {
                    formatted += '<br>';
                    return;
                }

                // Detectar secciones especiales
                if (line.includes('=== PATHPING') || line.includes('=== MTR') || line.includes('=== TRACEROUTE')) {
                    formatted += `<div style="background: #e3f2fd; padding: 10px; border-left: 4px solid #2196F3; margin: 10px 0; font-weight: bold; color: #1565C0;">${escapeHtml(line)}</div>`;
                    return;
                }

                // Detectar descripciones de método
                if (line.includes('PathPing combina') || line.includes('MTR proporciona') || line.includes('Muestra la ruta')) {
                    formatted += `<div style="background: #f8f9fa; padding: 8px; border-radius: 4px; margin: 5px 0; font-style: italic; color: #6c757d;">${escapeHtml(line)}</div>`;
                    return;
                }

                // Para PathPing, detectar la sección de estadísticas
                if (method === 'pathping' && (line.includes('Computing statistics') || line.includes('Source to Here') || line.includes('This Node/Link'))) {
                    inStatsSection = true;
                    formatted += `<div style="background: #fff3cd; padding: 10px; border-left: 4px solid #ffc107; margin: 10px 0; font-weight: bold; color: #856404;">${escapeHtml(line)}</div>`;
                    return;
                }

                // Formateo de líneas de salto
                if (line.match(/^\s*\d+/)) {
                    if (method === 'pathping' && inStatsSection) {
                        // Estadísticas de PathPing - resaltar pérdida de paquetes
                        if (line.includes('%')) {
                            const lossMatch = line.match(/(\d+)%/);
                            if (lossMatch) {
                                const lossPercent = parseInt(lossMatch[1]);
                                let color = '#28a745'; // Verde para 0% pérdida
                                let bgColor = 'rgba(40,167,69,0.1)';

                                if (lossPercent > 0 && lossPercent <= 5) {
                                    color = '#ffc107'; // Amarillo para pérdida baja
                                    bgColor = 'rgba(255,193,7,0.1)';
                                } else if (lossPercent > 5 && lossPercent <= 15) {
                                    color = '#fd7e14'; // Naranja para pérdida media
                                    bgColor = 'rgba(253,126,20,0.1)';
                                } else if (lossPercent > 15) {
                                    color = '#dc3545'; // Rojo para pérdida alta
                                    bgColor = 'rgba(220,53,69,0.1)';
                                }

                                formatted += `<div style="color: ${color}; font-weight: bold; background: ${bgColor}; padding: 8px; border-radius: 3px; margin: 2px 0; font-family: monospace; border-left: 3px solid ${color};">${escapeHtml(line)}</div>`;
                            } else {
                                formatted += `<div style="font-family: monospace; padding: 4px; color: var(--text-color);">${escapeHtml(line)}</div>`;
                            }
                        } else {
                            formatted += `<div style="font-family: monospace; padding: 4px; color: var(--text-color);">${escapeHtml(line)}</div>`;
                        }
                    } else {
                        // Líneas de salto normales (traceroute)
                        formatted += `<div class="hop-line" style="color: var(--primary-color); font-weight: bold; font-family: monospace; padding: 4px; background: rgba(0,123,255,0.1); border-radius: 3px; margin: 2px 0;">${escapeHtml(line)}</div>`;
                    }
                } else if (line.includes('ms') && !line.includes('=')) {
                    // Líneas con tiempos de respuesta
                    formatted += `<div style="color: var(--success-color); font-family: monospace; padding: 2px 8px;">${escapeHtml(line)}</div>`;
                } else if (line.includes('*') || line.includes('timeout') || line.includes('Request timed out')) {
                    // Timeouts
                    formatted += `<div style="color: var(--danger-color); font-family: monospace; padding: 2px 8px; background: rgba(220,53,69,0.1);">${escapeHtml(line)}</div>`;
                } else {
                    // Líneas normales
                    formatted += `<div style="font-family: monospace; padding: 2px 8px; color: var(--text-color);">${escapeHtml(line)}</div>`;
                }
            });

            return formatted || escapeHtml(output);
        }
        
        function escapeHtml(text) {
            const div = document.createElement('div');
            div.textContent = text;
            return div.innerHTML;
        }
        
        function stopTraceroute() {
            // Nota: El traceroute del backend ya se habrá completado
            // Este botón es principalmente visual
            stopTracerouteBtn.disabled = true;
            tracerouteStatus.textContent = 'Deteniendo...';
        }
        
        // Funciones para historial de traceroute
        const viewHistoryBtn = document.getElementById('viewHistoryBtn');
        const clearHistoryBtn = document.getElementById('clearHistoryBtn');
        const tracerouteHistoryModal = document.getElementById('tracerouteHistoryModal');
        const closeHistoryModal = document.getElementById('closeHistoryModal');
        const refreshHistoryBtn = document.getElementById('refreshHistoryBtn');
        const exportHistoryBtn = document.getElementById('exportHistoryBtn');
        const historyFilterMethod = document.getElementById('historyFilterMethod');
        const historyContent = document.getElementById('historyContent');

        async function loadTracerouteHistory() {
            try {
                const method = historyFilterMethod.value;
                const url = `/api/traceroute-history?website=${encodeURIComponent(websiteName)}&limit=20${method ? '&method=' + method : ''}`;
                const response = await fetch(url);
                const data = await response.json();

                if (data.history && data.history.length > 0) {
                    let html = '';

                    data.history.forEach((entry, index) => {
                        const date = new Date(entry.timestamp * 1000);
                        const methodIcon = entry.method === 'pathping' ? '📊' : '🛤️';
                        const statusIcon = entry.error ? '❌' : '✅';
                        const statusColor = entry.error ? '#dc3545' : '#28a745';

                        html += `
                            <div class="history-entry" style="border: 1px solid #dee2e6; border-radius: 6px; margin-bottom: 15px; background: white;">
                                <div class="history-header" style="background: #f8f9fa; padding: 12px; border-bottom: 1px solid #dee2e6; cursor: pointer;" onclick="toggleHistoryEntry(${index})">
                                    <div style="display: flex; justify-content: space-between; align-items: center;">
                                        <div>
                                            <strong>${methodIcon} ${entry.method.toUpperCase()}</strong> - ${entry.website}
                                            <span style="color: ${statusColor}; margin-left: 10px;">${statusIcon} ${entry.error ? 'Error' : 'Exitoso'}</span>
                                        </div>
                                        <div style="font-size: 14px; color: #6c757d;">
                                            ${date.toLocaleString()}
                                            <span id="toggleIcon${index}" style="margin-left: 10px;">▼</span>
                                        </div>
                                    </div>
                                </div>
                                <div id="historyEntryContent${index}" class="history-content" style="display: none; padding: 15px;">
                                    <pre style="background: #f8f9fa; padding: 12px; border-radius: 4px; font-size: 12px; overflow-x: auto; white-space: pre-wrap;">${escapeHtml(entry.output)}</pre>
                                </div>
                            </div>
                        `;
                    });

                    historyContent.innerHTML = html;
                } else {
                    historyContent.innerHTML = '<p>No hay historial disponible para este sitio.</p>';
                }
            } catch (error) {
                historyContent.innerHTML = `<p style="color: #dc3545;">Error cargando historial: ${error.message}</p>`;
            }
        }

        window.toggleHistoryEntry = function(index) {
            const content = document.getElementById(`historyEntryContent${index}`);
            const icon = document.getElementById(`toggleIcon${index}`);

            if (content.style.display === 'none') {
                content.style.display = 'block';
                icon.textContent = '▲';
            } else {
                content.style.display = 'none';
                icon.textContent = '▼';
            }
        }

        async function clearTracerouteHistory() {
            if (!confirm('¿Estás seguro de que quieres limpiar todo el historial de traceroute? Esta acción no se puede deshacer.')) {
                return;
            }

            try {
                const response = await fetch('/api/traceroute-history', { method: 'DELETE' });
                const result = await response.json();

                if (result.success) {
                    alert('✅ Historial limpiado correctamente');
                    if (tracerouteHistoryModal.style.display !== 'none') {
                        loadTracerouteHistory();
                    }
                } else {
                    alert('❌ Error limpiando historial: ' + result.error);
                }
            } catch (error) {
                alert('❌ Error de conexión: ' + error.message);
            }
        }

        async function exportTracerouteHistory() {
            try {
                const response = await fetch(`/api/traceroute-history?website=${encodeURIComponent(websiteName)}&limit=100`);
                const data = await response.json();

                if (data.history && data.history.length > 0) {
                    let exportText = `Historial de Traceroute/PathPing - ${websiteName}\n`;
                    exportText += `Exportado el: ${new Date().toLocaleString()}\n`;
                    exportText += '='.repeat(80) + '\n\n';

                    data.history.forEach((entry, index) => {
                        const date = new Date(entry.timestamp * 1000);
                        exportText += `[${index + 1}] ${entry.method.toUpperCase()} - ${date.toLocaleString()}\n`;
                        exportText += '-'.repeat(60) + '\n';
                        exportText += entry.output + '\n\n';
                    });

                    const blob = new Blob([exportText], { type: 'text/plain' });
                    const url = window.URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = `traceroute_history_${websiteName}_${new Date().getTime()}.txt`;
                    document.body.appendChild(a);
                    a.click();
                    window.URL.revokeObjectURL(url);
                    document.body.removeChild(a);
                } else {
                    alert('No hay historial para exportar');
                }
            } catch (error) {
                alert('❌ Error exportando historial: ' + error.message);
            }
        }

        // Event listeners
        runTracerouteBtn.addEventListener('click', runTraceroute);
        stopTracerouteBtn.addEventListener('click', stopTraceroute);
        viewHistoryBtn.addEventListener('click', () => {
            tracerouteHistoryModal.style.display = 'block';
            loadTracerouteHistory();
        });
        clearHistoryBtn.addEventListener('click', clearTracerouteHistory);
        closeHistoryModal.addEventListener('click', () => {
            tracerouteHistoryModal.style.display = 'none';
        });
        refreshHistoryBtn.addEventListener('click', loadTracerouteHistory);
        exportHistoryBtn.addEventListener('click', exportTracerouteHistory);
        historyFilterMethod.addEventListener('change', loadTracerouteHistory);
        
        // Inicialización
        document.addEventListener("DOMContentLoaded", () => {
            const theme = localStorage.getItem("theme");
            if (theme) {
                document.body.dataset.theme = theme;
            }
            fetchData();
        });
    </script>
</body>
</html>
''')

    with open("templates/speedtest_history.html", "w", encoding="utf-8") as f:
        f.write(r'''<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Historial de Pruebas de Velocidad</title>
    <style>
        :root {
            --bg-color: #f4f7f9;
            --text-color: #333;
            --card-bg: #fff;
            --table-header-bg: #34495e;
            --table-row-hover: #f1f3f5;
            --success-color: #2ecc71;
            --danger-color: #e74c3c;
            --warning-color: #f39c12;
            --info-color: #3498db;
            --primary-color: #3498db;
            --border-color: #ddd;
            --table-border: #ddd;
            --link-color: #2980b9;
        }
        [data-theme="dark"] {
            --bg-color: #2c3e50;
            --text-color: #ecf0f1;
            --card-bg: #34495e;
            --table-header-bg: #2c3e50;
            --table-row-hover: #3e5771;
            --success-color: #27ae60;
            --danger-color: #c0392b;
            --warning-color: #e67e22;
            --info-color: #5dade2;
            --primary-color: #5dade2;
            --border-color: #4a627a;
            --table-border: #4a627a;
            --link-color: #5dade2;
        }
        body {
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
            background-color: var(--bg-color);
            color: var(--text-color);
            margin: 0;
            padding: 20px;
        }
        h1 {
            text-align: center;
        }
        table {
            width: 100%;
            max-width: 800px;
            margin: 20px auto;
            border-collapse: collapse;
            background-color: var(--card-bg);
            box-shadow: 0 2px 15px rgba(0,0,0,.1);
        }
        th, td {
            padding: 12px 15px;
            text-align: left;
            border-bottom: 1px solid var(--table-border);
        }
        thead th {
            background-color: var(--table-header-bg);
            color: #fff;
        }
        button {
            background-color: var(--link-color);
            color: #fff;
            border: 0;
            padding: 10px 15px;
            border-radius: 5px;
            cursor: pointer;
            font-weight: 700;
            margin: 10px 0;
        }
        button:hover {
            opacity: .8;
        }
    </style>
</head>
<body>
    <div style="padding-left: 20px;">
        <button onclick="window.location.href='/'">← Volver al Inicio</button>
    </div>
    <h1>🏎️ Historial de Pruebas de Velocidad</h1>
    <table>
        <thead>
            <tr>
                <th>Fecha y Hora</th>
                <th>Bajada (Mbps)</th>
                <th>Subida (Mbps)</th>
                <th>Latencia (ms)</th>
            </tr>
        </thead>
        <tbody id="speedtest-body"></tbody>
    </table>
    
    <script>
        async function loadHistory() {
            try {
                const response = await fetch("/api/speedtest-results");
                const result = await response.json();
                const history = result.data.history;
                const tbody = document.getElementById("speedtest-body");
                
                tbody.innerHTML = "";
                
                history.reverse().forEach(test => {
                    const row = document.createElement("tr");
                    // ✅ CORRECCIÓN: Multiplicar timestamp por 1000
                    const testDate = new Date(test.timestamp * 1000).toLocaleString();
                    
                    row.innerHTML = `
                        <td>${testDate}</td>
                        <td>${test.download_mbps}</td>
                        <td>${test.upload_mbps}</td>
                        <td>${test.ping_ms}</td>
                    `;
                    
                    tbody.appendChild(row);
                });
            } catch (error) {
                console.error("Error:", error);
            }
        }
        
        document.addEventListener("DOMContentLoaded", () => {
            const theme = localStorage.getItem("theme");
            if (theme) {
                document.body.dataset.theme = theme;
            }
            loadHistory();
        });
    </script>
</body>
</html>
''')

    # Crear plantilla de pruebas avanzadas
    with open("templates/advanced_tests.html", "w", encoding="utf-8") as f:
        f.write(r'''<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Pruebas Avanzadas - Monitor de Red</title>
    <style>
        :root {
            --bg-color: #f4f7f9;
            --text-color: #333;
            --card-bg: #fff;
            --table-header-bg: #34495e;
            --table-row-hover: #f1f3f5;
            --success-color: #2ecc71;
            --danger-color: #e74c3c;
            --warning-color: #f39c12;
            --info-color: #3498db;
            --primary-color: #3498db;
            --border-color: #ddd;
            --link-color: #2980b9;
        }
        [data-theme="dark"] {
            --bg-color: #2c3e50;
            --text-color: #ecf0f1;
            --card-bg: #34495e;
            --table-header-bg: #2c3e50;
            --table-row-hover: #3e5771;
            --success-color: #27ae60;
            --danger-color: #c0392b;
            --warning-color: #e67e22;
            --info-color: #5dade2;
            --primary-color: #5dade2;
            --border-color: #4a627a;
            --link-color: #5dade2;
        }

        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }

        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background-color: var(--bg-color);
            min-height: 100vh;
            color: var(--text-color);
            transition: background-color 0.3s, color 0.3s;
        }

        .container {
            max-width: 1400px;
            margin: 0 auto;
            padding: 20px;
        }

        .header {
            background: var(--card-bg);
            border-radius: 12px;
            padding: 20px;
            margin-bottom: 20px;
            box-shadow: 0 4px 20px rgba(0,0,0,0.1);
            display: flex;
            justify-content: space-between;
            align-items: center;
        }

        .header h1 {
            color: var(--primary-color);
            font-size: 2rem;
            font-weight: 700;
        }

        .back-btn {
            background: var(--primary-color);
            color: white;
            border: none;
            border-radius: 8px;
            padding: 12px 24px;
            font-size: 16px;
            cursor: pointer;
            transition: background 0.3s;
            text-decoration: none;
            display: inline-flex;
            align-items: center;
            gap: 8px;
        }

        .back-btn:hover {
            background: #2563eb;
        }

        .tests-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(450px, 1fr));
            gap: 20px;
        }

        .test-card {
            background: var(--card-bg);
            border-radius: 12px;
            padding: 24px;
            box-shadow: 0 4px 20px rgba(0,0,0,0.1);
            transition: transform 0.3s, box-shadow 0.3s;
        }

        .test-card:hover {
            transform: translateY(-2px);
            box-shadow: 0 8px 30px rgba(0,0,0,0.15);
        }

        .test-card h3 {
            color: var(--primary-color);
            font-size: 1.5rem;
            margin-bottom: 12px;
            display: flex;
            align-items: center;
            gap: 10px;
        }

        .test-card p {
            color: #6b7280;
            margin-bottom: 20px;
            line-height: 1.6;
        }

        .test-controls {
            display: flex;
            flex-direction: column;
            gap: 12px;
        }

        .input-group {
            display: flex;
            flex-direction: column;
            gap: 6px;
        }

        .input-group label {
            font-weight: 600;
            color: var(--text-color);
            font-size: 14px;
        }

        .input-group input {
            padding: 10px 12px;
            border: 2px solid var(--border-color);
            border-radius: 8px;
            font-size: 14px;
            transition: border-color 0.3s;
        }

        .input-group input:focus {
            outline: none;
            border-color: var(--primary-color);
        }

        .btn {
            background: var(--primary-color);
            color: white;
            border: none;
            border-radius: 8px;
            padding: 12px 20px;
            font-size: 14px;
            font-weight: 600;
            cursor: pointer;
            transition: all 0.3s;
            display: flex;
            align-items: center;
            justify-content: center;
            gap: 8px;
        }

        .btn:hover {
            background: #2563eb;
            transform: translateY(-1px);
        }

        .btn:disabled {
            background: #9ca3af;
            cursor: not-allowed;
            transform: none;
        }

        .btn-success { background: var(--success-color); }
        .btn-success:hover { background: #059669; }

        .btn-warning { background: var(--warning-color); }
        .btn-warning:hover { background: #d97706; }

        .btn-info { background: var(--info-color); }
        .btn-info:hover { background: #0891b2; }

        .status-indicator {
            padding: 6px 12px;
            border-radius: 20px;
            font-size: 12px;
            font-weight: 600;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            margin-top: 10px;
        }

        .status-idle {
            background: #f3f4f6;
            color: #6b7280;
        }

        .status-running {
            background: #fef3c7;
            color: #d97706;
            animation: pulse 2s infinite;
        }

        .status-completed {
            background: #d1fae5;
            color: #065f46;
        }

        .status-error {
            background: #fee2e2;
            color: #dc2626;
        }

        @keyframes pulse {
            0%, 100% { opacity: 1; }
            50% { opacity: 0.7; }
        }

        .results-section {
            margin-top: 20px;
            padding: 16px;
            background: #f9fafb;
            border-radius: 8px;
            border-left: 4px solid var(--primary-color);
        }

        .results-section h4 {
            color: var(--primary-color);
            margin-bottom: 12px;
            font-size: 16px;
        }

        .tool-warning {
            background: #fff3cd;
            border: 1px solid #ffc107;
            border-radius: 8px;
            padding: 12px 16px;
            margin-top: 12px;
            display: flex;
            align-items: flex-start;
            gap: 10px;
        }

        .tool-warning-icon {
            font-size: 20px;
            flex-shrink: 0;
        }

        .tool-warning-content {
            flex: 1;
        }

        .tool-warning-title {
            font-weight: 600;
            color: #856404;
            margin-bottom: 4px;
        }

        .tool-warning-message {
            color: #856404;
            font-size: 13px;
            line-height: 1.5;
        }

        .tool-warning-command {
            background: #f8f9fa;
            border: 1px solid #dee2e6;
            border-radius: 4px;
            padding: 6px 10px;
            margin-top: 8px;
            font-family: 'Courier New', monospace;
            font-size: 12px;
            color: #333;
            word-break: break-all;
        }

        .test-card.disabled {
            opacity: 0.6;
            pointer-events: none;
        }

        .test-card.disabled h3 {
            color: #95a5a6;
        }

        .results-table {
            width: 100%;
            border-collapse: collapse;
            margin-top: 10px;
            background: white;
            border-radius: 6px;
            overflow: hidden;
            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
        }

        .results-table th,
        .results-table td {
            padding: 10px 12px;
            text-align: left;
            border-bottom: 1px solid var(--border-color);
            font-size: 13px;
        }

        .results-table th {
            background: #f8fafc;
            font-weight: 600;
            color: var(--text-color);
        }

        .results-table td {
            color: #4b5563;
        }

        .error-message {
            background: #fee2e2;
            color: #dc2626;
            padding: 12px;
            border-radius: 6px;
            border-left: 4px solid #dc2626;
            margin-top: 10px;
            font-size: 14px;
        }

        .success-metric {
            color: var(--success-color);
            font-weight: 600;
        }

        .warning-metric {
            color: var(--warning-color);
            font-weight: 600;
        }

        .danger-metric {
            color: var(--danger-color);
            font-weight: 600;
        }

        .loading-spinner {
            display: inline-block;
            width: 16px;
            height: 16px;
            border: 2px solid #f3f3f3;
            border-top: 2px solid var(--primary-color);
            border-radius: 50%;
            animation: spin 1s linear infinite;
        }

        @keyframes spin {
            0% { transform: rotate(0deg); }
            100% { transform: rotate(360deg); }
        }

        .device-grid {
            display: grid;
            grid-template-columns: repeat(auto-fill, minmax(250px, 1fr));
            gap: 12px;
            margin-top: 10px;
        }

        .device-card {
            background: white;
            border: 1px solid var(--border-color);
            border-radius: 6px;
            padding: 12px;
            font-size: 13px;
        }

        .device-card .device-ip {
            font-weight: 600;
            color: var(--primary-color);
            margin-bottom: 4px;
        }

        .device-card .device-info {
            color: #6b7280;
            font-size: 11px;
        }

        @media (max-width: 768px) {
            .tests-grid {
                grid-template-columns: 1fr;
            }

            .container {
                padding: 10px;
            }

            .header {
                flex-direction: column;
                gap: 15px;
                text-align: center;
            }
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>🔬 Pruebas Avanzadas de Red</h1>
            <a href="/" class="back-btn">
                ← Volver al Dashboard
            </a>
        </div>

        <div class="tests-grid">
            <!-- Test 1: DHCP Discovery -->
            <div class="test-card">
                <h3>🔍 Descubrimiento DHCP</h3>
                <p>Detecta todos los servidores DHCP activos en la red local mediante paquetes DHCP Discover.</p>
                <div class="test-controls">
                    <div class="input-group">
                        <label for="dhcpInterface">Interfaz de red:</label>
                        <select id="dhcpInterface" class="form-select">
                            <option value="auto">Cargando interfaces...</option>
                        </select>
                    </div>
                    <button id="btnDhcpDiscovery" class="btn btn-info">
                        <span class="btn-text">Ejecutar Descubrimiento</span>
                    </button>
                    <div id="statusDhcp" class="status-indicator status-idle">Inactivo</div>
                </div>
                <div id="resultsDhcp" class="results-section" style="display: none;">
                    <h4>Servidores DHCP Encontrados</h4>
                    <div id="dhcpContent"></div>
                </div>
            </div>

            <!-- Test 2: Network Discovery -->
            <div class="test-card">
                <h3>🌐 Descubrimiento de Red</h3>
                <p>Escanea la red local para encontrar dispositivos activos usando nmap.</p>
                <div class="test-controls">
                    <div class="input-group">
                        <label for="subnetInput">Subred (opcional):</label>
                        <input type="text" id="subnetInput" placeholder="ej: 192.168.1.0/24" />
                    </div>
                    <button id="btnNetworkDiscovery" class="btn btn-info">
                        <span class="btn-text">Escanear Red</span>
                    </button>
                    <div id="statusNetwork" class="status-indicator status-idle">Inactivo</div>
                </div>
                <div id="resultsNetwork" class="results-section" style="display: none;">
                    <h4>Dispositivos Encontrados</h4>
                    <div id="networkContent"></div>
                </div>
            </div>

            <!-- Test 3: PathPing/MTR Analysis -->
            <div class="test-card">
                <h3>🛣️ Análisis PathPing/Traceroute</h3>
                <p>Análisis de ruta con estadísticas de pérdida y latencia en cada salto usando PathPing (Windows) o MTR (Linux/Mac).</p>
                <div class="test-controls">
                    <div class="input-group">
                        <label for="mtrTarget">Host destino:</label>
                        <input type="text" id="mtrTarget" value="google.com" placeholder="ej: google.com" />
                    </div>
                    <button id="btnMtrAnalysis" class="btn btn-warning">
                        <span class="btn-text">Ejecutar Análisis</span>
                    </button>
                    <div id="statusMtr" class="status-indicator status-idle">Inactivo</div>
                </div>
                <div id="resultsMtr" class="results-section" style="display: none;">
                    <h4>Resultados PathPing/Traceroute</h4>
                    <div id="mtrContent"></div>
                </div>
            </div>

            <!-- Test 4: Netflix Speed -->
            <div class="test-card">
                <h3>🎬 Test Netflix CDN</h3>
                <p>Mide la velocidad de conexión específica a los servidores de Netflix usando fast.com.</p>
                <div class="test-controls">
                    <button id="btnNetflixSpeed" class="btn btn-success">
                        <span class="btn-text">Test Netflix</span>
                    </button>
                    <div id="statusNetflix" class="status-indicator status-idle">Inactivo</div>
                </div>
                <div id="resultsNetflix" class="results-section" style="display: none;">
                    <h4>Resultados Netflix CDN</h4>
                    <div id="netflixContent"></div>
                </div>
            </div>

            <!-- Test 5: UDP Jitter -->
            <div class="test-card">
                <h3>📞 Test Jitter UDP</h3>
                <p>Simula una videollamada midiendo jitter y pérdida de paquetes UDP.</p>
                <div class="test-controls">
                    <div class="input-group">
                        <label for="iperf3Server">Servidor iperf3:</label>
                        <input type="text" id="iperf3Server" value="iperf.scottlinux.com" placeholder="ej: iperf.scottlinux.com" />
                    </div>
                    <div class="input-group">
                        <label for="iperf3Port">Puerto:</label>
                        <input type="number" id="iperf3Port" value="5201" placeholder="5201" />
                    </div>
                    <div class="input-group">
                        <label for="testDuration">Duración (seg):</label>
                        <input type="number" id="testDuration" value="10" placeholder="10" />
                    </div>
                    <button id="btnUdpJitter" class="btn btn-warning">
                        <span class="btn-text">Test Jitter UDP</span>
                    </button>
                    <div id="statusUdp" class="status-indicator status-idle">Inactivo</div>
                </div>
                <div id="resultsUdp" class="results-section" style="display: none;">
                    <h4>Resultados Jitter UDP</h4>
                    <div id="udpContent"></div>
                </div>
            </div>

            <!-- Test 6: CDN DNS -->
            <div class="test-card">
                <h3>🌍 Test DNS CDN</h3>
                <p>Analiza la resolución DNS y latencia de dominios CDN populares. Selecciona un CDN o ingresa un dominio personalizado.</p>
                <div class="test-controls">
                    <div class="input-group">
                        <label for="cdnSelect">CDN Preconfigurado:</label>
                        <select id="cdnSelect" style="padding: 8px; border: 1px solid #ddd; border-radius: 4px; font-size: 14px; width: 100%; margin-bottom: 10px;">
                            <option value="">-- Seleccionar CDN --</option>
                            <option value="googlevideo.com">🎥 YouTube/Google Video</option>
                            <option value="cloudfront.net">☁️ Amazon CloudFront</option>
                            <option value="fastly.com">⚡ Fastly</option>
                            <option value="jsdelivr.net">📦 jsDelivr</option>
                            <option value="unpkg.com">📚 UNPKG</option>
                            <option value="cdnjs.cloudflare.com">🌩️ Cloudflare CDNJS</option>
                            <option value="stackpath.bootstrapcdn.com">🅱️ BootstrapCDN</option>
                            <option value="maxcdn.bootstrapcdn.com">🔧 MaxCDN Bootstrap</option>
                            <option value="ajax.googleapis.com">🔗 Google APIs</option>
                            <option value="fonts.googleapis.com">📝 Google Fonts</option>
                            <option value="code.jquery.com">💎 jQuery CDN</option>
                            <option value="cdn.jsdelivr.net">🚀 jsDelivr Main</option>
                            <option value="unpkg.com">📦 UNPKG Global</option>
                            <option value="cdnjs.com">💾 CDNJS Main</option>
                            <option value="rawgit.com">📄 RawGit</option>
                        </select>
                    </div>
                    <div class="input-group">
                        <label for="cdnCustomDomain">O Dominio Personalizado:</label>
                        <input type="text" id="cdnCustomDomain" placeholder="Ej: example.com, cdn.mysite.com" style="padding: 8px; border: 1px solid #ddd; border-radius: 4px; font-size: 14px; width: 100%;" />
                    </div>
                    <div style="margin: 10px 0; padding: 8px; background: #e3f2fd; border-left: 4px solid #2196F3; border-radius: 4px; font-size: 12px;">
                        💡 <strong>Tip:</strong> Los CDN preconfigurados son servicios populares optimizados para análisis. También puedes probar cualquier dominio personalizado.
                    </div>
                    <button id="btnCdnDns" class="btn btn-info">
                        <span class="btn-text">Test DNS CDN</span>
                    </button>
                    <div id="statusCdn" class="status-indicator status-idle">Inactivo</div>
                </div>
                <div id="resultsCdn" class="results-section" style="display: none;">
                    <h4>Resultados DNS CDN</h4>
                    <div id="cdnContent"></div>
                </div>
            </div>

            <!-- Test 7: MTU Verification -->
            <div class="test-card">
                <h3>📏 Verificación MTU</h3>
                <p>Detecta el MTU óptimo de tu conexión probando diferentes tamaños de paquetes sin fragmentación.</p>
                <div class="test-controls">
                    <div class="input-group">
                        <label for="mtuTarget">Host destino:</label>
                        <input type="text" id="mtuTarget" value="8.8.8.8" placeholder="ej: 8.8.8.8" />
                    </div>
                    <button id="btnMtuVerification" class="btn btn-info">
                        <span class="btn-text">Verificar MTU</span>
                    </button>
                    <div id="statusMtu" class="status-indicator status-idle">Inactivo</div>
                </div>
                <div id="resultsMtu" class="results-section" style="display: none;">
                    <h4>Resultados Verificación MTU</h4>
                    <div id="mtuContent"></div>
                </div>
            </div>

            <!-- Test 8: Sustained Load Test -->
            <div class="test-card">
                <h3>📊 Test de Carga Sostenida</h3>
                <p>Descarga escalonada de archivos (1MB, 10MB, 50MB, 100MB, 250MB, 1GB) para evaluar estabilidad y detectar throttling del ISP.</p>
                <div class="test-controls">
                    <button id="btnSustainedLoad" class="btn btn-warning">
                        <span class="btn-text">Ejecutar Test de Carga (6 archivos)</span>
                    </button>
                    <div id="statusLoad" class="status-indicator status-idle">Inactivo</div>
                </div>
                <div id="resultsLoad" class="results-section" style="display: none;">
                    <h4>Resultados Test de Carga</h4>
                    <div id="loadContent"></div>
                </div>
            </div>
        </div>
    </div>

    <script>
        // Funciones de utilidad
        function updateStatus(testName, status) {
            // Mapear nombres de pruebas a IDs de elementos
            const statusIds = {
                'dhcp_discovery': 'statusDhcp',
                'network_discovery': 'statusNetwork',
                'mtr_analysis': 'statusMtr',
                'netflix_speed': 'statusNetflix',
                'udp_jitter': 'statusUdp',
                'cdn_dns': 'statusCdn',
                'mtu_verification': 'statusMtu',
                'sustained_load': 'statusLoad'
            };

            const statusElement = document.getElementById(statusIds[testName]);
            if (!statusElement) {
                console.error(`No se encontró elemento de estado para: ${testName}, ID esperado: ${statusIds[testName]}`);
                return;
            }

            const statusClasses = ['status-idle', 'status-running', 'status-completed', 'status-error'];

            statusClasses.forEach(cls => statusElement.classList.remove(cls));

            switch(status) {
                case 'running':
                    statusElement.classList.add('status-running');
                    statusElement.textContent = 'Ejecutando...';
                    break;
                case 'completed':
                    statusElement.classList.add('status-completed');
                    statusElement.textContent = 'Completado';
                    break;
                case 'error':
                    statusElement.classList.add('status-error');
                    statusElement.textContent = 'Error';
                    break;
                default:
                    statusElement.classList.add('status-idle');
                    statusElement.textContent = 'Inactivo';
            }
        }

        function updateButtonState(testName, isRunning) {
            // Mapear nombres de pruebas a IDs de botones
            const buttonIds = {
                'dhcp_discovery': 'btnDhcpDiscovery',
                'network_discovery': 'btnNetworkDiscovery',
                'mtr_analysis': 'btnMtrAnalysis',
                'netflix_speed': 'btnNetflixSpeed',
                'udp_jitter': 'btnUdpJitter',
                'cdn_dns': 'btnCdnDns',
                'mtu_verification': 'btnMtuVerification',
                'sustained_load': 'btnSustainedLoad'
            };

            const btn = document.getElementById(buttonIds[testName]);
            if (!btn) {
                console.error(`No se encontró botón para: ${testName}, ID esperado: ${buttonIds[testName]}`);
                return;
            }

            const btnText = btn.querySelector('.btn-text');
            if (!btnText) {
                console.error(`No se encontró texto del botón para: ${testName}`);
                return;
            }

            if (isRunning) {
                btn.disabled = true;
                btnText.innerHTML = '<span class="loading-spinner"></span> Ejecutando...';
            } else {
                btn.disabled = false;
                btnText.innerHTML = btn.getAttribute('data-original-text') || btnText.textContent.replace('Ejecutando...', '').trim();
            }
        }

        function showResults(testName, show = true) {
            // Mapear nombres de pruebas a IDs de resultados
            const resultsIds = {
                'dhcp_discovery': 'resultsDhcp',
                'network_discovery': 'resultsNetwork',
                'mtr_analysis': 'resultsMtr',
                'netflix_speed': 'resultsNetflix',
                'udp_jitter': 'resultsUdp',
                'cdn_dns': 'resultsCdn',
                'mtu_verification': 'resultsMtu',
                'sustained_load': 'resultsLoad'
            };

            const resultsElement = document.getElementById(resultsIds[testName]);
            if (!resultsElement) {
                console.error(`No se encontró elemento de resultados para: ${testName}, ID esperado: ${resultsIds[testName]}`);
                return;
            }

            resultsElement.style.display = show ? 'block' : 'none';
        }

        function formatError(error, help = '') {
            let html = `<div class="error-message" style="background: #f8d7da; border: 1px solid #f5c6cb; padding: 15px; border-radius: 6px;">
                <h4 style="margin: 0 0 10px 0; color: #721c24;">❌ Error</h4>
                <p style="margin: 0; color: #721c24;">${error}</p>`;

            if (help) {
                html += `
                    <div style="background: #fff3cd; border: 1px solid #ffc107; border-radius: 4px; padding: 10px; margin-top: 10px;">
                        <strong style="color: #856404;">💡 Ayuda:</strong>
                        <p style="margin: 5px 0 0 0; color: #856404;">${help}</p>
                    </div>`;
            }

            html += `</div>`;
            return html;
        }

        function getContentDiv(testName) {
            // Mapear nombres de pruebas a IDs de contenido
            const contentIds = {
                'dhcp_discovery': 'dhcpContent',
                'network_discovery': 'networkContent',
                'mtr_analysis': 'mtrContent',
                'netflix_speed': 'netflixContent',
                'udp_jitter': 'udpContent',
                'cdn_dns': 'cdnContent',
                'mtu_verification': 'mtuContent',
                'sustained_load': 'loadContent'
            };

            const contentDiv = document.getElementById(contentIds[testName]);
            if (!contentDiv) {
                console.error(`No se encontró elemento de contenido para: ${testName}, ID esperado: ${contentIds[testName]}`);
            }
            return contentDiv;
        }

        // Función genérica para ejecutar pruebas
        async function runTest(testName, endpoint, postData = {}) {
            console.log(`Ejecutando prueba: ${testName}, endpoint: ${endpoint}`);
            updateStatus(testName, 'running');
            updateButtonState(testName, true);

            try {
                console.log(`Enviando request a: ${endpoint}`);
                const response = await fetch(endpoint, {
                    method: 'POST',
                    headers: {
                        'Content-Type': 'application/json'
                    },
                    body: JSON.stringify(postData)
                });

                console.log(`Response status: ${response.status}`);
                const result = await response.json();
                console.log(`Response result:`, result);

                if (result.success) {
                    // Polling para obtener resultados
                    pollForResults(testName);
                } else {
                    updateStatus(testName, 'error');
                    const contentDiv = getContentDiv(testName);
                    if (contentDiv) {
                        contentDiv.innerHTML = formatError(result.error || 'Error desconocido');
                        showResults(testName);
                    }
                }
            } catch (error) {
                console.error('Error en runTest:', error);
                updateStatus(testName, 'error');
                const contentDiv = getContentDiv(testName);
                if (contentDiv) {
                    contentDiv.innerHTML = formatError(`Error de conexión: ${error.message}`);
                    showResults(testName);
                }
            } finally {
                updateButtonState(testName, false);
            }
        }

        // Polling para obtener resultados
        function pollForResults(testName) {
            const poll = async () => {
                try {
                    const response = await fetch(`/api/advanced-tests/results/${testName}`);
                    const data = await response.json();

                    if (data.status === 'completed') {
                        updateStatus(testName, 'completed');
                        displayResults(testName, data.results);
                        showResults(testName);
                        return;
                    } else if (data.status === 'running') {
                        setTimeout(poll, 2000);
                    } else {
                        updateStatus(testName, 'error');
                        const contentDiv = getContentDiv(testName);
                        if (contentDiv) {
                            contentDiv.innerHTML = formatError('Error en la prueba');
                            showResults(testName);
                        }
                    }
                } catch (error) {
                    updateStatus(testName, 'error');
                    const contentDiv = getContentDiv(testName);
                    if (contentDiv) {
                        contentDiv.innerHTML = formatError(`Error obteniendo resultados: ${error.message}`);
                        showResults(testName);
                    }
                }
            };

            setTimeout(poll, 2000);
        }

        // Mapeo de IDs para contenido de resultados
        const contentIds = {
            'dhcp_discovery': 'dhcpContent',
            'network_discovery': 'networkContent',
            'mtr_analysis': 'mtrContent',
            'netflix_speed': 'netflixContent',
            'udp_jitter': 'udpContent',
            'cdn_dns': 'cdnContent',
            'mtu_verification': 'mtuContent',
            'sustained_load': 'loadContent'
        };

        // Funciones para mostrar resultados específicos
        function displayResults(testName, results) {
            const contentDiv = document.getElementById(contentIds[testName]);

            if (!contentDiv) {
                console.error(`No se encontró el contenedor de resultados para: ${testName}`);
                return;
            }

            if (!results.success) {
                contentDiv.innerHTML = formatError(results.error, results.help || '');
                return;
            }

            // Formatear resultados según el tipo de test
            if (testName === 'cdn_dns') {
                contentDiv.innerHTML = formatCdnDnsResults(results);
            } else if (testName === 'dhcp_discovery') {
                contentDiv.innerHTML = formatDhcpResults(results);
            } else if (testName === 'network_discovery') {
                contentDiv.innerHTML = formatNetworkDiscoveryResults(results);
            } else if (testName === 'mtu_verification') {
                contentDiv.innerHTML = formatMtuResults(results);
            } else if (testName === 'sustained_load') {
                contentDiv.innerHTML = formatSustainedLoadResults(results);
            } else if (testName === 'mtr_analysis') {
                contentDiv.innerHTML = formatPathPingResults(results);
            } else {
                // Mostrar resultados básicos como JSON para otros tests
                contentDiv.innerHTML = `<pre style="background: #f8f9fa; padding: 12px; border-radius: 4px; font-size: 12px; overflow-x: auto;">${JSON.stringify(results, null, 2)}</pre>`;
            }
        }

        // Obtener información del CDN basado en el dominio
        function getCdnInfo(domain) {
            const cdnDatabase = {
                'googlevideo.com': { name: 'YouTube/Google Video', icon: '🎥', description: 'Streaming de video y contenido multimedia' },
                'cloudfront.net': { name: 'Amazon CloudFront', icon: '☁️', description: 'CDN global de Amazon Web Services' },
                'fastly.com': { name: 'Fastly', icon: '⚡', description: 'CDN de alta velocidad y edge computing' },
                'jsdelivr.net': { name: 'jsDelivr', icon: '📦', description: 'CDN gratuito para librerías JavaScript' },
                'unpkg.com': { name: 'UNPKG', icon: '📚', description: 'CDN para paquetes npm' },
                'cdnjs.cloudflare.com': { name: 'Cloudflare CDNJS', icon: '🌩️', description: 'Biblioteca de librerías JavaScript' },
                'stackpath.bootstrapcdn.com': { name: 'BootstrapCDN', icon: '🅱️', description: 'CDN oficial de Bootstrap' },
                'maxcdn.bootstrapcdn.com': { name: 'MaxCDN Bootstrap', icon: '🔧', description: 'CDN alternativo para Bootstrap' },
                'ajax.googleapis.com': { name: 'Google APIs', icon: '🔗', description: 'APIs y librerías de Google' },
                'fonts.googleapis.com': { name: 'Google Fonts', icon: '📝', description: 'Fuentes web de Google' },
                'code.jquery.com': { name: 'jQuery CDN', icon: '💎', description: 'CDN oficial de jQuery' },
                'cdn.jsdelivr.net': { name: 'jsDelivr Main', icon: '🚀', description: 'CDN principal de jsDelivr' },
                'cdnjs.com': { name: 'CDNJS Main', icon: '💾', description: 'Repositorio principal de CDNJS' },
                'rawgit.com': { name: 'RawGit', icon: '📄', description: 'CDN para archivos de GitHub' }
            };

            return cdnDatabase[domain] || {
                name: 'Dominio personalizado',
                icon: '🌐',
                description: 'Análisis de dominio personalizado'
            };
        }

        // Análisis de rendimiento DNS
        function getPerformanceAnalysis(dnsTime) {
            let analysis = '';
            let color = '';
            let icon = '';

            if (dnsTime <= 20) {
                analysis = 'Excelente - Tiempo de resolución muy rápido';
                color = '#28a745';
                icon = '🚀';
            } else if (dnsTime <= 50) {
                analysis = 'Bueno - Tiempo de resolución aceptable';
                color = '#28a745';
                icon = '✅';
            } else if (dnsTime <= 100) {
                analysis = 'Regular - Tiempo de resolución moderado';
                color = '#ffc107';
                icon = '⚠️';
            } else {
                analysis = 'Lento - Tiempo de resolución elevado';
                color = '#dc3545';
                icon = '🐌';
            }

            return `<p style="color: ${color}; font-weight: bold;">${icon} ${analysis}</p>`;
        }

        // Formatear resultados del test CDN DNS
        function formatCdnDnsResults(results) {
            // Identificar el tipo de CDN basado en el dominio
            const cdnInfo = getCdnInfo(results.domain);

            let html = `
                <div style="background: #f8f9fa; padding: 15px; border-radius: 6px; margin-bottom: 15px;">
                    <h4 style="margin: 0 0 10px 0; color: #28a745;">✅ Test CDN DNS Completado</h4>
                    <p><strong>Dominio:</strong> ${results.domain}</p>
                    <p><strong>Proveedor:</strong> ${cdnInfo.name} ${cdnInfo.icon}</p>
                    <p><strong>Tipo de servicio:</strong> ${cdnInfo.description}</p>
                    <p><strong>Tiempo DNS:</strong> ${results.dns_resolution_time.toFixed(2)} ms</p>
                    <p><strong>IPs resueltas:</strong> ${results.resolved_ips.length}</p>
                    ${getPerformanceAnalysis(results.dns_resolution_time)}
                </div>

                <div style="background: white; border: 1px solid #dee2e6; border-radius: 6px; overflow: hidden;">
                    <div style="background: #e9ecef; padding: 12px; font-weight: bold; border-bottom: 1px solid #dee2e6;">
                        📊 Resultados de Latencia
                    </div>
                    <div style="overflow-x: auto;">
                        <table style="width: 100%; border-collapse: collapse;">
                            <thead>
                                <tr style="background: #f8f9fa;">
                                    <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">IP</th>
                                    <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Método</th>
                                    <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Estado</th>
                                    <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Latencia Promedio</th>
                                    <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Min/Max</th>
                                </tr>
                            </thead>
                            <tbody>
            `;

            results.ping_results.forEach(ping => {
                const statusIcon = ping.success ? '✅' : '❌';
                const statusColor = ping.success ? '#28a745' : '#dc3545';
                const latencyText = ping.success ? `${ping.avg_latency.toFixed(2)} ms` : 'N/A';
                const minMaxText = ping.success ? `${ping.min_latency.toFixed(2)} / ${ping.max_latency.toFixed(2)} ms` : 'N/A';

                html += `
                    <tr>
                        <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;"><code>${ping.ip}</code></td>
                        <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;">${ping.method}</td>
                        <td style="padding: 10px; border-bottom: 1px solid #f1f3f4; color: ${statusColor};">${statusIcon} ${ping.success ? 'Exitoso' : 'Falló'}</td>
                        <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;">${latencyText}</td>
                        <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;">${minMaxText}</td>
                    </tr>
                `;
            });

            html += `
                            </tbody>
                        </table>
                    </div>
                </div>
            `;

            return html;
        }

        // Formatear resultados DHCP Discovery
        function formatDhcpResults(results) {
            if (results.dhcp_servers && results.dhcp_servers.length > 0) {
                let html = `
                    <div style="background: #f8f9fa; padding: 15px; border-radius: 6px; margin-bottom: 15px;">
                        <h4 style="margin: 0 0 10px 0; color: #28a745;">✅ Servidores DHCP Encontrados: ${results.total_found}</h4>
                    </div>
                `;

                results.dhcp_servers.forEach(server => {
                    html += `
                        <div style="background: white; border: 1px solid #dee2e6; border-radius: 6px; padding: 15px; margin-bottom: 10px;">
                            <h5 style="margin: 0 0 10px 0;">🖥️ Servidor DHCP</h5>
                            <p><strong>IP del servidor:</strong> <code>${server.server_ip}</code></p>
                            <p><strong>IP ofrecida:</strong> <code>${server.offered_ip}</code></p>
                            <p><strong>Tiempo de respuesta:</strong> ${server.response_time_ms.toFixed(2)} ms</p>
                            <p><strong>Detectado en:</strong> ${new Date(server.timestamp * 1000).toLocaleString()}</p>
                        </div>
                    `;
                });

                return html;
            } else {
                return `
                    <div style="background: #fff3cd; border: 1px solid #ffeaa7; padding: 15px; border-radius: 6px;">
                        <h4 style="margin: 0; color: #856404;">⚠️ No se encontraron servidores DHCP</h4>
                        <p style="margin: 10px 0 0 0;">No se detectaron respuestas DHCP en la red actual.</p>
                    </div>
                `;
            }
        }

        // Formatear resultados Network Discovery
        function formatNetworkDiscoveryResults(results) {
            let html = `
                <div style="background: #f8f9fa; padding: 15px; border-radius: 6px; margin-bottom: 15px;">
                    <h4 style="margin: 0 0 10px 0; color: #28a745;">✅ Escaneo de Red Completado</h4>
                    <p><strong>Subred escaneada:</strong> ${results.subnet_scanned}</p>
                    <p><strong>Dispositivos encontrados:</strong> ${results.devices_found}</p>
                </div>
            `;

            if (results.devices && results.devices.length > 0) {
                html += `
                    <div style="background: white; border: 1px solid #dee2e6; border-radius: 6px; overflow: hidden;">
                        <div style="background: #e9ecef; padding: 12px; font-weight: bold; border-bottom: 1px solid #dee2e6;">
                            🌐 Dispositivos Detectados
                        </div>
                        <div style="overflow-x: auto;">
                            <table style="width: 100%; border-collapse: collapse;">
                                <thead>
                                    <tr style="background: #f8f9fa;">
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">IP</th>
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Estado</th>
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Hostname</th>
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">MAC</th>
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Fabricante</th>
                                    </tr>
                                </thead>
                                <tbody>
                `;

                results.devices.forEach(device => {
                    html += `
                        <tr>
                            <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;"><code>${device.ip}</code></td>
                            <td style="padding: 10px; border-bottom: 1px solid #f1f3f4; color: #28a745;">✅ Activo</td>
                            <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;">${device.hostname || 'N/A'}</td>
                            <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;">${device.mac || 'N/A'}</td>
                            <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;">${device.vendor || 'N/A'}</td>
                        </tr>
                    `;
                });

                html += `
                                </tbody>
                            </table>
                        </div>
                    </div>
                `;
            }

            return html;
        }

        // Formatear resultados de Verificación MTU
        function formatMtuResults(results) {
            let html = `
                <div style="background: #f8f9fa; padding: 15px; border-radius: 6px; margin-bottom: 15px;">
                    <h4 style="margin: 0 0 10px 0; color: #28a745;">✅ Verificación MTU Completada</h4>
                    <p><strong>Host destino:</strong> ${results.target_host}</p>
                    <p><strong>MTU Óptimo detectado:</strong> <code style="background: #fff; padding: 4px 8px; border-radius: 4px;">${results.optimal_mtu || 'No determinado'} bytes</code></p>
                </div>
            `;

            // Mostrar recomendaciones
            if (results.recommendations && results.recommendations.length > 0) {
                html += `
                    <div style="background: #fff3cd; border: 1px solid #ffc107; border-radius: 6px; padding: 15px; margin-bottom: 15px;">
                        <h5 style="margin: 0 0 10px 0;">💡 Recomendaciones</h5>
                `;
                results.recommendations.forEach(rec => {
                    html += `<p style="margin: 5px 0;">• ${rec}</p>`;
                });
                html += `</div>`;
            }

            // Tabla de resultados por tamaño MTU
            if (results.results && results.results.length > 0) {
                html += `
                    <div style="background: white; border: 1px solid #dee2e6; border-radius: 6px; overflow: hidden;">
                        <div style="background: #e9ecef; padding: 12px; font-weight: bold; border-bottom: 1px solid #dee2e6;">
                            📊 Resultados por Tamaño MTU
                        </div>
                        <div style="overflow-x: auto;">
                            <table style="width: 100%; border-collapse: collapse;">
                                <thead>
                                    <tr style="background: #f8f9fa;">
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">MTU (bytes)</th>
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Payload (bytes)</th>
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Estado</th>
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Tiempo</th>
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Resultado</th>
                                    </tr>
                                </thead>
                                <tbody>
                `;

                results.results.forEach(test => {
                    const statusIcon = test.success ? '✅' : '❌';
                    const statusColor = test.success ? '#28a745' : (test.fragmented ? '#ffc107' : '#dc3545');
                    const statusText = test.success ? 'OK' : (test.fragmented ? 'Fragmentado' : 'Falló');
                    const timeText = test.success ? `${test.response_time_ms} ms` : 'N/A';
                    const isOptimal = test.mtu === results.optimal_mtu;

                    html += `
                        <tr style="${isOptimal ? 'background: #d4edda; font-weight: bold;' : ''}">
                            <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;">${test.mtu}${isOptimal ? ' 🎯' : ''}</td>
                            <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;">${test.payload_size}</td>
                            <td style="padding: 10px; border-bottom: 1px solid #f1f3f4; color: ${statusColor};">${statusIcon} ${statusText}</td>
                            <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;">${timeText}</td>
                            <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;"><small>${test.status}</small></td>
                        </tr>
                    `;
                });

                html += `
                                </tbody>
                            </table>
                        </div>
                    </div>
                `;
            }

            return html;
        }

        // Formatear resultados de Test de Carga Sostenida
        function formatSustainedLoadResults(results) {
            let html = `
                <div style="background: #f8f9fa; padding: 15px; border-radius: 6px; margin-bottom: 15px;">
                    <h4 style="margin: 0 0 10px 0; color: #28a745;">✅ Test de Carga Sostenida Completado</h4>
                    <p><strong>Duración total:</strong> ${results.total_duration_seconds.toFixed(1)} segundos</p>
                    <p><strong>Datos descargados:</strong> ${results.total_gb >= 1 ? results.total_gb.toFixed(2) + ' GB' : results.total_mb.toFixed(2) + ' MB'}</p>
                    <p><strong>Archivos probados:</strong> ${results.files_tested} (${results.files_successful} exitosos, ${results.files_failed} fallidos)</p>
                    <p><strong>Velocidad promedio general:</strong> <code style="background: #fff; padding: 4px 8px; border-radius: 4px;">${results.overall_avg_speed_mbps.toFixed(2)} Mbps</code></p>
                    <p><strong>Degradación de velocidad:</strong> ${results.speed_degradation_percent.toFixed(1)}%</p>
                    <p><strong>Estabilidad promedio:</strong> ${results.avg_stability_percent.toFixed(1)}%</p>
                </div>
            `;

            // Tabla de resultados por archivo
            if (results.file_results && results.file_results.length > 0) {
                html += `
                    <div style="background: white; border: 1px solid #dee2e6; border-radius: 6px; overflow: hidden; margin-bottom: 15px;">
                        <div style="background: #e9ecef; padding: 12px; font-weight: bold; border-bottom: 1px solid #dee2e6;">
                            📁 Resultados por Tamaño de Archivo
                        </div>
                        <div style="overflow-x: auto;">
                            <table style="width: 100%; border-collapse: collapse;">
                                <thead>
                                    <tr style="background: #f8f9fa;">
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Tamaño</th>
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Estado</th>
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Duración</th>
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Descargado</th>
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Vel. Promedio</th>
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Vel. Mín/Máx</th>
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Estabilidad</th>
                                    </tr>
                                </thead>
                                <tbody>
                `;

                results.file_results.forEach(file => {
                    const statusIcon = file.success ? '✅' : '❌';
                    const statusColor = file.success ? '#28a745' : '#dc3545';
                    const statusText = file.success ? 'OK' : file.error;

                    if (file.success) {
                        const stabilityColor = file.stability_percent >= 85 ? '#28a745' : (file.stability_percent >= 70 ? '#ffc107' : '#dc3545');

                        html += `
                            <tr>
                                <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;"><strong>${file.size}</strong></td>
                                <td style="padding: 10px; border-bottom: 1px solid #f1f3f4; color: ${statusColor};">${statusIcon} ${statusText}</td>
                                <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;">${file.duration_seconds.toFixed(1)}s</td>
                                <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;">${file.mb_downloaded.toFixed(2)} MB</td>
                                <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;"><strong>${file.avg_speed_mbps.toFixed(2)} Mbps</strong></td>
                                <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;">${file.min_speed_mbps.toFixed(2)} / ${file.max_speed_mbps.toFixed(2)} Mbps</td>
                                <td style="padding: 10px; border-bottom: 1px solid #f1f3f4; color: ${stabilityColor};">${file.stability_percent.toFixed(1)}%</td>
                            </tr>
                        `;
                    } else {
                        html += `
                            <tr>
                                <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;"><strong>${file.size}</strong></td>
                                <td style="padding: 10px; border-bottom: 1px solid #f1f3f4; color: ${statusColor};" colspan="6">${statusIcon} ${statusText}</td>
                            </tr>
                        `;
                    }
                });

                html += `
                                </tbody>
                            </table>
                        </div>
                    </div>
                `;
            }

            // Mostrar recomendaciones
            if (results.recommendations && results.recommendations.length > 0) {
                html += `
                    <div style="background: ${results.recommendations[0].includes('✅') ? '#d4edda' : '#fff3cd'}; border: 1px solid ${results.recommendations[0].includes('✅') ? '#28a745' : '#ffc107'}; border-radius: 6px; padding: 15px; margin-bottom: 15px;">
                        <h5 style="margin: 0 0 10px 0;">💡 Análisis de Rendimiento</h5>
                `;
                results.recommendations.forEach(rec => {
                    html += `<p style="margin: 5px 0;">${rec}</p>`;
                });
                html += `</div>`;
            }

            return html;
        }

        // Formatear resultados de PathPing/MTR
        function formatPathPingResults(results) {
            const method = results.method === 'pathping' ? 'PathPing' : 'MTR';

            let html = `
                <div style="background: #f8f9fa; padding: 15px; border-radius: 6px; margin-bottom: 15px;">
                    <h4 style="margin: 0 0 10px 0; color: #28a745;">✅ Análisis ${method} Completado</h4>
                    <p><strong>Host destino:</strong> ${results.target}</p>
                    <p><strong>Método:</strong> ${method} (${results.method === 'pathping' ? 'Windows' : 'Linux/Mac'})</p>
                    <p><strong>Total de saltos:</strong> ${results.total_hops}</p>
                </div>
            `;

            if (results.hops && results.hops.length > 0) {
                html += `
                    <div style="background: white; border: 1px solid #dee2e6; border-radius: 6px; overflow: hidden;">
                        <div style="background: #e9ecef; padding: 12px; font-weight: bold; border-bottom: 1px solid #dee2e6;">
                            🛣️ Ruta Detallada con Estadísticas
                        </div>
                        <div style="overflow-x: auto;">
                            <table style="width: 100%; border-collapse: collapse;">
                                <thead>
                                    <tr style="background: #f8f9fa;">
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Salto</th>
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Host/IP</th>
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Pérdida</th>
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Enviados</th>
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Mejor</th>
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Promedio</th>
                                        <th style="padding: 10px; text-align: left; border-bottom: 1px solid #dee2e6;">Peor</th>
                                    </tr>
                                </thead>
                                <tbody>
                `;

                results.hops.forEach(hop => {
                    const lossColor = hop.loss_percent === 0 ? '#28a745' : (hop.loss_percent < 10 ? '#ffc107' : '#dc3545');
                    const lossIcon = hop.loss_percent === 0 ? '✅' : (hop.loss_percent < 10 ? '⚠️' : '❌');

                    html += `
                        <tr>
                            <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;"><strong>${hop.hop}</strong></td>
                            <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;"><code>${hop.hostname}</code></td>
                            <td style="padding: 10px; border-bottom: 1px solid #f1f3f4; color: ${lossColor};">
                                ${lossIcon} ${hop.loss_percent}%
                            </td>
                            <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;">${hop.packets_sent || 'N/A'}</td>
                            <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;">${hop.best_ms > 0 ? hop.best_ms + ' ms' : '-'}</td>
                            <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;"><strong>${hop.avg_ms > 0 ? hop.avg_ms + ' ms' : '-'}</strong></td>
                            <td style="padding: 10px; border-bottom: 1px solid #f1f3f4;">${hop.worst_ms > 0 ? hop.worst_ms + ' ms' : '-'}</td>
                        </tr>
                    `;
                });

                html += `
                                </tbody>
                            </table>
                        </div>
                    </div>
                `;

                // Análisis de problemas
                const problemHops = results.hops.filter(h => h.loss_percent > 5);
                if (problemHops.length > 0) {
                    html += `
                        <div style="background: #fff3cd; border: 1px solid #ffc107; border-radius: 6px; padding: 15px; margin-top: 15px;">
                            <h5 style="margin: 0 0 10px 0;">⚠️ Saltos con Pérdida de Paquetes Detectada</h5>
                            <p style="margin: 5px 0;">Se detectaron ${problemHops.length} salto(s) con pérdida > 5%:</p>
                            <ul style="margin: 10px 0; padding-left: 20px;">
                    `;
                    problemHops.forEach(hop => {
                        html += `<li>Salto ${hop.hop} (${hop.hostname}): ${hop.loss_percent}% de pérdida</li>`;
                    });
                    html += `
                            </ul>
                            <p style="margin: 5px 0; font-size: 13px; color: #856404;">
                                💡 Pérdida de paquetes puede indicar congestión o problemas en ese segmento de la ruta.
                            </p>
                        </div>
                    `;
                }
            } else {
                html += `
                    <div style="background: #fff3cd; border: 1px solid #ffc107; border-radius: 6px; padding: 15px;">
                        <h4 style="margin: 0; color: #856404;">⚠️ No se pudieron obtener estadísticas</h4>
                        <p style="margin: 10px 0 0 0;">No se recibieron datos de saltos en la ruta.</p>
                    </div>
                `;
            }

            return html;
        }

        // Event listeners
        document.addEventListener('DOMContentLoaded', function() {
            // Mapeo de nombres de tests a elementos del DOM
            const testElementMap = {
                'dhcp_discovery': {
                    card: document.querySelector('.test-card:nth-child(1)'),
                    button: document.getElementById('btnDhcpDiscovery'),
                    controls: document.querySelector('.test-card:nth-child(1) .test-controls')
                },
                'network_discovery': {
                    card: document.querySelector('.test-card:nth-child(2)'),
                    button: document.getElementById('btnNetworkDiscovery'),
                    controls: document.querySelector('.test-card:nth-child(2) .test-controls')
                },
                'mtr_analysis': {
                    card: document.querySelector('.test-card:nth-child(3)'),
                    button: document.getElementById('btnMtrAnalysis'),
                    controls: document.querySelector('.test-card:nth-child(3) .test-controls')
                },
                'netflix_speed': {
                    card: document.querySelector('.test-card:nth-child(4)'),
                    button: document.getElementById('btnNetflixSpeed'),
                    controls: document.querySelector('.test-card:nth-child(4) .test-controls')
                },
                'udp_jitter': {
                    card: document.querySelector('.test-card:nth-child(5)'),
                    button: document.getElementById('btnUdpJitter'),
                    controls: document.querySelector('.test-card:nth-child(5) .test-controls')
                },
                'cdn_dns': {
                    card: document.querySelector('.test-card:nth-child(6)'),
                    button: document.getElementById('btnCdnDns'),
                    controls: document.querySelector('.test-card:nth-child(6) .test-controls')
                },
                'mtu_verification': {
                    card: document.querySelector('.test-card:nth-child(7)'),
                    button: document.getElementById('btnMtuVerification'),
                    controls: document.querySelector('.test-card:nth-child(7) .test-controls')
                },
                'sustained_load': {
                    card: document.querySelector('.test-card:nth-child(8)'),
                    button: document.getElementById('btnSustainedLoad'),
                    controls: document.querySelector('.test-card:nth-child(8) .test-controls')
                }
            };

            // Verificar disponibilidad de herramientas
            fetch('/api/tools-availability')
                .then(response => response.json())
                .then(data => {
                    if (data.success && data.tools) {
                        Object.keys(data.tools).forEach(testName => {
                            const tool = data.tools[testName];
                            const elements = testElementMap[testName];

                            if (!tool.available && elements && elements.button) {
                                // Deshabilitar el botón
                                elements.button.disabled = true;

                                // Crear mensaje de advertencia
                                const warningDiv = document.createElement('div');
                                warningDiv.className = 'tool-warning';
                                warningDiv.innerHTML = `
                                    <div class="tool-warning-icon">⚠️</div>
                                    <div class="tool-warning-content">
                                        <div class="tool-warning-title">Herramienta no disponible</div>
                                        <div class="tool-warning-message">
                                            Falta instalar: ${tool.missing_tools.join(', ')}
                                        </div>
                                        ${tool.install_command ? `<div class="tool-warning-command">${tool.install_command}</div>` : ''}
                                    </div>
                                `;

                                // Insertar advertencia después de los controles
                                if (elements.controls) {
                                    elements.controls.parentElement.insertBefore(warningDiv, elements.controls.nextSibling);
                                }
                            }
                        });
                    }
                })
                .catch(error => {
                    console.error('Error verificando disponibilidad de herramientas:', error);
                });

            // Cargar interfaces de red disponibles para DHCP
            fetch('/api/network-interfaces')
                .then(response => response.json())
                .then(data => {
                    const interfaceSelect = document.getElementById('dhcpInterface');
                    interfaceSelect.innerHTML = ''; // Limpiar opciones

                    if (data.success && data.interfaces && data.interfaces.length > 0) {
                        data.interfaces.forEach(iface => {
                            const option = document.createElement('option');
                            option.value = iface.name;
                            option.textContent = iface.display_name;
                            interfaceSelect.appendChild(option);
                        });
                    } else {
                        const option = document.createElement('option');
                        option.value = 'auto';
                        option.textContent = 'Auto-detectar interfaz';
                        interfaceSelect.appendChild(option);
                    }
                })
                .catch(error => {
                    console.error('Error cargando interfaces:', error);
                    const interfaceSelect = document.getElementById('dhcpInterface');
                    interfaceSelect.innerHTML = '<option value="auto">Auto-detectar interfaz</option>';
                });

            // DHCP Discovery
            document.getElementById('btnDhcpDiscovery').addEventListener('click', function() {
                console.log('Iniciando test DHCP Discovery...');
                const interfaceName = document.getElementById('dhcpInterface').value;
                runTest('dhcp_discovery', '/api/advanced-tests/dhcp-discovery', { interface: interfaceName });
            });

            // Network Discovery
            document.getElementById('btnNetworkDiscovery').addEventListener('click', function() {
                const subnet = document.getElementById('subnetInput').value.trim();
                runTest('network_discovery', '/api/advanced-tests/network-discovery', { subnet: subnet || null });
            });

            // MTR Analysis
            document.getElementById('btnMtrAnalysis').addEventListener('click', function() {
                const target = document.getElementById('mtrTarget').value.trim();
                if (!target) {
                    alert('Por favor ingresa un host destino');
                    return;
                }
                runTest('mtr_analysis', '/api/advanced-tests/mtr-analysis', { target_host: target });
            });

            // Netflix Speed
            document.getElementById('btnNetflixSpeed').addEventListener('click', function() {
                runTest('netflix_speed', '/api/advanced-tests/netflix-speed');
            });

            // UDP Jitter
            document.getElementById('btnUdpJitter').addEventListener('click', function() {
                const server = document.getElementById('iperf3Server').value.trim();
                const port = parseInt(document.getElementById('iperf3Port').value);
                const duration = parseInt(document.getElementById('testDuration').value);

                if (!server) {
                    alert('Por favor ingresa un servidor iperf3');
                    return;
                }

                runTest('udp_jitter', '/api/advanced-tests/udp-jitter', {
                    server_host: server,
                    server_port: port,
                    duration: duration
                });
            });

            // CDN DNS - Manejo de selección
            const cdnSelect = document.getElementById('cdnSelect');
            const cdnCustomDomain = document.getElementById('cdnCustomDomain');

            // Limpiar input personalizado cuando se selecciona del dropdown
            cdnSelect.addEventListener('change', function() {
                if (this.value) {
                    cdnCustomDomain.value = '';
                }
            });

            // Limpiar dropdown cuando se escribe en el input personalizado
            cdnCustomDomain.addEventListener('input', function() {
                if (this.value.trim()) {
                    cdnSelect.value = '';
                }
            });

            // CDN DNS - Ejecutar test
            document.getElementById('btnCdnDns').addEventListener('click', function() {
                const selectedCdn = cdnSelect.value.trim();
                const customDomain = cdnCustomDomain.value.trim();

                let domain = '';
                let source = '';

                if (selectedCdn) {
                    domain = selectedCdn;
                    // Obtener el texto descriptivo del option seleccionado
                    const selectedOption = cdnSelect.options[cdnSelect.selectedIndex];
                    source = selectedOption.textContent;
                } else if (customDomain) {
                    domain = customDomain;
                    source = 'Dominio personalizado';
                } else {
                    alert('Por favor selecciona un CDN de la lista o ingresa un dominio personalizado');
                    return;
                }

                console.log(`Ejecutando test DNS CDN para: ${domain} (${source})`);
                runTest('cdn_dns', '/api/advanced-tests/cdn-dns', { domain: domain });
            });

            // MTU Verification
            document.getElementById('btnMtuVerification').addEventListener('click', function() {
                const target = document.getElementById('mtuTarget').value.trim();
                if (!target) {
                    alert('Por favor ingresa un host destino');
                    return;
                }
                console.log(`Ejecutando verificación de MTU hacia: ${target}`);
                runTest('mtu_verification', '/api/advanced-tests/mtu-verification', { target_host: target });
            });

            // Sustained Load Test
            document.getElementById('btnSustainedLoad').addEventListener('click', function() {
                console.log('Ejecutando test de carga sostenida con múltiples archivos...');
                runTest('sustained_load', '/api/advanced-tests/sustained-load', {});
            });

            // Guardar textos originales de botones
            document.querySelectorAll('.btn .btn-text').forEach(btnText => {
                btnText.parentElement.setAttribute('data-original-text', btnText.textContent);
            });
        });
    </script>
</body>
</html>
''')


# --- ENDPOINTS DE GESTIÓN DE BACKUPS ---

@app.route("/api/backup/create", methods=["POST"])
def api_create_backup():
    """Crear backup manual de los datos"""
    try:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        backup_filename = f"backup_{timestamp}.json"
        backup_path = os.path.join(BACKUP_DIR, backup_filename)
        
        if not os.path.exists(BACKUP_DIR):
            os.makedirs(BACKUP_DIR)
        
        with data_lock:
            with open(backup_path, 'w', encoding='utf-8') as f:
                json.dump(ping_results_data, f, indent=2)
        
        return jsonify({
            "success": True,
            "message": f"Backup creado: {backup_filename}",
            "filename": backup_filename
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/backup/list", methods=["GET"])
def api_list_backups():
    """Listar todos los backups disponibles"""
    try:
        if not os.path.exists(BACKUP_DIR):
            return jsonify({"backups": []})
        
        backups = []
        for filename in os.listdir(BACKUP_DIR):
            if filename.endswith('.json'):
                filepath = os.path.join(BACKUP_DIR, filename)
                size = os.path.getsize(filepath)
                mtime = os.path.getmtime(filepath)
                backups.append({
                    "filename": filename,
                    "size": size,
                    "date": datetime.fromtimestamp(mtime).strftime('%Y-%m-%d %H:%M:%S')
                })
        
        backups.sort(key=lambda x: x['date'], reverse=True)
        return jsonify({"backups": backups})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/backup/restore/<filename>", methods=["POST"])
def api_restore_backup(filename):
    """Restaurar un backup específico"""
    try:
        backup_path = os.path.join(BACKUP_DIR, filename)
        
        if not os.path.exists(backup_path):
            return jsonify({"success": False, "error": "Backup no encontrado"}), 404
        
        with open(backup_path, 'r', encoding='utf-8') as f:
            backup_data = json.load(f)
        
        with data_lock:
            global ping_results_data
            ping_results_data = backup_data
            
            # Guardar como archivo principal
            with open(JSON_OUTPUT_FILE, 'w', encoding='utf-8') as f:
                json.dump(ping_results_data, f, indent=2)
        
        return jsonify({
            "success": True,
            "message": f"Backup restaurado: {filename}"
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/backup/delete/<filename>", methods=["DELETE"])
def api_delete_backup(filename):
    """Eliminar un backup específico"""
    try:
        backup_path = os.path.join(BACKUP_DIR, filename)
        
        if not os.path.exists(backup_path):
            return jsonify({"success": False, "error": "Backup no encontrado"}), 404
        
        os.remove(backup_path)
        
        return jsonify({
            "success": True,
            "message": f"Backup eliminado: {filename}"
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/data/clear", methods=["POST"])
def api_clear_data():
    """Limpiar todos los datos de monitoreo"""
    try:
        # Crear backup automático antes de limpiar
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        backup_filename = f"backup_before_clear_{timestamp}.json"
        backup_path = os.path.join(BACKUP_DIR, backup_filename)
        
        if not os.path.exists(BACKUP_DIR):
            os.makedirs(BACKUP_DIR)
        
        with data_lock:
            # Backup
            with open(backup_path, 'w', encoding='utf-8') as f:
                json.dump(ping_results_data, f, indent=2)
            
            # Limpiar
            ping_results_data.clear()
            
            # Guardar vacío
            with open(JSON_OUTPUT_FILE, 'w', encoding='utf-8') as f:
                json.dump(ping_results_data, f, indent=2)
        
        return jsonify({
            "success": True,
            "message": "Datos limpiados correctamente",
            "backup": backup_filename
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

# --- ENDPOINTS PARA PRUEBAS AVANZADAS ---

@app.route("/advanced-tests")
def advanced_tests_page():
    """Página de pruebas avanzadas"""
    return render_template("advanced_tests.html")

@app.route("/api/advanced-tests/dhcp-discovery", methods=["POST"])
def api_dhcp_discovery():
    """API para descubrimiento de servidores DHCP"""
    data = request.get_json() or {}
    interface_name = data.get("interface", "auto")

    def run_dhcp_test():
        with advanced_tests_lock:
            advanced_tests_results["dhcp_discovery"]["status"] = "running"

        result = discover_dhcp_servers(interface_name)

        with advanced_tests_lock:
            advanced_tests_results["dhcp_discovery"]["results"] = result
            advanced_tests_results["dhcp_discovery"]["last_run"] = datetime.now().isoformat()
            advanced_tests_results["dhcp_discovery"]["status"] = "completed"

    # Ejecutar en thread separado
    thread = threading.Thread(target=run_dhcp_test, daemon=True)
    thread.start()

    return jsonify({"success": True, "message": f"Iniciando descubrimiento DHCP en interfaz: {interface_name}..."})

@app.route("/api/network-interfaces", methods=["GET"])
def api_get_network_interfaces():
    """API para obtener lista de interfaces de red disponibles"""
    try:
        interfaces = get_network_interfaces()
        return jsonify({"success": True, "interfaces": interfaces})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

@app.route("/api/tools-availability", methods=["GET"])
def api_get_tools_availability():
    """API para obtener el estado de disponibilidad de las herramientas"""
    try:
        tools_status = check_tool_availability()
        return jsonify({"success": True, "tools": tools_status})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

@app.route("/api/advanced-tests/network-discovery", methods=["POST"])
def api_network_discovery():
    """API para descubrimiento de dispositivos en red"""
    data = request.get_json() or {}
    subnet = data.get("subnet")

    def run_network_scan():
        with advanced_tests_lock:
            advanced_tests_results["network_discovery"]["status"] = "running"

        result = discover_network_devices(subnet)

        with advanced_tests_lock:
            advanced_tests_results["network_discovery"]["results"] = result
            advanced_tests_results["network_discovery"]["last_run"] = datetime.now().isoformat()
            advanced_tests_results["network_discovery"]["status"] = "completed"

    # Ejecutar en thread separado
    thread = threading.Thread(target=run_network_scan, daemon=True)
    thread.start()

    return jsonify({"success": True, "message": "Iniciando escaneo de red..."})

@app.route("/api/advanced-tests/mtr-analysis", methods=["POST"])
def api_mtr_analysis():
    """API para análisis MTR"""
    data = request.get_json() or {}
    target_host = data.get("target_host", "google.com")

    def run_mtr_test():
        with advanced_tests_lock:
            advanced_tests_results["mtr_analysis"]["status"] = "running"

        result = run_mtr_analysis(target_host)

        with advanced_tests_lock:
            advanced_tests_results["mtr_analysis"]["results"] = result
            advanced_tests_results["mtr_analysis"]["last_run"] = datetime.now().isoformat()
            advanced_tests_results["mtr_analysis"]["status"] = "completed"

    # Ejecutar en thread separado
    thread = threading.Thread(target=run_mtr_test, daemon=True)
    thread.start()

    return jsonify({"success": True, "message": f"Iniciando análisis MTR para {target_host}..."})

@app.route("/api/advanced-tests/netflix-speed", methods=["POST"])
def api_netflix_speed():
    """API para test de velocidad Netflix"""
    def run_netflix_test():
        with advanced_tests_lock:
            advanced_tests_results["netflix_speed"]["status"] = "running"

        result = run_netflix_speed_test()

        with advanced_tests_lock:
            advanced_tests_results["netflix_speed"]["results"] = result
            advanced_tests_results["netflix_speed"]["last_run"] = datetime.now().isoformat()
            advanced_tests_results["netflix_speed"]["status"] = "completed"

    # Ejecutar en thread separado
    thread = threading.Thread(target=run_netflix_test, daemon=True)
    thread.start()

    return jsonify({"success": True, "message": "Iniciando test de velocidad Netflix..."})

@app.route("/api/advanced-tests/udp-jitter", methods=["POST"])
def api_udp_jitter():
    """API para test de jitter UDP"""
    data = request.get_json() or {}
    server_host = data.get("server_host", "iperf.scottlinux.com")
    server_port = data.get("server_port", 5201)
    duration = data.get("duration", 10)

    def run_udp_test():
        with advanced_tests_lock:
            advanced_tests_results["udp_jitter"]["status"] = "running"

        result = run_udp_jitter_test(server_host, server_port, duration)

        with advanced_tests_lock:
            advanced_tests_results["udp_jitter"]["results"] = result
            advanced_tests_results["udp_jitter"]["last_run"] = datetime.now().isoformat()
            advanced_tests_results["udp_jitter"]["status"] = "completed"

    # Ejecutar en thread separado
    thread = threading.Thread(target=run_udp_test, daemon=True)
    thread.start()

    return jsonify({"success": True, "message": f"Iniciando test UDP con {server_host}..."})

@app.route("/api/advanced-tests/cdn-dns", methods=["POST"])
def api_cdn_dns():
    """API para test DNS de CDN"""
    data = request.get_json() or {}
    domain = data.get("domain", "googlevideo.com")

    def run_cdn_test():
        with advanced_tests_lock:
            advanced_tests_results["cdn_dns"]["status"] = "running"

        result = run_cdn_dns_test(domain)

        with advanced_tests_lock:
            advanced_tests_results["cdn_dns"]["results"] = result
            advanced_tests_results["cdn_dns"]["last_run"] = datetime.now().isoformat()
            advanced_tests_results["cdn_dns"]["status"] = "completed"

    # Ejecutar en thread separado
    thread = threading.Thread(target=run_cdn_test, daemon=True)
    thread.start()

    return jsonify({"success": True, "message": f"Iniciando test DNS para {domain}..."})

@app.route("/api/advanced-tests/mtu-verification", methods=["POST"])
def api_mtu_verification():
    """API para verificación de MTU"""
    data = request.get_json() or {}
    target_host = data.get("target_host", "8.8.8.8")

    def run_mtu_test():
        with advanced_tests_lock:
            advanced_tests_results["mtu_verification"]["status"] = "running"

        result = verify_mtu(target_host)

        with advanced_tests_lock:
            advanced_tests_results["mtu_verification"]["results"] = result
            advanced_tests_results["mtu_verification"]["last_run"] = datetime.now().isoformat()
            advanced_tests_results["mtu_verification"]["status"] = "completed"

    # Ejecutar en thread separado
    thread = threading.Thread(target=run_mtu_test, daemon=True)
    thread.start()

    return jsonify({"success": True, "message": f"Iniciando verificación de MTU hacia {target_host}..."})

@app.route("/api/advanced-tests/sustained-load", methods=["POST"])
def api_sustained_load():
    """API para test de carga sostenida con múltiples archivos"""
    data = request.get_json() or {}
    test_server = data.get("test_server", "http://speedtest.tele2.net")

    def run_load_test():
        with advanced_tests_lock:
            advanced_tests_results["sustained_load"]["status"] = "running"

        result = run_sustained_load_test(test_server=test_server)

        with advanced_tests_lock:
            advanced_tests_results["sustained_load"]["results"] = result
            advanced_tests_results["sustained_load"]["last_run"] = datetime.now().isoformat()
            advanced_tests_results["sustained_load"]["status"] = "completed"

    # Ejecutar en thread separado
    thread = threading.Thread(target=run_load_test, daemon=True)
    thread.start()

    return jsonify({"success": True, "message": "Iniciando test de carga sostenida con 6 archivos (1MB a 1GB)..."})

@app.route("/api/advanced-tests/status")
def api_advanced_tests_status():
    """API para obtener el estado de todas las pruebas avanzadas"""
    with advanced_tests_lock:
        return jsonify(advanced_tests_results)

@app.route("/api/advanced-tests/results/<test_name>")
def api_advanced_test_result(test_name):
    """API para obtener resultados específicos de una prueba"""
    if test_name not in advanced_tests_results:
        return jsonify({"success": False, "error": "Prueba no encontrada"}), 404

    with advanced_tests_lock:
        return jsonify(advanced_tests_results[test_name])

if __name__ == "__main__":
    print("Creando archivos de plantilla si no existen...")
    create_template_files()

    print("Cargando historial de traceroute...")
    load_traceroute_history()

    print("Iniciando hilo de monitoreo principal...")
    main_check_thread = threading.Thread(target=main_check_loop, daemon=True, name="PingCheckLoop")
    main_check_thread.start()

    print("Iniciando hilo de speedtest automático...")
    speedtest_thread = threading.Thread(target=speedtest_loop, daemon=True, name="SpeedtestLoop")
    speedtest_thread.start()

    print("Esperando a que la primera ronda de pings finalice...")
    first_round_completed = first_ping_round_done.wait(timeout=120)
    if not first_round_completed:
         print("ADVERTENCIA: Timeout esperando primera ronda. Servidor iniciará, pero datos iniciales pueden faltar.")
    else:
         print("Primera ronda completada.")

    url = f"http://127.0.0.1:{FLASK_PORT}"
    print(f"🌍 Servidor web iniciado en http://0.0.0.0:{FLASK_PORT}")
    print(f"   Abre tu navegador en {url}")
    
    try:
        can_open_browser = (platform.system() == "Windows" or platform.system() == "Darwin" or ('DISPLAY' in os.environ and os.environ['DISPLAY']))
        if can_open_browser:
            webbrowser.open(url)
        else:
            print("(No se abrirá navegador automáticamente.)")
    except Exception as e:
        print(f"No se pudo abrir navegador: {e}")

    print("Iniciando servidor Waitress...")
    serve(app, host="0.0.0.0", port=FLASK_PORT, threads=16)
